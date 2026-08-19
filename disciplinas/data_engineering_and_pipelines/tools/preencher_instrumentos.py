#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Preenche os templates de AVALIACAO (questoes por unidade, avaliacao final) e o de
ENTREGA DE TRABALHO, reutilizando os helpers de preencher_docx.py. Preserva o
template (só troca orientacao/exemplos por conteudo). Nas questoes de multipla
escolha, a alternativa CORRETA mantem o '*' na frente (renderizado em destaque).

Uso:
    python tools/preencher_instrumentos.py            # tudo
    python tools/preencher_instrumentos.py questoes   # so as 40 Questoes UNI1-4
    python tools/preencher_instrumentos.py avaliacao  # so a avaliacao final
    python tools/preencher_instrumentos.py entrega    # so a entrega de trabalho
"""
import sys
import re
import shutil
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml.ns import qn

TOOLS = Path(__file__).resolve().parent
sys.path.insert(0, str(TOOLS))
import preencher_docx as pf
from preencher_docx import (render_blocks, find_table, build_formula_index,
                            DISCIPLINA, CONTEUDISTA, DESKTOP, PRISTINE, DISC)


# ----------------------------------------------------------------- helpers
def load_pristine(docx_path: Path) -> Document:
    PRISTINE.mkdir(parents=True, exist_ok=True)
    pristine = PRISTINE / docx_path.name
    if not pristine.exists():
        shutil.copy2(docx_path, pristine)
        print("  backup pristine criado:", pristine.name)
    return Document(pristine)


def set_placeholder(doc, prefixos, valor):
    for p in doc.paragraphs:
        low = p.text.strip().lower()
        if any(low.startswith(pre) for pre in prefixos):
            p.add_run(" " + valor)
            return True
    return False


def cut_body_after(doc, pred, keep_matched=True):
    """Remove paragrafos body-level a partir do 1o que casa `pred` (mantendo-o
    se keep_matched). Retorna qtd removida."""
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        if pred(p.text.strip()):
            start = i
            break
    if start is None:
        return 0
    if keep_matched:
        start += 1
    removed = 0
    for p in paras[start:]:
        p._p.getparent().remove(p._p)
        removed += 1
    return removed


def lines_from_heading(md, pat):
    lines = md.split("\n")
    for i, l in enumerate(lines):
        if re.match(pat, l):
            return lines[i:]
    return lines


def parse_h2(md):
    out = []
    title, buf = None, []
    for l in md.split("\n"):
        if re.match(r"^##\s+", l):
            if title is not None:
                out.append((title, buf))
            title = l[2:].strip(); buf = []
        else:
            if title is not None:
                buf.append(l)
    if title is not None:
        out.append((title, buf))
    return out


IS_PROF = lambda t: t.lower().startswith(("professor-conteudista", "conteudista:"))


def unlock_rows_and_breaks(doc):
    for tbl in doc.element.body.iter(qn("w:tbl")):
        for tr in tbl.findall(qn("w:tr")):
            trPr = tr.find(qn("w:trPr"))
            if trPr is not None:
                for tag in ("w:cantSplit", "w:tblHeader"):
                    el = trPr.find(qn(tag))
                    if el is not None:
                        trPr.remove(el)
    for p in doc.element.body.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            for tag in ("w:pageBreakBefore", "w:keepNext", "w:keepLines"):
                el = pPr.find(qn(tag))
                if el is not None:
                    pPr.remove(el)


# ----------------------------------------------------------------- fills
def fill_linear(docx_path: Path, md_path: Path, heading_pat: str, label: str):
    if not docx_path.exists():
        print(f"  ! docx nao encontrado: {docx_path.name}"); return
    doc = load_pristine(docx_path)
    md = md_path.read_text(encoding="utf-8")
    print(f"  renderizando formulas de {label}...")
    formula_pngs = build_formula_index(md)

    set_placeholder(doc, ("disciplina:",), DISCIPLINA)
    set_placeholder(doc, ("professor-conteudista:", "conteudista:"), CONTEUDISTA)
    removed = cut_body_after(doc, IS_PROF, keep_matched=True)

    body_lines = lines_from_heading(md, heading_pat)
    render_blocks(doc, body_lines, formula_pngs, md_path.parent, first=False, alts=True)
    unlock_rows_and_breaks(doc)
    doc.save(docx_path)
    stars = sum(1 for l in body_lines if re.match(r"^\*[a-eA-E]\.\s", l.strip()))
    print(f"  OK -> {docx_path.name}  (orientacao removida: {removed} par.; alternativas corretas c/ '*': {stars})")


def fill_entrega():
    docx_path = DESKTOP / "Instrumentos Avaliativos" / f"TEMPLATE ENTREGA DE TRABALHO - {DISCIPLINA}.docx"
    md_path = DISC / "instrumentos_avaliativos/entrega_trabalho.md"
    if not docx_path.exists():
        print(f"  ! docx nao encontrado: {docx_path.name}"); return
    doc = load_pristine(docx_path)
    md = md_path.read_text(encoding="utf-8")
    formula_pngs = build_formula_index(md)
    sections = parse_h2(md)

    boxmap = [("título", "título"), ("desafio", "desafio"),
              ("fontes de pesquisa", "fonte de pesquisa"),
              ("entregável", "entregável"), ("solução", "solução")]
    filled = []
    for title, body in sections:
        tl = title.lower()
        done = False
        for key, needle in boxmap:
            if key in tl:
                t = find_table(doc, needle)
                if t:
                    render_blocks(t.rows[1].cells[0], body, formula_pngs, md_path.parent, first=True)
                    filled.append(key)
                done = True
                break
        if not done and "roteiro do estudante" in tl:
            cut_body_after(doc, lambda x: x.strip().lower() == "roteiro do estudante", keep_matched=True)
            render_blocks(doc, body, formula_pngs, md_path.parent, first=False)
            filled.append("roteiro")

    unlock_rows_and_breaks(doc)
    doc.save(docx_path)
    print(f"  OK -> {docx_path.name}  (secoes: {', '.join(filled)})")


# ----------------------------------------------------------------- main
def do_questoes():
    print("QUESTOES (40 Questoes - UNI1-4):")
    for u in (1, 2, 3, 4):
        dx = DESKTOP / f"Unidade {u}" / f"40 Questões - UNI{u}_{DISCIPLINA}.docx"
        md = DISC / f"unidade_{u}/questoes_uni{u}.md"
        fill_linear(dx, md, r"^##\s+Quest", f"UNI{u}")


def do_avaliacao():
    print("AVALIACAO FINAL:")
    dx = DESKTOP / "Instrumentos Avaliativos" / f"Avaliação final_(10 discursivas)_{DISCIPLINA}.docx"
    md = DISC / "instrumentos_avaliativos/avaliacao_final.md"
    fill_linear(dx, md, r"^##\s+Quest", "avaliacao final")


def do_entrega():
    print("ENTREGA DE TRABALHO:")
    fill_entrega()


def main():
    what = sys.argv[1] if len(sys.argv) > 1 else "all"
    print("Disciplina:", DISCIPLINA, "| Navegador:", pf.mg.BROWSER or "-")
    print("-" * 60)
    if what in ("all", "questoes"):
        do_questoes()
    if what in ("all", "avaliacao"):
        do_avaliacao()
    if what in ("all", "entrega"):
        do_entrega()
    print("-" * 60)
    print("Pronto. Abra os .docx na Area de Trabalho para conferir.")


if __name__ == "__main__":
    main()
