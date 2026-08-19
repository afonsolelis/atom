#!/usr/bin/env python3
"""Gera e valida os oito DOCX institucionais das Unidades 1 a 4.

Uso:
    PYTHONPATH=/tmp/dse-docx-libs python scripts/generate_institutional_docx.py

O gerador parte dos modelos originais em ``documentos/`` sem modificá-los. Ele
reaproveita capa, imagens, cabeçalho, rodapé, estilos, seção e caixas coloridas,
substitui as instruções do modelo pelo Markdown autoral e conserva, de forma
deliberada, apenas o bloco institucional autorizado da Biblioteca Virtual.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import sys
import tempfile
import unicodedata
from collections import Counter
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Sequence
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "entrega_final" / "docx"
DISCIPLINE = "Distributed Systems Engineering"
AUTHOR = "Afonso Cesar Lelis Brandão"
GENERATION_DATE = datetime(2026, 8, 1, 12, 0, 0, tzinfo=timezone.utc)

UNIT_SPECS = {
    unit: {
        "lessons": list(range((unit - 1) * 4 + 1, unit * 4 + 1)),
        "unit_md": ROOT / f"unidade_{unit}" / f"unidade_{unit}.md",
        "script_md": ROOT / f"unidade_{unit}" / "roteiros_20min.md",
        "questions_md": ROOT / f"unidade_{unit}" / f"questoes_uni{unit}.md",
        "unit_template": ROOT
        / "documentos"
        / f"Unidade {unit}"
        / f"TEMPLATE - Unidade {unit}_nome da disciplina.docx",
        "questions_template": ROOT
        / "documentos"
        / f"Unidade {unit}"
        / f"40 Questões - UNI{unit}_nomedadisciplina.docx",
        "unit_output": OUTPUT_DIR
        / f"Unidade_{unit}_Material_Didatico_Distributed_Systems_Engineering.docx",
        "questions_output": OUTPUT_DIR
        / f"Unidade_{unit}_Questionario_Distributed_Systems_Engineering.docx",
    }
    for unit in range(1, 5)
}

INTERNAL_INSTRUCTION_MARKERS = (
    "tabela para uso exclusivo do(a) coordenador(a)",
    "orientações para produção de material escrito",
    "orientacoes para producao de material escrito",
    "a pessoa conteudista deve utilizar este template",
    "apagar as orientações após leitura",
    "apagar as orientacoes apos leitura",
    "caro coordenador, insira o plano de ensino",
    "insira duas questões para testar",
    "insira duas questoes para testar",
    "elabore uma questão dissertativa",
    "elabore uma questao dissertativa",
    "exemplo de questão:",
    "exemplo de questao:",
    "exemplo de feedback:",
    "título do vídeo (elabore um título",
    "titulo do video (elabore um titulo",
    "elabore um texto provocativo para",
    "questionário unidade 1questões autorais e de banca",
)

PLACEHOLDER_MARKERS = (
    "nome da disciplina: xxxx",
    "conteudista: xxxx",
    "nomedadisciplina",
    "xxxx",
)

CREDENTIAL_PATTERN = re.compile(
    r"(?im)^\s*(?:senha|password|login|usu[aá]rio)\s*[:=]\s*\S+"
)

INLINE_TOKEN_RE = re.compile(
    r"(\[[^\]]+\]\(https?://[^)]+\)|<https?://[^>]+>|\*\*\*.+?\*\*\*|\*\*.+?\*\*|`[^`]+`|\*[^*\n]+?\*)"
)

LATEX_RESIDUAL_RE = re.compile(r"\\[A-Za-z]+|\\\s|\$|[_^]\{")
MARKDOWN_RESIDUAL_RE = re.compile(
    r"\*\*|(?<!\*)\*[^*\n]{1,160}\*(?!\*)|`|\[[^\]]+\]\(https?://"
)

RENDER_FONT_NAME = "Times New Roman"
RENDER_FONT_SIZE = Pt(12)
RENDER_LINE_SPACING = 1.15
RENDER_SPACE_AFTER = Pt(0)


def normalized(text: str) -> str:
    """Normaliza texto para comparação sem revelar seu conteúdo."""

    return " ".join(text.split())


def fold(text: str) -> str:
    """Normaliza caixa e diacríticos para verificações editoriais."""

    decomposed = unicodedata.normalize("NFKD", text.casefold())
    return "".join(ch for ch in decomposed if not unicodedata.combining(ch))


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def body_tables(document: Document) -> list[Table]:
    """Inclui tabelas dentro de controles de conteúdo (SDT)."""

    return [
        Table(element, document._body)
        for element in document.element.body.iter(qn("w:tbl"))
    ]


def table_text(table: Table) -> str:
    return "\n".join(
        paragraph.text
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )


def find_table(document: Document, predicate) -> Table:
    for table in body_tables(document):
        if predicate(normalized(table_text(table))):
            return table
    raise ValueError("Caixa institucional não encontrada no modelo")


def table_clone(document: Document, predicate):
    return deepcopy(find_table(document, predicate)._tbl)


def replace_text_nodes(element, replacements: dict[str, str]) -> None:
    for text_node in element.iter(qn("w:t")):
        value = text_node.text or ""
        for old, new in replacements.items():
            value = value.replace(old, new)
        text_node.text = value


def update_headers(document: Document, unit: int, questions: bool = False) -> None:
    replacements = {
        "TEMPLATE ÁTOMO 3.0": "MATERIAL DIDÁTICO ÁTOMO 3.0",
        "TEMPLATE ATOMO 3.0": "MATERIAL DIDÁTICO ÁTOMO 3.0",
    }
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            replace_text_nodes(paragraph._p, replacements)
        for table in section.header.tables:
            replace_text_nodes(table._tbl, replacements)


def preserve_cover_elements(document: Document, questions: bool) -> list:
    """Mantém os elementos da capa anteriores à tabela de validação interna."""

    selected = []
    for child in document.element.body:
        if child.tag in {qn("w:tbl"), qn("w:sdt"), qn("w:sectPr")}:
            break
        selected.append(deepcopy(child))
    if questions and not selected:
        # Os questionários institucionais trazem os logos num parágrafo vazio.
        first = next(iter(document.element.body), None)
        if first is not None and first.tag == qn("w:p"):
            selected.append(deepcopy(first))
    return selected


def clear_document_body(document: Document) -> None:
    body = document.element.body
    for child in list(body):
        body.remove(child)


def set_core_properties(document: Document, title: str, subject: str) -> None:
    props = document.core_properties
    props.title = title
    props.subject = subject
    props.author = "UniFECAF EAD"
    props.last_modified_by = "UniFECAF EAD"
    props.comments = ""
    props.keywords = "sistemas distribuídos; material didático; UniFECAF"
    props.created = GENERATION_DATE
    props.modified = GENERATION_DATE


def set_render_profile(document: Document, font_name: str) -> None:
    """Define o perfil tipográfico do conteúdo inserido neste documento."""

    global RENDER_FONT_NAME
    RENDER_FONT_NAME = font_name
    for style in document.styles:
        if not style.name or style.name.casefold() != "normal":
            continue
        style.font.name = font_name
        style.font.size = RENDER_FONT_SIZE
        run_properties = style.element.get_or_add_rPr()
        fonts = run_properties.rFonts
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            run_properties.insert(0, fonts)
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), font_name)


def format_run(run, *, size: Pt | None = None, font_name: str | None = None) -> None:
    selected_font = font_name or RENDER_FONT_NAME
    run.font.name = selected_font
    run.font.size = size or RENDER_FONT_SIZE
    run_properties = run._r.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), selected_font)


def format_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = RENDER_LINE_SPACING
    paragraph.paragraph_format.space_after = RENDER_SPACE_AFTER
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_cover_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    format_run(run, size=Pt(18))
    run.bold = True

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    format_run(run, size=Pt(14))
    run.bold = True

    for label, value in (
        ("Disciplina", DISCIPLINE),
        ("Professor-conteudista", AUTHOR),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lead = paragraph.add_run(f"{label}: ")
        format_run(lead)
        lead.bold = True
        value_run = paragraph.add_run(value)
        format_run(value_run)


def add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = paragraph.add_run(label)
    format_run(run)
    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    run.font.underline = True
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def add_emphasis_runs(
    paragraph, text: str, *, bold: bool = False, italic: bool = False
) -> None:
    """Adiciona texto já delimitado, preservando ênfase simples aninhada."""

    position = 0
    nested = re.compile(r"\*([^*\n]+)\*|`([^`\n]+)`")
    for match in nested.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            format_run(run)
            run.bold = bold
            run.italic = italic
        value = match.group(1) if match.group(1) is not None else match.group(2)
        run = paragraph.add_run(value)
        format_run(run)
        run.bold = bold
        run.italic = italic or match.group(1) is not None
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        format_run(run)
        run.bold = bold
        run.italic = italic


def add_inline(paragraph, text: str) -> None:
    """Renderiza a marcação inline usada pelo material sem perder o * do gabarito."""

    # O Markdown autoral usa delimitadores LaTeX para fórmulas. O Word não
    # interpreta essa sintaxe quando ela é inserida como texto comum; portanto,
    # convertemos cada expressão para uma notação Unicode legível antes de
    # processar negrito, itálico e hyperlinks. A conversão antecipada também
    # alcança fórmulas que aparecem dentro de uma indicação editorial em itálico.
    text = re.sub(r"\$([^$\n]+)\$", lambda match: clean_formula(match.group(1)), text)

    if re.match(r"^\*[a-e]\.", text):
        run = paragraph.add_run("*")
        format_run(run)
        text = text[1:]

    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            format_run(run)
        token = match.group(0)
        markdown_link = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
        angle_link = re.fullmatch(r"<(https?://[^>]+)>", token)
        if markdown_link:
            add_hyperlink(paragraph, markdown_link.group(1), markdown_link.group(2))
        elif angle_link:
            add_hyperlink(paragraph, angle_link.group(1), angle_link.group(1))
        elif token.startswith("***"):
            add_emphasis_runs(paragraph, token[3:-3], bold=True, italic=True)
        elif token.startswith("**"):
            add_emphasis_runs(paragraph, token[2:-2], bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            format_run(run)
        elif token.startswith("*"):
            add_emphasis_runs(paragraph, token[1:-1], italic=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        format_run(run)


def add_paragraph(container, text: str = "", style: str | None = None):
    try:
        paragraph = container.add_paragraph(style=style) if style else container.add_paragraph()
    except (KeyError, ValueError):
        paragraph = container.add_paragraph()
    format_paragraph(paragraph)
    if text:
        add_inline(paragraph, text)
    return paragraph


def shade_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def clean_formula(expression: str) -> str:
    """Converte o subconjunto LaTeX das fontes para texto Unicode legível."""

    result = " ".join(expression.split())

    # Resolva primeiro comandos aninhados, para que frações com subscritos
    # (por exemplo, ``\frac{\lambda_{\text{pico}}}{C}``) fiquem simples.
    for _ in range(5):
        updated = re.sub(r"\\(?:text|mathrm)\{([^{}]*)\}", r"\1", result)
        if updated == result:
            break
        result = updated
    result = re.sub(r"_\{([^{}]*)\}", r"_\1", result)
    result = re.sub(r"\^\{([^{}]*)\}", r"^\1", result)
    result = result.replace("{,}", ",")
    for _ in range(5):
        updated = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1) / (\2)", result)
        if updated == result:
            break
        result = updated
    replacements = {
        r"\leftarrow": "←",
        r"\Rightarrow": "⇒",
        r"\rightarrow": "→",
        r"\times": "×",
        r"\cdot": "·",
        r"\approx": "≈",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\lceil": "⌈",
        r"\rceil": "⌉",
        r"\lfloor": "⌊",
        r"\rfloor": "⌋",
        r"\sum": "∑",
        r"\Delta": "Δ",
        r"\delta": "δ",
        r"\lambda": "λ",
        r"\rho": "ρ",
        r"\max": "max",
        r"\min": "min",
        r"\quad": " ",
        r"\left": "",
        r"\right": "",
    }
    for old, new in replacements.items():
        result = result.replace(old, new)
    result = result.replace(r"\%", "%").replace(r"\,", " ").replace(r"\ ", " ")
    result = result.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", result).strip()


def is_markdown_table(lines: Sequence[str], index: int) -> bool:
    if index + 1 >= len(lines) or "|" not in lines[index]:
        return False
    return bool(re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1]))


def split_table_row(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def add_markdown_table(container, rows: list[list[str]]) -> None:
    column_count = max(len(row) for row in rows)
    if isinstance(container, _Cell):
        table = container.add_table(rows=len(rows), cols=column_count)
    else:
        table = container.add_table(rows=len(rows), cols=column_count)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row_index, source_row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            format_paragraph(paragraph)
            value = source_row[column_index] if column_index < len(source_row) else ""
            add_inline(paragraph, value)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def render_markdown(
    container,
    markdown: str,
    *,
    page_break_before_headings: Iterable[str] = (),
) -> None:
    lines = markdown.strip().splitlines()
    page_break_set = {fold(value) for value in page_break_before_headings}
    index = 0
    emitted = False
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip())
                index += 1
            index += 1
            paragraph = add_paragraph(container)
            run = paragraph.add_run("\n".join(code_lines))
            format_run(run)
            shade_paragraph(paragraph, "F2F2F2")
            if language:
                paragraph.paragraph_format.keep_with_next = False
            emitted = True
            continue

        if stripped == "$$":
            formula_lines = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                formula_lines.append(lines[index].strip())
                index += 1
            index += 1
            paragraph = add_paragraph(container)
            run = paragraph.add_run(clean_formula(" ".join(formula_lines)))
            format_run(run)
            run.italic = True
            emitted = True
            continue

        if is_markdown_table(lines, index):
            rows = [split_table_row(lines[index])]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(split_table_row(lines[index]))
                index += 1
            add_markdown_table(container, rows)
            emitted = True
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 6)
            heading_text = heading.group(2).strip()
            if fold(heading_text) in page_break_set and hasattr(container, "add_page_break"):
                container.add_page_break()
            paragraph = add_paragraph(container, heading_text, f"Heading {level}")
            paragraph.paragraph_format.keep_with_next = True
            emitted = True
            index += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip().lstrip(">").strip())
                index += 1
            paragraph = add_paragraph(container, " ".join(quote_lines))
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.1)
            for run in paragraph.runs:
                run.italic = True
            shade_paragraph(paragraph, "F4F7FA")
            emitted = True
            continue

        bullet = re.match(r"^[-+]\s+(.+)$", stripped)
        if bullet:
            paragraph = add_paragraph(container, "• " + bullet.group(1))
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.first_line_indent = Inches(-0.14)
            emitted = True
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            paragraph = add_paragraph(container, stripped)
            paragraph.paragraph_format.left_indent = Inches(0.16)
            paragraph.paragraph_format.first_line_indent = Inches(-0.16)
            emitted = True
            index += 1
            continue

        paragraph = add_paragraph(container, stripped)
        if re.match(r"^(PORQUE|OU)$", stripped):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
        emitted = True
        index += 1

    if isinstance(container, _Cell) and not emitted:
        container.add_paragraph()


def remove_all_cell_content(cell: _Cell) -> None:
    for child in list(cell._tc):
        if child.tag != qn("w:tcPr"):
            cell._tc.remove(child)


def replace_header_text(cell: _Cell, title: str) -> None:
    text_nodes = list(cell._tc.iter(qn("w:t")))
    if text_nodes:
        text_nodes[0].text = title
        for node in text_nodes[1:]:
            node.text = ""
    else:
        cell.text = title
        for run in cell.paragraphs[0].runs:
            run.bold = True


def allow_table_to_split(table_element) -> None:
    for cant_split in list(table_element.iter(qn("w:cantSplit"))):
        parent = cant_split.getparent()
        if parent is not None:
            parent.remove(cant_split)


def trim_table_rows(table: Table, wanted: int) -> None:
    while len(table.rows) > wanted:
        row = table.rows[-1]
        table._tbl.remove(row._tr)


def append_box(
    document: Document,
    template_element,
    title: str,
    markdown: str,
    *,
    rows_to_keep: int = 2,
    preserve_last_row: bool = False,
) -> Table:
    element = deepcopy(template_element)
    allow_table_to_split(element)
    document.element.body.append(element)
    table = Table(element, document._body)
    replace_header_text(table.cell(0, 0), title)
    trim_table_rows(table, rows_to_keep)
    if len(table.rows) < 2:
        raise ValueError(f"Caixa sem linha de conteúdo: {title}")
    remove_all_cell_content(table.cell(1, 0))
    render_markdown(table.cell(1, 0), markdown)
    if preserve_last_row and len(table.rows) >= 3:
        # A última linha é o bloco de acesso institucional autorizado. Não tocar.
        pass
    document.add_paragraph()
    return table


def append_library_box(
    document: Document, template_element, direct_source_markdown: str
) -> Table:
    element = deepcopy(template_element)
    allow_table_to_split(element)
    document.element.body.append(element)
    table = Table(element, document._body)
    replace_header_text(table.cell(0, 0), "DIRETO DA FONTE — BIBLIOTECA VIRTUAL")
    trim_table_rows(table, 3)
    remove_all_cell_content(table.cell(1, 0))
    render_markdown(table.cell(1, 0), direct_source_markdown)
    # table.rows[2] permanece byte-a-byte no XML clonado.
    document.add_paragraph()
    return table


def add_box_page_break(document: Document) -> None:
    add_page_break(document)


def split_unit_source(text: str, lesson_numbers: Sequence[int]):
    title_match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
    if not title_match:
        raise ValueError("Título da unidade ausente")
    unit_title = title_match.group(1).strip()

    lesson_matches = list(
        re.finditer(r"^##\s+Aula\s+(\d+)\s+—\s+(.+)$", text, re.MULTILINE)
    )
    lesson_by_number = {int(match.group(1)): match for match in lesson_matches}
    if set(lesson_by_number) != set(lesson_numbers):
        raise ValueError(
            f"Aulas encontradas {sorted(lesson_by_number)}; esperadas {lesson_numbers}"
        )

    first_lesson = min(match.start() for match in lesson_matches)
    intro = text[title_match.end() : first_lesson].strip()
    tail_match = re.search(
        r"^##\s+Atividades, síntese e material complementar\s*$",
        text,
        re.MULTILINE | re.IGNORECASE,
    )
    tail_start = tail_match.start() if tail_match else len(text)

    lessons = {}
    sorted_matches = sorted(lesson_matches, key=lambda match: match.start())
    for index, match in enumerate(sorted_matches):
        number = int(match.group(1))
        end = (
            sorted_matches[index + 1].start()
            if index + 1 < len(sorted_matches)
            else tail_start
        )
        chunk = text[match.end() : end].strip()
        chunk = re.sub(
            r"(?ms)^###\s+Roteiro da Videoaula.*?(?=^###\s+|\Z)", "", chunk
        ).strip()
        lessons[number] = {"title": match.group(2).strip(), "content": chunk}

    tail = text[tail_match.end() :].strip() if tail_match else ""
    return unit_title, intro, lessons, tail


def split_script_source(text: str, lesson_numbers: Sequence[int]):
    matches = list(
        re.finditer(
            r"^##\s+(?:Roteiro da\s+)?Videoaula\s+(\d+)\s+—\s+(.+)$",
            text,
            re.MULTILINE,
        )
    )
    by_number = {int(match.group(1)): match for match in matches}
    if set(by_number) != set(lesson_numbers):
        raise ValueError(
            f"Videoaulas encontradas {sorted(by_number)}; esperadas {lesson_numbers}"
        )
    scripts = {}
    sorted_matches = sorted(matches, key=lambda match: match.start())
    for index, match in enumerate(sorted_matches):
        end = sorted_matches[index + 1].start() if index + 1 < len(matches) else len(text)
        content = text[match.end() : end].strip()
        content = re.sub(r"(?m)^---\s*$", "", content).strip()
        scripts[int(match.group(1))] = {
            "title": match.group(2).strip().strip("\"“”"),
            "content": content,
        }
    return scripts


def extract_subsection(text: str, heading: str) -> str:
    match = re.search(
        rf"^###\s+{heading}\s*$", text, re.MULTILINE | re.IGNORECASE
    )
    if not match:
        return ""
    next_heading = re.search(r"^###\s+", text[match.end() :], re.MULTILINE)
    end = match.end() + next_heading.start() if next_heading else len(text)
    return text[match.end() : end].strip()


def remove_subsections(text: str, headings: Sequence[str]) -> str:
    result = text
    for heading in headings:
        result = re.sub(
            rf"(?ms)^###\s+{heading}\s*$.*?(?=^###\s+|\Z)", "", result
        )
    return result.strip()


def complementary_entries(material: str) -> dict[str, str]:
    marker = re.compile(
        r"^\*\*(Direto da Fonte|Para Mergulhar|Podcast(?:/vídeo)?|Artigo científico)[^*]*\*\*\s*",
        re.MULTILINE | re.IGNORECASE,
    )
    matches = list(marker.finditer(material))
    entries: dict[str, str] = {}
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(material)
        label = fold(match.group(1))
        value = material[match.end() : end].strip()
        if label.startswith("direto da fonte"):
            entries["direct"] = value
        elif label.startswith("para mergulhar"):
            entries["deeper"] = value
        elif label.startswith("podcast"):
            entries["podcast"] = value
        elif label.startswith("artigo cientifico"):
            entries["article"] = value
    return entries


def spoken_word_counts(script_text: str) -> dict[int, dict[str, int]]:
    scripts = split_script_source(
        script_text,
        [int(value) for value in re.findall(r"^##\s+(?:Roteiro da\s+)?Videoaula\s+(\d+)", script_text, re.MULTILINE)],
    )
    counts = {}
    for number, script in scripts.items():
        content = script["content"]
        start = re.search(r"^###\s+Abertura contextualizada\s*$", content, re.MULTILINE)
        end = re.search(r"^###\s+Indicações de edição", content, re.MULTILINE)
        narrative = content[start.end() : end.start()] if start and end else content
        narrative = re.sub(r"(?m)^\*\[indicação de edição:.*?\]\*\s*$", "", narrative)
        spoken = len(re.findall(r"\b[\wÀ-ÿ][\wÀ-ÿ’'/-]*\b", narrative))
        conservative = content[: end.start()] if end else content
        conservative = re.sub(
            r"(?m)^\*\[indicação de edição:.*?\]\*\s*$", "", conservative
        )
        conservative_count = len(
            re.findall(r"\b[\wÀ-ÿ][\wÀ-ÿ’'/-]*\b", conservative)
        )
        editing_section = re.search(
            r"(?ms)^###\s+Indicações de edição e recursos visuais\s*$"
            r"(.*?)(?=^###\s+|\Z)",
            content,
        )
        editing_timestamps = []
        if editing_section:
            editing_timestamps = [
                int(minutes) * 60 + int(seconds)
                for minutes, seconds in re.findall(
                    r"(?<!\d)(\d{1,2}):([0-5]\d)", editing_section.group(1)
                )
            ]
        if any(
            current < previous
            for previous, current in zip(editing_timestamps, editing_timestamps[1:])
        ):
            raise ValueError(
                f"Timestamps de edição fora de ordem na Videoaula {number}: "
                f"{editing_timestamps}"
            )
        if editing_timestamps and max(editing_timestamps) > 600:
            raise ValueError(
                f"Timestamp acima de 10 minutos na Videoaula {number}: "
                f"{editing_timestamps}"
            )
        counts[number] = {
            "spoken_words": spoken,
            "conservative_words": conservative_count,
            "editing_timestamps_seconds": editing_timestamps,
        }
    return counts


def question_source_metrics(text: str) -> dict:
    split = re.split(r"^##\s+Gabarito e feedbacks\s*$", text, maxsplit=1, flags=re.MULTILINE)
    if len(split) != 2:
        raise ValueError("Seção de gabarito ausente")
    questions, feedbacks = split
    numbered_questions = len(
        re.findall(r"(?m)^(?:[1-9]|[1-3]\d|40)\.\s+", questions)
    )
    labeled_questions = len(
        re.findall(r"(?m)^\*\*Questão\s+(?:[1-9]|[1-3]\d|40)\b", questions)
    )
    bold_numbered_questions = len(
        re.findall(r"(?m)^\*\*(?:[1-9]|[1-3]\d|40)\.\*\*\s+", questions)
    )
    question_count = numbered_questions + labeled_questions + bold_numbered_questions
    alternatives = re.findall(r"(?m)^\*?([a-e])\.\s+", questions)
    correct = re.findall(r"(?m)^\*([a-e])\.\s+", questions)
    feedback_headings = re.findall(r"(?m)^\*\*Questão\s+(\d+)\*\*", feedbacks)
    feedback_items = re.findall(r"(?m)^-\s+([a-e])\.\s+", feedbacks)
    distribution = Counter(correct)
    metrics = {
        "questions": question_count,
        "alternatives": len(alternatives),
        "correct_markers": len(correct),
        "feedback_sections": len(feedback_headings),
        "feedback_items": len(feedback_items),
        "answer_distribution": {letter: distribution[letter] for letter in "abcde"},
    }
    expected = {
        "questions": 40,
        "alternatives": 200,
        "correct_markers": 40,
        "feedback_sections": 40,
        "feedback_items": 200,
        "answer_distribution": {letter: 8 for letter in "abcde"},
    }
    if metrics != expected:
        raise ValueError(f"Banco de questões fora do padrão: {metrics}")
    return metrics


def credential_row_hash(table: Table) -> str:
    if len(table.rows) < 3:
        raise ValueError("Bloco da Biblioteca Virtual incompleto")
    value = normalized(" ".join(cell.text for cell in table.rows[2].cells))
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def build_unit_document(unit: int, spec: dict) -> tuple[Path, dict]:
    unit_text = read_text(spec["unit_md"])
    script_text = read_text(spec["script_md"])
    unit_title, introduction, lessons, tail = split_unit_source(
        unit_text, spec["lessons"]
    )
    scripts = split_script_source(script_text, spec["lessons"])
    counts = spoken_word_counts(script_text)

    document = Document(spec["unit_template"])
    set_render_profile(document, "Times New Roman")
    update_headers(document, unit)
    cover = preserve_cover_elements(document, questions=False)
    section_properties = deepcopy(document.element.body.sectPr)

    text_box_template = table_clone(
        document, lambda text: fold(text).startswith("texto base aula")
    )
    script_box_template = table_clone(
        document, lambda text: fold(text).startswith("roteiro videoaula")
    )
    quiz_box_template = table_clone(
        document, lambda text: fold(text).startswith("quiz nao avaliativo")
    )
    library_box_template = table_clone(
        document, lambda text: "biblioteca virtual" in fold(text)
    )
    library_source_hash = credential_row_hash(
        find_table(document, lambda text: "biblioteca virtual" in fold(text))
    )
    deeper_box_template = table_clone(
        document, lambda text: fold(text).startswith("para mergulhar")
    )
    podcast_box_template = table_clone(
        document, lambda text: fold(text).startswith("podcast")
    )
    article_box_template = table_clone(
        document, lambda text: fold(text).startswith("artigo cientifico")
    )
    try:
        aai_box_template = table_clone(
            document, lambda text: fold(text).startswith("aai")
        )
    except ValueError:
        aai_box_template = None

    clear_document_body(document)
    replacements = {
        "TEMPLATE DE PRODUÇÃO": f"MATERIAL DIDÁTICO — UNIDADE {unit}",
        "Nome da disciplina: XXXX": f"Nome da disciplina: {DISCIPLINE}",
        "Conteudista: XXXX": f"Conteudista: {AUTHOR}",
        "XXXX": DISCIPLINE,
    }
    for element in cover:
        replace_text_nodes(element, replacements)
        document.element.body.append(element)

    add_cover_title(
        document,
        unit_title,
        "Texto-base e roteiros de videoaula",
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    duration_run = paragraph.add_run("Videoaulas dimensionadas para 20 minutos.")
    format_run(duration_run)
    duration_run.italic = True
    add_page_break(document)

    append_box(
        document,
        text_box_template,
        f"APRESENTAÇÃO DA UNIDADE {unit}",
        introduction,
    )

    for lesson_number in spec["lessons"]:
        add_box_page_break(document)
        lesson = lessons[lesson_number]
        append_box(
            document,
            text_box_template,
            f"TEXTO BASE — AULA {lesson_number}: {lesson['title']}",
            lesson["content"],
        )
        add_box_page_break(document)
        script = scripts[lesson_number]
        append_box(
            document,
            script_box_template,
            f"ROTEIRO DA VIDEOAULA {lesson_number}: {script['title']}",
            script["content"],
        )

    quiz = extract_subsection(tail, r"Quiz não avaliativo")
    aai = extract_subsection(tail, r"Atividade Avaliativa Individual \(AAI\)")
    synthesis = remove_subsections(
        tail,
        (
            r"Quiz não avaliativo",
            r"Atividade Avaliativa Individual \(AAI\)",
            r"Material complementar",
        ),
    )
    material = extract_subsection(tail, r"Material complementar")
    entries = complementary_entries(material)

    if synthesis:
        add_box_page_break(document)
        append_box(
            document,
            text_box_template,
            f"SÍNTESE DA UNIDADE {unit}",
            synthesis,
        )
    if quiz:
        add_box_page_break(document)
        append_box(
            document,
            quiz_box_template,
            "QUIZ NÃO AVALIATIVO",
            quiz,
            rows_to_keep=2,
        )
    if aai and aai_box_template is not None:
        append_box(
            document,
            aai_box_template,
            "ATIVIDADE AVALIATIVA INDIVIDUAL (AAI)",
            aai,
            rows_to_keep=2,
        )

    add_box_page_break(document)
    append_library_box(document, library_box_template, entries.get("direct", material))
    append_box(
        document,
        deeper_box_template,
        "PARA MERGULHAR NO ASSUNTO (SAIBA MAIS)",
        entries.get("deeper", ""),
        rows_to_keep=2,
    )
    append_box(
        document,
        podcast_box_template,
        "PODCAST / VÍDEO (CURADORIA DE ATÉ 45 MIN)",
        entries.get("podcast", ""),
        rows_to_keep=2,
    )
    append_box(
        document,
        article_box_template,
        "ARTIGO CIENTÍFICO",
        entries.get("article", ""),
        rows_to_keep=2,
    )

    document.element.body.append(section_properties)
    set_core_properties(
        document,
        f"Unidade {unit} — {DISCIPLINE}",
        "Material didático institucional com textos-base e roteiros de videoaula",
    )
    save_and_sanitize(document, spec["unit_output"])
    metrics = validate_unit_document(
        spec["unit_output"], unit, spec["lessons"], library_source_hash, counts
    )
    return spec["unit_output"], metrics


def build_question_document(unit: int, spec: dict) -> tuple[Path, dict]:
    source = read_text(spec["questions_md"])
    source_metrics = question_source_metrics(source)
    document = Document(spec["questions_template"])
    set_render_profile(document, "Arial")
    update_headers(document, unit, questions=True)
    cover = preserve_cover_elements(document, questions=True)
    section_properties = deepcopy(document.element.body.sectPr)
    clear_document_body(document)
    for element in cover:
        document.element.body.append(element)

    add_cover_title(
        document,
        f"QUESTIONÁRIO — UNIDADE {unit}",
        "Banco autoral de 40 questões com gabarito comentado",
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = paragraph.add_run(
        "20 questões de asserção-razão e 20 questões de interpretação; "
        "cinco alternativas e feedback por alternativa."
    )
    format_run(summary_run)
    add_page_break(document)

    source_without_title = re.sub(r"^#\s+.*?\n", "", source, count=1)
    render_markdown(
        document,
        source_without_title,
        page_break_before_headings=("Gabarito e feedbacks",),
    )
    document.element.body.append(section_properties)
    set_core_properties(
        document,
        f"Questionário — Unidade {unit} — {DISCIPLINE}",
        "Banco institucional de questões com gabarito e feedbacks",
    )
    save_and_sanitize(document, spec["questions_output"])
    metrics = validate_question_document(spec["questions_output"], unit, source_metrics)
    return spec["questions_output"], metrics


def save_and_sanitize(document: Document, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    sanitize_package(output)


def sanitize_package(path: Path) -> None:
    """Remove metadados auxiliares, comentários e Custom XML do pacote final."""

    removable_prefixes = (
        "customXml/",
        "word/comments",
        "word/people",
    )
    removable_names = {"docProps/custom.xml"}
    with ZipFile(path, "r") as archive:
        members = {
            info.filename: archive.read(info.filename)
            for info in archive.infolist()
            if info.filename not in removable_names
            and not info.filename.startswith(removable_prefixes)
        }

    relationship_namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    content_type_namespace = "http://schemas.openxmlformats.org/package/2006/content-types"

    for name, data in list(members.items()):
        if name.endswith(".rels"):
            from lxml import etree

            root = etree.fromstring(data)
            changed = False
            for relationship in list(root):
                target = (relationship.get("Target") or "").casefold()
                rel_type = (relationship.get("Type") or "").casefold()
                if (
                    "customxml" in target
                    or "custom-properties" in rel_type
                    or "comments" in target
                    or "people" in target
                ):
                    root.remove(relationship)
                    changed = True
            if changed:
                members[name] = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
        elif name == "[Content_Types].xml":
            from lxml import etree

            root = etree.fromstring(data)
            changed = False
            for child in list(root):
                part_name = (child.get("PartName") or "").casefold()
                if (
                    part_name.startswith("/customxml/")
                    or part_name == "/docprops/custom.xml"
                    or part_name.startswith("/word/comments")
                    or part_name.startswith("/word/people")
                ):
                    root.remove(child)
                    changed = True
            if changed:
                members[name] = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )

    file_descriptor, temporary_name = tempfile.mkstemp(
        prefix=path.stem + "-", suffix=".docx", dir=path.parent
    )
    os.close(file_descriptor)
    temporary = Path(temporary_name)
    try:
        with ZipFile(temporary, "w", compression=ZIP_DEFLATED) as archive:
            for name in sorted(members):
                archive.writestr(name, members[name])
        os.replace(temporary, path)
        os.chmod(path, 0o644)
    finally:
        if temporary.exists():
            temporary.unlink()


def document_non_library_text(document: Document) -> str:
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for section in document.sections:
        pieces.extend(paragraph.text for paragraph in section.header.paragraphs)
        pieces.extend(paragraph.text for paragraph in section.footer.paragraphs)
        for table in section.header.tables:
            pieces.append(table_text(table))
        for table in section.footer.tables:
            pieces.append(table_text(table))
    for table in document.tables:
        text = table_text(table)
        if "biblioteca virtual" not in fold(text):
            pieces.append(text)
    return "\n".join(pieces)


def package_metrics(path: Path) -> dict:
    with ZipFile(path, "r") as archive:
        names = archive.namelist()
        from lxml import etree

        xml_text_parts = []
        for name in names:
            if not name.endswith(".xml"):
                continue
            try:
                root = etree.fromstring(archive.read(name))
            except etree.XMLSyntaxError:
                continue
            xml_text_parts.append(" ".join(root.itertext()))
    xml_text = fold(" ".join(xml_text_parts))
    return {
        "custom_xml_parts": sum(name.startswith("customXml/") for name in names),
        "comment_parts": sum(name.startswith("word/comments") for name in names),
        "xml_internal_instruction_hits": sum(
            fold(marker) in xml_text for marker in INTERNAL_INSTRUCTION_MARKERS
        ),
        "xml_placeholder_hits": sum(
            fold(marker) in xml_text for marker in PLACEHOLDER_MARKERS
        ),
    }


def common_validation(document: Document, path: Path, unit: int) -> dict:
    non_library = document_non_library_text(document)
    full_body_text = "\n".join(
        [*(paragraph.text for paragraph in document.paragraphs),
         *(table_text(table) for table in document.tables)]
    )
    folded = fold(non_library)
    instruction_hits = [
        marker for marker in INTERNAL_INSTRUCTION_MARKERS if fold(marker) in folded
    ]
    placeholder_hits = [marker for marker in PLACEHOLDER_MARKERS if fold(marker) in folded]
    unexpected_credentials = len(CREDENTIAL_PATTERN.findall(non_library))
    header_text = " ".join(
        paragraph.text
        for section in document.sections
        for paragraph in section.header.paragraphs
    )
    metrics = {
        "reopened": True,
        "sections": len(document.sections),
        "inline_shapes": len(document.inline_shapes),
        "body_paragraphs": len(document.paragraphs),
        "top_level_tables": len(document.tables),
        "header_unit_correct": f"UNIDADE {unit}" in header_text.upper(),
        "internal_instruction_hits": len(instruction_hits),
        "placeholder_hits": len(placeholder_hits),
        "unexpected_credential_hits": unexpected_credentials,
        "latex_residual_hits": len(LATEX_RESIDUAL_RE.findall(full_body_text)),
        "markdown_residual_hits": len(MARKDOWN_RESIDUAL_RE.findall(full_body_text)),
        "contains_20_minute_wording": bool(
            re.search(r"(?i)(?:vinte|20)\s+minutos?", non_library)
        ),
        "package": package_metrics(path),
        "size_bytes": path.stat().st_size,
    }
    if not metrics["header_unit_correct"]:
        raise ValueError(f"Cabeçalho incorreto: {path.name}")
    if metrics["internal_instruction_hits"]:
        raise ValueError(f"Instrução interna remanescente: {path.name}")
    if metrics["placeholder_hits"]:
        raise ValueError(f"Placeholder remanescente: {path.name}")
    if metrics["unexpected_credential_hits"]:
        raise ValueError(f"Credencial fora do bloco autorizado: {path.name}")
    if metrics["latex_residual_hits"]:
        raise ValueError(f"Sintaxe LaTeX residual no DOCX: {path.name}")
    if metrics["markdown_residual_hits"]:
        raise ValueError(f"Marcação Markdown residual no DOCX: {path.name}")
    if metrics["contains_20_minute_wording"]:
        raise ValueError(f"Duração antiga encontrada: {path.name}")
    if (
        metrics["package"]["custom_xml_parts"]
        or metrics["package"]["comment_parts"]
        or metrics["package"]["xml_internal_instruction_hits"]
        or metrics["package"]["xml_placeholder_hits"]
    ):
        raise ValueError(f"Parte XML interna remanescente: {path.name}")
    return metrics


def validate_unit_document(
    path: Path,
    unit: int,
    lesson_numbers: Sequence[int],
    expected_library_hash: str,
    script_counts: dict[int, dict[str, int]],
) -> dict:
    document = Document(path)
    metrics = common_validation(document, path, unit)
    tables = document.tables
    all_text = "\n".join(table_text(table) for table in tables)
    text_labels = [
        number
        for number in lesson_numbers
        if re.search(rf"TEXTO BASE\s+—\s+AULA {number}\b", all_text, re.IGNORECASE)
    ]
    script_labels = [
        number
        for number in lesson_numbers
        if re.search(
            rf"ROTEIRO DA VIDEOAULA {number}\b", all_text, re.IGNORECASE
        )
    ]
    library_tables = [
        table for table in tables if "biblioteca virtual" in fold(table_text(table))
    ]
    if len(library_tables) != 1:
        raise ValueError(f"Quantidade inválida de blocos da Biblioteca: {path.name}")
    library_preserved = credential_row_hash(library_tables[0]) == expected_library_hash
    metrics.update(
        {
            "document_type": "material_didatico",
            "lesson_text_labels": text_labels,
            "lesson_script_labels": script_labels,
            "authorized_library_blocks": len(library_tables),
            "authorized_library_block_preserved": library_preserved,
            "script_word_counts": {
                str(number): script_counts[number] for number in lesson_numbers
            },
            "typography": typography_metrics(
                unit_content_paragraphs(document), "Times New Roman"
            ),
        }
    )
    if text_labels != list(lesson_numbers) or script_labels != list(lesson_numbers):
        raise ValueError(f"Rótulos de aula incorretos: {path.name}")
    if not library_preserved:
        raise ValueError(f"Bloco autorizado alterado: {path.name}")
    if unit in (1, 2):
        for number in lesson_numbers:
            if script_counts[number]["conservative_words"] > 1200:
                raise ValueError(f"Roteiro acima de 1.200 palavras: aula {number}")
    return metrics


def question_doc_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def iter_cell_paragraphs(cell: _Cell):
    yield from cell.paragraphs
    for nested_table in cell.tables:
        for row in nested_table.rows:
            seen_cells = set()
            for nested_cell in row.cells:
                identity = id(nested_cell._tc)
                if identity in seen_cells:
                    continue
                seen_cells.add(identity)
                yield from iter_cell_paragraphs(nested_cell)


def unit_content_paragraphs(document: Document):
    for table in document.tables:
        is_library = "biblioteca virtual" in fold(table_text(table))
        row_indexes = [1] if is_library else list(range(1, len(table.rows)))
        for row_index in row_indexes:
            seen_cells = set()
            for cell in table.rows[row_index].cells:
                identity = id(cell._tc)
                if identity in seen_cells:
                    continue
                seen_cells.add(identity)
                yield from iter_cell_paragraphs(cell)


def question_content_paragraphs(document: Document):
    started = False
    for paragraph in document.paragraphs:
        if paragraph._p.xpath(".//w:br[@w:type='page']"):
            started = True
            continue
        if started:
            yield paragraph


def typography_metrics(paragraphs, expected_font: str) -> dict:
    checked_runs = 0
    wrong_font_runs = 0
    wrong_size_runs = 0
    checked_paragraphs = 0
    wrong_line_spacing = 0
    wrong_alignment = 0
    nonzero_space_after = 0

    for paragraph in paragraphs:
        text = "".join(paragraph._p.itertext()).strip()
        if not text:
            continue
        checked_paragraphs += 1
        line_spacing = paragraph.paragraph_format.line_spacing
        if not isinstance(line_spacing, float) or abs(line_spacing - 1.15) > 0.001:
            wrong_line_spacing += 1
        if paragraph.text.strip() not in {"PORQUE", "OU"} and paragraph.alignment not in (
            WD_ALIGN_PARAGRAPH.LEFT,
            None,
        ):
            wrong_alignment += 1
        after = paragraph.paragraph_format.space_after
        if after is not None and abs(after.pt) > 0.001:
            nonzero_space_after += 1

        for run_element in paragraph._p.iter(qn("w:r")):
            run_text = "".join(run_element.itertext()).strip()
            if not run_text:
                continue
            checked_runs += 1
            properties = run_element.find(qn("w:rPr"))
            fonts = properties.find(qn("w:rFonts")) if properties is not None else None
            size = properties.find(qn("w:sz")) if properties is not None else None
            ascii_font = fonts.get(qn("w:ascii")) if fonts is not None else None
            hansi_font = fonts.get(qn("w:hAnsi")) if fonts is not None else None
            size_value = size.get(qn("w:val")) if size is not None else None
            if ascii_font != expected_font or hansi_font != expected_font:
                wrong_font_runs += 1
            if size_value != "24":
                wrong_size_runs += 1

    metrics = {
        "expected_font": expected_font,
        "expected_size_pt": 12,
        "expected_line_spacing": 1.15,
        "expected_alignment": "left",
        "expected_space_after_pt": 0,
        "checked_runs": checked_runs,
        "wrong_font_runs": wrong_font_runs,
        "wrong_size_runs": wrong_size_runs,
        "checked_paragraphs": checked_paragraphs,
        "wrong_line_spacing_paragraphs": wrong_line_spacing,
        "wrong_alignment_paragraphs": wrong_alignment,
        "nonzero_space_after_paragraphs": nonzero_space_after,
    }
    if any(
        metrics[key]
        for key in (
            "wrong_font_runs",
            "wrong_size_runs",
            "wrong_line_spacing_paragraphs",
            "wrong_alignment_paragraphs",
            "nonzero_space_after_paragraphs",
        )
    ):
        raise ValueError(f"Tipografia fora do padrão: {metrics}")
    return metrics


def validate_question_document(path: Path, unit: int, source_metrics: dict) -> dict:
    document = Document(path)
    metrics = common_validation(document, path, unit)
    text = question_doc_text(document)
    split = re.split(r"(?im)^Gabarito e feedbacks\s*$", text, maxsplit=1)
    if len(split) != 2:
        raise ValueError(f"Gabarito ausente no DOCX: {path.name}")
    questions, feedbacks = split
    generated = {
        "questions": len(
            re.findall(r"(?m)^(?:[1-9]|[1-3]\d|40)\.\s+", questions)
        )
        + len(
            re.findall(
                r"(?m)^Questão\s+(?:[1-9]|[1-3]\d|40)\b", questions
            )
        ),
        "alternatives": len(re.findall(r"(?m)^\*?[a-e]\.\s+", questions)),
        "correct_markers": len(re.findall(r"(?m)^\*[a-e]\.\s+", questions)),
        "feedback_sections": len(
            re.findall(r"(?m)^Questão\s+(?:[1-9]|[1-3]\d|40)\b", feedbacks)
        ),
        "feedback_items": len(re.findall(r"(?m)^•\s+[a-e]\.\s+", feedbacks)),
        "answer_distribution": {
            letter: len(re.findall(rf"(?m)^\*{letter}\.\s+", questions))
            for letter in "abcde"
        },
    }
    metrics.update(
        {
            "document_type": "questionario",
            "question_metrics": generated,
            "source_metrics_match": generated == source_metrics,
            "typography": typography_metrics(
                question_content_paragraphs(document), "Arial"
            ),
        }
    )
    if generated != source_metrics:
        raise ValueError(
            f"Contagem do DOCX diverge do Markdown em {path.name}: {generated}"
        )
    return metrics


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    report = {
        "generated_at": "2026-08-01",
        "generator": str(Path(__file__).relative_to(ROOT)),
        "discipline": DISCIPLINE,
        "files": {},
        "template_exceptions": {
            "unit_2": (
                "O modelo institucional da Unidade 2 trazia internamente rótulos das Aulas 1–4; "
                "as mesmas caixas e sua formatação foram preservadas e corretamente remapeadas "
                "para as Aulas 5–8."
            )
        },
        "security": {
            "authorized_library_block": (
                "Preservado em cada material didático por autorização; nenhum valor é incluído neste relatório."
            ),
            "custom_xml_removed": True,
        },
        "all_passed": False,
    }

    for unit, spec in UNIT_SPECS.items():
        for required in (
            spec["unit_md"],
            spec["script_md"],
            spec["questions_md"],
            spec["unit_template"],
            spec["questions_template"],
        ):
            if not required.is_file():
                raise FileNotFoundError(required)

        unit_path, unit_metrics = build_unit_document(unit, spec)
        report["files"][unit_path.name] = unit_metrics
        question_path, question_metrics = build_question_document(unit, spec)
        report["files"][question_path.name] = question_metrics

    report["all_passed"] = len(report["files"]) == 8
    report_path = OUTPUT_DIR / "validacao_docx.json"
    report_path.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )

    print(f"DOCX gerados e validados: {len(report['files'])}")
    for name, metrics in report["files"].items():
        if metrics["document_type"] == "questionario":
            detail = "40 questões; 200 alternativas; 40 gabaritos; 200 feedbacks"
        else:
            detail = (
                f"aulas {metrics['lesson_text_labels'][0]}–"
                f"{metrics['lesson_text_labels'][-1]}; bloco institucional autorizado preservado"
            )
        print(f"- {name}: OK — {detail}")
    print(f"Relatório: {report_path.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:  # relatório de falha sem conteúdo sensível
        print(f"ERRO: {type(exc).__name__}: {exc}", file=sys.stderr)
        raise
