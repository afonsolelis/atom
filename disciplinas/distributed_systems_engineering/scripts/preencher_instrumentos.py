#!/usr/bin/env python3
"""Preenche os modelos de avaliação: 40 questões por unidade, avaliação final e PBL.

Nas questões objetivas, a alternativa correta mantém o asterisco na frente da
letra, como exige o modelo institucional; ele é renderizado em destaque para
facilitar a conferência da coordenação.

Uso:
    PYTHONPATH=/tmp/dse-docx-libs python3 scripts/preencher_instrumentos.py
    PYTHONPATH=/tmp/dse-docx-libs python3 scripts/preencher_instrumentos.py questoes
    PYTHONPATH=/tmp/dse-docx-libs python3 scripts/preencher_instrumentos.py avaliacao
    PYTHONPATH=/tmp/dse-docx-libs python3 scripts/preencher_instrumentos.py entrega
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx_comum import (  # noqa: E402
    CONTEUDISTA,
    DISCIPLINA,
    DOCUMENTOS,
    FONTE_AVALIACAO,
    RAIZ,
    achar_caixa,
    arquivo_avaliacao,
    arquivo_entrega_trabalho,
    arquivo_questoes,
    carregar,
    destravar_linhas_e_quebras,
    formatar_run,
    remover_caixa,
    renderizar,
    secoes,
)

INSTRUMENTOS = DOCUMENTOS / "Instrumentos Avaliativos"


def preencher_identificacao(documento, rotulos: tuple[str, ...], valor: str) -> bool:
    for paragrafo in documento.paragraphs:
        if paragrafo.text.strip().lower().startswith(rotulos):
            formatar_run(paragrafo.add_run(" " + valor), FONTE_AVALIACAO)
            return True
    return False


def cortar_a_partir_de(documento, condicao, manter_casado: bool = True) -> int:
    """Apaga os parágrafos do corpo a partir do primeiro que casa `condicao`."""
    paragrafos = documento.paragraphs
    inicio = next(
        (i for i, p in enumerate(paragrafos) if condicao(p.text.strip())), None
    )
    if inicio is None:
        return 0
    if manter_casado:
        inicio += 1
    removidos = 0
    for paragrafo in paragrafos[inicio:]:
        paragrafo._p.getparent().remove(paragrafo._p)
        removidos += 1
    return removidos


def linhas_a_partir_de(markdown: str, padrao: str) -> list[str]:
    linhas = markdown.split("\n")
    for indice, linha in enumerate(linhas):
        if re.match(padrao, linha):
            return linhas[indice:]
    return linhas


E_CONTEUDISTA = lambda texto: texto.lower().startswith(  # noqa: E731
    ("professor-conteudista", "conteudista:")
)


def preencher_linear(modelo: Path, fonte_md: Path, padrao: str, destino: Path) -> Path:
    documento = Document(modelo)
    markdown = carregar(fonte_md)

    preencher_identificacao(documento, ("disciplina:",), DISCIPLINA)
    preencher_identificacao(
        documento, ("professor-conteudista:", "conteudista:"), CONTEUDISTA
    )
    removidos = cortar_a_partir_de(documento, E_CONTEUDISTA)

    corpo = linhas_a_partir_de(markdown, padrao)
    renderizar(documento, corpo, fonte=FONTE_AVALIACAO, alternativas=True)
    destravar_linhas_e_quebras(documento)

    documento.core_properties.subject = DISCIPLINA
    documento.core_properties.author = CONTEUDISTA
    documento.core_properties.comments = ""

    destino.parent.mkdir(parents=True, exist_ok=True)
    documento.save(destino)
    corretas = sum(1 for l in corpo if re.match(r"^\*[a-eA-E]\.\s", l.strip()))
    print(
        f"  OK -> {destino.name}  "
        f"(orientações removidas: {removidos} par.; alternativas com '*': {corretas})"
    )
    return destino


def fazer_questoes() -> list[Path]:
    print("QUESTIONÁRIOS (40 questões por unidade):")
    saidas = []
    for unidade in (1, 2, 3, 4):
        modelo = (
            DOCUMENTOS
            / f"Unidade {unidade}"
            / f"40 Questões - UNI{unidade}_nomedadisciplina.docx"
        )
        destino = arquivo_questoes(unidade)
        saidas.append(
            preencher_linear(
                modelo,
                RAIZ / f"unidade_{unidade}" / f"questoes_uni{unidade}.md",
                r"^##\s+Quest",
                destino,
            )
        )
    return saidas


def fazer_avaliacao() -> Path:
    print("AVALIAÇÃO FINAL (10 discursivas):")
    modelo = INSTRUMENTOS / "Avaliação final_(10 discursivas)_nomedadisciplina.docx"
    destino = arquivo_avaliacao()
    return preencher_linear(
        modelo,
        RAIZ / "instrumentos_avaliativos" / "avaliacao_dissertativa.md",
        r"^#\s+Parte A",
        destino,
    )


CAIXAS_PBL = (
    ("1. título", "título"),
    ("2. desafio", "desafio"),
    ("3. fontes de pesquisa", "fonte de pesquisa"),
    ("4. componentes avaliativos", "entregável"),
    ("solução esperada", "solução"),
)


def fazer_entrega() -> Path:
    print("ENTREGA DE TRABALHO (PBL):")
    modelo = INSTRUMENTOS / "TEMPLATE ENTREGA DE TRABALHO - nomedadisciplina.docx"
    destino = arquivo_entrega_trabalho()
    fonte_md = RAIZ / "instrumentos_avaliativos" / "entrega_trabalho.md"

    documento = Document(modelo)
    markdown = carregar(fonte_md)
    blocos = secoes(markdown)

    # Orientações do modelo que não devem chegar à versão final.
    caixa_orientacao = achar_caixa(documento, "o case existe para que o estudante")
    if caixa_orientacao is not None:
        remover_caixa(caixa_orientacao)
    for paragrafo in list(documento.paragraphs):
        texto = paragrafo.text.strip().lower()
        if texto.startswith(
            ("atenção:", "excluir as orientações", "todos os itens são obrigatórios")
        ):
            paragrafo._p.getparent().remove(paragrafo._p)

    preenchidas = []
    for titulo, corpo in blocos:
        minusculo = titulo.lower()
        for chave, trecho in CAIXAS_PBL:
            if not minusculo.startswith(chave):
                continue
            caixa = achar_caixa(documento, trecho)
            if caixa is not None:
                renderizar(
                    caixa.rows[1].cells[0], corpo, fonte=FONTE_AVALIACAO, limpar=True
                )
                preenchidas.append(chave)
            break

    # O roteiro do estudante fica no corpo, depois das caixas.
    cortar_a_partir_de(
        documento, lambda t: t.strip().lower() == "roteiro do estudante"
    )
    roteiro = next(
        (c for t, c in blocos if t.lower().startswith("roteiro do estudante")), []
    )
    if roteiro:
        renderizar(documento, roteiro, fonte=FONTE_AVALIACAO)
        preenchidas.append("roteiro do estudante")

    destravar_linhas_e_quebras(documento)
    documento.core_properties.subject = DISCIPLINA
    documento.core_properties.author = CONTEUDISTA
    documento.core_properties.comments = ""

    destino.parent.mkdir(parents=True, exist_ok=True)
    documento.save(destino)
    print(f"  OK -> {destino.name}  (seções: {', '.join(preenchidas)})")
    return destino


def main() -> int:
    alvo = sys.argv[1] if len(sys.argv) > 1 else "tudo"
    print("Disciplina:", DISCIPLINA)
    print("-" * 66)
    if alvo in ("tudo", "questoes"):
        fazer_questoes()
    if alvo in ("tudo", "avaliacao"):
        fazer_avaliacao()
    if alvo in ("tudo", "entrega"):
        fazer_entrega()
    print("-" * 66)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
