#!/usr/bin/env python3
"""Valida os DOCX gerados em `entrega_final/docx/` e grava um relatório JSON.

Confere integridade do pacote, presença das caixas institucionais, tipografia,
ausência de orientações do modelo e de sintaxe residual de Markdown/LaTeX, além
das contagens exigidas pelo contrato (40 questões, 200 alternativas, 40 corretas
marcadas com `*`, 8 por letra, 10 discursivas com rubrica).

Uso:
    PYTHONPATH=/tmp/dse-docx-libs python3 scripts/validar_entrega.py
"""

from __future__ import annotations

import hashlib
import json
import re
import sys
import unicodedata
from collections import Counter
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx_comum import (  # noqa: E402
    ENTREGA,
    FONTE_AVALIACAO,
    FONTE_UNIDADE,
    RAIZ,
    arquivo_avaliacao,
    arquivo_entrega_trabalho,
    arquivo_questoes,
    arquivo_unidade,
)

AULAS = {1: (1, 4), 2: (5, 8), 3: (9, 12), 4: (13, 16)}

# Trechos de orientação do modelo que não podem sobreviver na versão final.
ORIENTACOES_PROIBIDAS = (
    "apagar as orientacoes apos leitura",
    "elabore um titulo que gere curiosidade",
    "insira duas questoes para testar",
    "elabore uma questao dissertativa",
    "exemplo de questao",
    "exemplo de feedback",
    "excluir os exemplos",
    "excluir as orientacoes em vermelho",
    "o conteudista deve",
    "titulo do video (elabore",
    "insira o link do artigo aqui",
    "sobrenome, nome. titulo do artigo",
    "nenhuma das anteriores",
    "primeira alternativa",
)

# Sintaxe que deveria ter sido convertida na renderização.
RESIDUOS = {
    "markdown_negrito": re.compile(r"\*\*"),
    "markdown_cabecalho": re.compile(r"(?m)^#{1,6}\s"),
    "markdown_link": re.compile(r"\[[^\]]+\]\([^)]+\)"),
    "latex_comando": re.compile(r"\\[a-zA-Z]{2,}"),
    "latex_delimitador": re.compile(r"\$\$|(?<!\w)\$(?=[^\s$])"),
}

problemas: list[str] = []


def falha(mensagem: str) -> None:
    problemas.append(mensagem)


def dobrar(texto: str) -> str:
    sem_acento = unicodedata.normalize("NFKD", texto)
    sem_acento = "".join(c for c in sem_acento if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", sem_acento.lower())


def caixas(documento) -> list[Table]:
    corpo = documento.element.body
    vistas, resultado = set(), []
    for elemento in corpo.iter(qn("w:tbl")):
        ancestral = elemento.getparent()
        aninhada = False
        while ancestral is not None and ancestral is not corpo:
            if ancestral.tag == qn("w:tbl") and id(ancestral) in vistas:
                aninhada = True
                break
            ancestral = ancestral.getparent()
        if aninhada:
            continue
        vistas.add(id(elemento))
        resultado.append(Table(elemento, documento))
    return resultado


def texto_integral(documento) -> str:
    partes = [p.text for p in documento.paragraphs]
    for tabela in caixas(documento):
        for linha in tabela.rows:
            for celula in linha.cells:
                partes.append(celula.text)
    return "\n".join(partes)


def paragrafos_de_conteudo(documento):
    """Parágrafos autorais: dentro das caixas de conteúdo e no corpo."""
    for tabela in caixas(documento):
        cabecalho = dobrar(tabela.rows[0].cells[0].text)
        if cabecalho.startswith(("tabela para uso exclusivo", "data da validacao", "validado")):
            continue
        for linha in tabela.rows[1:]:
            for celula in linha.cells:
                yield from celula.paragraphs
    yield from documento.paragraphs


def conferir_integridade(caminho: Path) -> dict:
    dados = caminho.read_bytes()
    try:
        with ZipFile(caminho) as pacote:
            corrompido = pacote.testzip()
    except BadZipFile:
        falha(f"{caminho.name}: pacote ZIP inválido")
        corrompido = "ilegível"
    if corrompido:
        falha(f"{caminho.name}: entrada corrompida no pacote ({corrompido})")
    return {
        "bytes": len(dados),
        "sha256": hashlib.sha256(dados).hexdigest(),
        "zip_ok": corrompido is None,
    }


def conferir_orientacoes(caminho: Path, texto: str) -> list[str]:
    dobrado = dobrar(texto)
    encontrados = [t for t in ORIENTACOES_PROIBIDAS if t in dobrado]
    for trecho in encontrados:
        falha(f"{caminho.name}: orientação do modelo remanescente ({trecho!r})")
    return encontrados


def conferir_residuos(caminho: Path, texto: str) -> dict[str, int]:
    achados = {}
    for rotulo, padrao in RESIDUOS.items():
        # o asterisco isolado da alternativa correta é exigido pelo modelo
        ocorrencias = len(padrao.findall(texto))
        if ocorrencias:
            achados[rotulo] = ocorrencias
            falha(f"{caminho.name}: {ocorrencias} ocorrência(s) de {rotulo}")
    return achados


def conferir_tipografia(caminho: Path, documento, fonte: str) -> dict:
    contagem: Counter = Counter()
    fora = 0
    for paragrafo in paragrafos_de_conteudo(documento):
        for run in paragrafo.runs:
            if not run.text.strip():
                continue
            nome = run.font.name
            tamanho = run.font.size.pt if run.font.size else None
            contagem[(nome, tamanho)] += 1
            if nome not in (fonte, "Consolas", None):
                fora += 1
    if fora:
        falha(f"{caminho.name}: {fora} run(s) fora da fonte {fonte}")
    return {
        "runs_por_fonte": {f"{n}/{t}": q for (n, t), q in contagem.most_common(6)},
        "runs_fora_do_padrao": fora,
    }


def validar_unidade(caminho: Path, unidade: int) -> dict:
    documento = Document(caminho)
    texto = texto_integral(documento)
    primeira, ultima = AULAS[unidade]
    esperadas = list(range(primeira, ultima + 1))

    cabecalhos = [c.rows[0].cells[0].text.strip() for c in caixas(documento)]
    dobrados = [dobrar(c) for c in cabecalhos]

    faltando = []
    for numero in esperadas:
        if not any(re.search(rf"texto base aula {numero}(?!\d)", c) for c in dobrados):
            faltando.append(f"TEXTO BASE AULA {numero}")
        if not any(re.search(rf"roteiro videoaula {numero}(?!\d)", c) for c in dobrados):
            faltando.append(f"ROTEIRO VIDEOAULA {numero}")
    for obrigatoria in ("quiz nao avaliativo", "direto da fonte", "para mergulhar", "podcast", "artigo cientifico"):
        if not any(obrigatoria in c for c in dobrados):
            faltando.append(obrigatoria)
    if unidade == 1 and not any("aai" in c for c in dobrados):
        faltando.append("AAI")
    if unidade != 1 and any("aai" in c for c in dobrados):
        falha(f"{caminho.name}: caixa de AAI presente fora da Unidade 1")
    for ausente in faltando:
        falha(f"{caminho.name}: caixa ausente ({ausente})")

    vazias = [
        c.rows[0].cells[0].text.strip()[:40]
        for c in caixas(documento)
        if not dobrar(c.rows[0].cells[0].text).startswith(
            ("tabela para uso exclusivo", "data da validacao", "validado")
        )
        and sum(len(l.cells[0].text.split()) for l in c.rows[1:]) < 10
    ]
    for vazia in vazias:
        falha(f"{caminho.name}: caixa praticamente vazia ({vazia!r})")

    if "distributed systems engineering" not in dobrar(texto):
        falha(f"{caminho.name}: nome da disciplina ausente")

    return {
        "tipo": "unidade",
        "unidade": unidade,
        "caixas": len(cabecalhos),
        "aulas_esperadas": esperadas,
        "caixas_ausentes": faltando,
        "caixas_vazias": vazias,
        "palavras": len(texto.split()),
        "orientacoes_remanescentes": conferir_orientacoes(caminho, texto),
        "residuos": conferir_residuos(caminho, texto),
        "tipografia": conferir_tipografia(caminho, documento, FONTE_UNIDADE),
        **conferir_integridade(caminho),
    }


def validar_questionario(caminho: Path, unidade: int) -> dict:
    documento = Document(caminho)
    linhas = [p.text.strip() for p in documento.paragraphs]
    texto = texto_integral(documento)

    enunciados = [l for l in linhas if re.match(r"^\d+\.\s+\S", l)]
    alternativas = [l for l in linhas if re.match(r"^\*?[a-e]\.\s+\S", l)]
    corretas = [l for l in linhas if re.match(r"^\*[a-e]\.\s+\S", l)]
    distribuicao = Counter(l[1] for l in corretas)
    gabaritos = [l for l in linhas if re.match(r"^Questão \d+\s*\(correta:", l)]
    # as devolutivas por alternativa saem como itens de lista ("•  a. Correta: …")
    devolutivas = [l for l in linhas if re.match(r"^(?:•\s+)?[a-e]\.\s+(Correta|Incorreta)", l)]

    esperado = {
        "questoes": 40,
        "alternativas": 200,
        "corretas": 40,
        "gabaritos": 40,
        "devolutivas": 200,
    }
    obtido = {
        "questoes": len(enunciados),
        "alternativas": len(alternativas),
        "corretas": len(corretas),
        "gabaritos": len(gabaritos),
        "devolutivas": len(devolutivas),
    }
    for chave, alvo in esperado.items():
        if obtido[chave] != alvo:
            falha(f"{caminho.name}: {chave} = {obtido[chave]}, esperado {alvo}")
    for letra in "abcde":
        if distribuicao.get(letra, 0) != 8:
            falha(
                f"{caminho.name}: letra '{letra}' correta em "
                f"{distribuicao.get(letra, 0)} questões, esperado 8"
            )

    return {
        "tipo": "questionario",
        "unidade": unidade,
        **obtido,
        "distribuicao_da_correta": dict(sorted(distribuicao.items())),
        "orientacoes_remanescentes": conferir_orientacoes(caminho, texto),
        "residuos": conferir_residuos(caminho, texto),
        "tipografia": conferir_tipografia(caminho, documento, FONTE_AVALIACAO),
        **conferir_integridade(caminho),
    }


def validar_avaliacao(caminho: Path) -> dict:
    documento = Document(caminho)
    linhas = [p.text.strip() for p in documento.paragraphs]
    texto = texto_integral(documento)

    enunciados = [l for l in linhas if re.match(r"^Questão \d+ —", l)]
    respostas = [l for l in linhas if re.match(r"^Questão \d+$", l)]
    rubricas = len(re.findall(r"(?i)\b10 pontos\b", texto))

    if len(enunciados) != 10:
        falha(f"{caminho.name}: {len(enunciados)} enunciados, esperado 10")
    if len(respostas) < 10:
        falha(f"{caminho.name}: {len(respostas)} respostas esperadas, esperado 10")
    if re.search(r"(?m)^\*[a-e]\.\s", texto):
        falha(f"{caminho.name}: a avaliação final não deve ter questões objetivas")

    return {
        "tipo": "avaliacao_final",
        "enunciados": len(enunciados),
        "respostas_esperadas": len(respostas),
        "mencoes_a_rubrica_de_10_pontos": rubricas,
        "orientacoes_remanescentes": conferir_orientacoes(caminho, texto),
        "residuos": conferir_residuos(caminho, texto),
        "tipografia": conferir_tipografia(caminho, documento, FONTE_AVALIACAO),
        **conferir_integridade(caminho),
    }


def validar_pbl(caminho: Path) -> dict:
    documento = Document(caminho)
    texto = texto_integral(documento)
    dobrados = [dobrar(c.rows[0].cells[0].text) for c in caixas(documento)]

    obrigatorias = ("titulo", "desafio", "fonte de pesquisa", "entregavel", "solucao")
    ausentes = [o for o in obrigatorias if not any(o in c for c in dobrados)]
    for ausente in ausentes:
        falha(f"{caminho.name}: caixa obrigatória ausente ({ausente})")
    if "roteiro do estudante" not in dobrar(texto):
        falha(f"{caminho.name}: roteiro do estudante ausente")

    fontes = len(re.findall(r"https?://", texto))
    if fontes < 4:
        falha(f"{caminho.name}: apenas {fontes} fontes com link, esperado ao menos 4")

    return {
        "tipo": "entrega_trabalho",
        "caixas_obrigatorias_ausentes": ausentes,
        "links": fontes,
        "orientacoes_remanescentes": conferir_orientacoes(caminho, texto),
        "residuos": conferir_residuos(caminho, texto),
        "tipografia": conferir_tipografia(caminho, documento, FONTE_AVALIACAO),
        **conferir_integridade(caminho),
    }


def main() -> int:
    relatorio: dict[str, dict] = {}

    for unidade in (1, 2, 3, 4):
        caminho = arquivo_unidade(unidade)
        if not caminho.exists():
            falha(f"{caminho.name}: arquivo não encontrado")
            continue
        relatorio[caminho.name] = validar_unidade(caminho, unidade)

    for unidade in (1, 2, 3, 4):
        caminho = arquivo_questoes(unidade)
        if not caminho.exists():
            falha(f"{caminho.name}: arquivo não encontrado")
            continue
        relatorio[caminho.name] = validar_questionario(caminho, unidade)

    caminho = arquivo_avaliacao()
    if caminho.exists():
        relatorio[caminho.name] = validar_avaliacao(caminho)
    else:
        falha(f"{caminho.name}: arquivo não encontrado")

    caminho = arquivo_entrega_trabalho()
    if caminho.exists():
        relatorio[caminho.name] = validar_pbl(caminho)
    else:
        falha(f"{caminho.name}: arquivo não encontrado")

    destino = ENTREGA / "validacao_docx.json"
    destino.write_text(
        json.dumps(
            {"arquivos": relatorio, "problemas": problemas},
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )

    for nome, dados in relatorio.items():
        resumo = {
            "unidade": f"{dados.get('caixas')} caixas, {dados.get('palavras')} palavras",
            "questionario": (
                f"{dados.get('questoes')} questões, {dados.get('alternativas')} alternativas, "
                f"{dados.get('corretas')} corretas, distribuição {dados.get('distribuicao_da_correta')}"
            ),
            "avaliacao_final": f"{dados.get('enunciados')} discursivas",
            "entrega_trabalho": f"{dados.get('links')} links",
        }[dados["tipo"]]
        print(f"  {nome}\n      {resumo}")

    print("-" * 66)
    if problemas:
        print(f"{len(problemas)} problema(s):")
        for problema in problemas:
            print("  -", problema)
        return 1
    print(f"{len(relatorio)} arquivos validados sem problemas.")
    print(f"Relatório: {destino.relative_to(RAIZ)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
