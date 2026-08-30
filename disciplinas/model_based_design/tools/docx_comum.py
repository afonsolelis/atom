#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Utilitários compartilhados para preencher os templates institucionais Word
(UniFECAF / Núcleo das Engenharias e Tecnologia) da disciplina Model-Based
Design for Cyber-Physical Systems, a partir do conteúdo em Markdown.

Reaproveita a lógica de conversão Markdown -> runs do Word (negrito, itálico,
listas, tabelas, blockquotes) da referência em
disciplinas/data_engineering_and_pipelines/tools/preencher_docx.py, adaptada
para:
  - o renderizador de fórmulas local (tools/formulas.py: coletar_formulas /
    renderizar), que devolve PNGs em disco + dimensões em pontos, em vez do
    pipeline por navegador headless usado na outra disciplina;
  - a formatação institucional exigida para esta entrega: Times New Roman 12,
    espaçamento 1,15, alinhamento à esquerda, sem recuo de primeira linha;
  - blockquotes que contêm tabelas Markdown aninhadas (o desafio do PBL usa
    esse padrão) e itens de lista com fórmulas de bloco embutidas (o memorial
    de cálculo do PBL usa esse padrão).

Importado por tools/preencher_questoes.py e tools/preencher_instrumentos.py.
"""
from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table
from docx.text.paragraph import Paragraph

TOOLS = Path(__file__).resolve().parent
DISC = TOOLS.parent
PRISTINE = TOOLS / "_templates_pristine"
SAIDA = DISC / "entrega_docx"

DISCIPLINA = "Model-Based Design for Cyber-Physical Systems"
CONTEUDISTA = "Afonso Cesar Lelis Brandão"

# Os nomes de arquivo dos templates pristinos usam "Model Based Design" (sem
# hífen), diferente do nome da disciplina usado no CONTEÚDO ("Model-Based
# Design..."). Mantém os dois separados para não quebrar a busca de arquivo.
DISCIPLINA_ARQUIVO = "Model Based Design for Cyber-Physical Systems"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = 1.15

import formulas as fm  # tools/formulas.py: coletar_formulas / renderizar


# ============================================================ fórmulas
def build_formula_index(md_text: str) -> dict:
    """Coleta e renderiza todas as fórmulas de um Markdown.

    Devolve {(tex, display): {"png": Path, "largura_pt": float, "altura_pt": float}}.
    """
    formulas = fm.coletar_formulas(md_text)
    if not formulas:
        return {}
    return fm.renderizar(formulas)


def new_stats() -> dict:
    return {"bloco_ok": 0, "linha_ok": 0, "faltantes": []}


def stats_total_imagens(stats: dict) -> int:
    return stats["bloco_ok"] + stats["linha_ok"]


# ============================================================ formatação base
def apply_para_format(p):
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.first_line_indent = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def new_para(container):
    """Cria um parágrafo já formatado (Times New Roman 12 / 1,15 / esquerda /
    sem recuo de primeira linha) em uma célula OU no corpo do documento."""
    p = container.add_paragraph()
    apply_para_format(p)
    return p


def style_run(r, bold=None, italic=None):
    r.font.name = FONT_NAME
    r.font.size = FONT_SIZE
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic
    return r


# ============================================================ docx helpers
def clear_container(cell):
    """Remove tudo de uma célula (parágrafos, tabelas aninhadas, sdt) menos
    as propriedades da célula (w:tcPr)."""
    tc = cell._tc
    tr = tc.getparent()
    if tr is not None and tr.tag == qn("w:tr"):
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None:
            for tag in ("w:cantSplit", "w:tblHeader"):
                el = trPr.find(qn(tag))
                if el is not None:
                    trPr.remove(el)
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)


def ensure_trailing_p(cell):
    tc = cell._tc
    if len(tc) == 0 or tc[-1].tag != qn("w:p"):
        cell.add_paragraph("")


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color")
    col.set(qn("w:val"), "0563C1")
    rpr.append(col)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


INLINE_RE = re.compile(
    r"\*\*(?P<b>.+?)\*\*"
    r"|`(?P<c>[^`]+)`"
    r"|\[(?P<lt>[^\]]+)\]\((?P<lu>[^)]+)\)"
    r"|(?<![\*\w])\*(?P<i>[^*]+?)\*(?![\*\w])"
)


def add_md_runs(p, text):
    """Adiciona runs com **negrito**, `código` (sem monoespaçado — mantém
    Times New Roman 12 por exigência de formatação), [link](url), *itálico*."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            style_run(p.add_run(text[pos:m.start()]))
        if m.group("b") is not None:
            style_run(p.add_run(m.group("b")), bold=True)
        elif m.group("c") is not None:
            style_run(p.add_run(m.group("c")))
        elif m.group("lt") is not None:
            add_hyperlink(p, m.group("lu"), m.group("lt"))
        elif m.group("i") is not None:
            style_run(p.add_run(m.group("i")), italic=True)
        pos = m.end()
    if pos < len(text):
        style_run(p.add_run(text[pos:]))


INLINE_MATH_RE = re.compile(r"(?<!\$)\$([^$\n]+)\$(?!\$)")


def iter_math(text):
    """Divide uma linha em partes ('text', conteudo) / ('inline', latex)."""
    pos = 0
    for m in INLINE_MATH_RE.finditer(text):
        if m.start() > pos:
            yield ("text", text[pos:m.start()])
        yield ("inline", m.group(1))
        pos = m.end()
    if pos < len(text):
        yield ("text", text[pos:])


def add_inline(p, text, formula_index, stats):
    """Como add_md_runs, mas tratando $fórmula$ inline como imagem."""
    for kind, content in iter_math(text):
        if kind == "text":
            add_md_runs(p, content.replace(r"\$", "$"))
        else:
            data = formula_index.get((content.strip(), False))
            if data:
                r = p.add_run()
                r.add_picture(io.BytesIO(data["png"].read_bytes()),
                              width=Pt(data["largura_pt"]), height=Pt(data["altura_pt"]))
                stats["linha_ok"] += 1
            else:
                style_run(p.add_run("$%s$" % content))
                stats["faltantes"].append(content.strip())


MAX_BLOCK_WIDTH_PT = 420.0


def add_display_math(container, latex, formula_index, stats):
    p = new_para(container)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data = formula_index.get((latex.strip(), True))
    if data:
        w = data["largura_pt"]
        h = data["altura_pt"]
        if w > MAX_BLOCK_WIDTH_PT:
            scale = MAX_BLOCK_WIDTH_PT / w
            w, h = MAX_BLOCK_WIDTH_PT, h * scale
        p.add_run().add_picture(io.BytesIO(data["png"].read_bytes()), width=Pt(w), height=Pt(h))
        stats["bloco_ok"] += 1
    else:
        style_run(p.add_run(latex.strip()))
        stats["faltantes"].append(latex.strip())
    return p


# ============================================================ tabelas markdown
def set_table_full_width(t, ncols):
    tblPr = t._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "autofit")
    col_w = Inches(6.0 / ncols)
    for row in t.rows:
        for c in row.cells:
            c.width = col_w


def add_markdown_table(container, tbl_lines, formula_index, stats):
    rows = []
    for k, ln in enumerate(tbl_lines):
        if k == 1:
            continue  # separador |---|
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = container.add_table(rows=len(rows), cols=ncols)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    set_table_full_width(t, ncols)
    for ri, r in enumerate(rows):
        for ci in range(ncols):
            tc = t.rows[ri].cells[ci]
            tc.text = ""
            p = tc.paragraphs[0]
            apply_para_format(p)
            val = r[ci] if ci < len(r) else ""
            add_inline(p, val, formula_index, stats)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    new_para(container)


def _is_table_start(lines, i):
    return (i < len(lines) and lines[i].strip().startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]))


# ============================================================ blockquote (com tabelas aninhadas)
def render_blockquote(container, quote_lines, formula_index, stats):
    i, n = 0, len(quote_lines)
    para_buf = []

    def flush():
        if not para_buf:
            return
        p = new_para(container)
        p.paragraph_format.left_indent = Inches(0.3)
        add_inline(p, " ".join(para_buf), formula_index, stats)
        for r in p.runs:
            r.italic = True
        para_buf.clear()

    while i < n:
        line = quote_lines[i]
        s = line.strip()
        if s == "":
            flush()
            i += 1
            continue
        if _is_table_start(quote_lines, i):
            flush()
            tbl_lines = []
            while i < n and quote_lines[i].strip().startswith("|"):
                tbl_lines.append(quote_lines[i].strip())
                i += 1
            add_markdown_table(container, tbl_lines, formula_index, stats)
            continue
        para_buf.append(s)
        i += 1
    flush()


# ============================================================ renderizador de blocos markdown
def render_blocks(container, lines, formula_index, stats, first=False, alts=False):
    """Renderiza linhas de Markdown em `container` (célula OU corpo do doc).

    - first=True: limpa o container antes de renderizar (uso em células).
    - alts=True: linhas 'a.'/'*a.' (alternativas de questão) preservam o
      asterisco que marca a alternativa correta.

      O modelo institucional é explícito quanto a isso: "Cada questão deve ter
      5 alternativas. Uma delas será a correta, que deve ser destacada com um
      asterisco (*) antes da alternativa correta", e os dois exemplos do
      próprio modelo trazem `*a. …`. O asterisco não é convenção do
      Markdown-fonte: é a marcação de gabarito exigida no entregável.
    """
    if first:
        clear_container(container)
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if s == "":
            i += 1
            continue

        # --- heading (markdown # a ######) -> parágrafo em negrito, sem
        #     mudar o estilo do parágrafo (não altera o catálogo de estilos
        #     do template)
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            p = new_para(container)
            add_inline(p, m.group(2), formula_index, stats)
            for r in p.runs:
                r.bold = True
            i += 1
            continue

        # --- horizontal rule
        if re.match(r"^-{3,}$", s):
            i += 1
            continue

        # --- display math: bloco $$ ... $$
        if s.startswith("$$") and s.endswith("$$") and len(s) > 4:
            add_display_math(container, s[2:-2], formula_index, stats)
            i += 1
            continue
        if s == "$$":
            i += 1
            expr = []
            while i < n and lines[i].strip() != "$$":
                expr.append(lines[i])
                i += 1
            i += 1
            # $$...$$ preserva quebras de linha internas (coletor.py usa
            # regex com DOTALL sobre o texto bruto, sem juntar por espaço).
            add_display_math(container, "\n".join(expr), formula_index, stats)
            continue

        # --- display math: bloco delimitado por linhas soltas com "$"
        if s == "$":
            i += 1
            expr = []
            while i < n and lines[i].strip() != "$":
                expr.append(lines[i])
                i += 1
            i += 1
            add_display_math(container, " ".join(expr), formula_index, stats)
            continue

        # --- blockquote
        if s.startswith(">"):
            quote = []
            while i < n and (lines[i].strip().startswith(">") or lines[i].strip() == ""):
                if lines[i].strip() == "":
                    if i + 1 < n and lines[i + 1].strip().startswith(">"):
                        quote.append("")
                        i += 1
                        continue
                    break
                quote.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            render_blockquote(container, quote, formula_index, stats)
            continue

        # --- tabela markdown
        if _is_table_start(lines, i):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(container, tbl_lines, formula_index, stats)
            continue

        # --- alternativa de questão ('*a.' / 'a.') — o asterisco da alternativa
        #     correta é PRESERVADO, conforme exige o modelo institucional
        if alts:
            malt = re.match(r"^(\*?)([a-eA-E])\.\s+(.*)$", s)
            if malt:
                marca, letter, rest = malt.group(1), malt.group(2), malt.group(3)
                p = new_para(container)
                style_run(p.add_run(marca + letter.lower() + ". "))
                add_inline(p, rest, formula_index, stats)
                i += 1
                continue

        # --- lista (marcadores / numerada), possivelmente multi-linha, com
        #     suporte a fórmula de bloco embutida em um item de lista
        if re.match(r"^[-*]\s+", s) or re.match(r"^\d+\.\s+", s):
            while i < n and (re.match(r"^\s*[-*]\s+", lines[i]) or re.match(r"^\s*\d+\.\s+", lines[i])
                             or (lines[i].strip() and lines[i].startswith("  "))):
                raw = lines[i]
                indent = len(raw) - len(raw.lstrip(" "))
                item = raw.strip()

                # fórmula de bloco embutida no meio de um item de lista
                # (linha isolada "$", ex.: memorial de cálculo do PBL)
                if item == "$":
                    i += 1
                    expr = []
                    while i < n and lines[i].strip() != "$":
                        expr.append(lines[i])
                        i += 1
                    i += 1
                    add_display_math(container, " ".join(expr), formula_index, stats)
                    continue

                mnum = re.match(r"^(\d+)\.\s+(.*)$", item)
                mbul = re.match(r"^[-*]\s+(.*)$", item)
                p = new_para(container)
                p.paragraph_format.left_indent = Inches(0.25 + (0.25 if indent >= 2 else 0))
                if mbul:
                    style_run(p.add_run("•  "))
                    add_inline(p, mbul.group(1), formula_index, stats)
                elif mnum:
                    style_run(p.add_run(mnum.group(1) + ".  "))
                    add_inline(p, mnum.group(2), formula_index, stats)
                else:
                    add_inline(p, item, formula_index, stats)
                i += 1
            continue

        # --- parágrafo normal
        p = new_para(container)
        add_inline(p, s, formula_index, stats)
        i += 1

    if hasattr(container, "_tc"):
        ensure_trailing_p(container)


# ============================================================ gabarito (feedbacks das 40 questões)
QUESTAO_HDR_RE = re.compile(r"^\*\*Quest(?:ã|a)o\s+(\d+)\*\*\s*\(correta:\s*([a-eA-E])\)\s*$")
FEEDBACK_ALT_RE = re.compile(r"^-\s+([a-eA-E])\.\s+(.*)$")


def render_gabarito(container, lines, formula_index, stats):
    """Renderiza a seção 'Gabarito e feedbacks': cabeçalho em negrito
    'Questão N (correta: X)' seguido das 5 linhas de feedback (uma por
    alternativa), sem qualquer asterisco. Devolve a quantidade de linhas de
    feedback (alternativas) renderizadas."""
    i, n = 0, len(lines)
    feedback_count = 0
    while i < n:
        s = lines[i].strip()
        if s == "":
            i += 1
            continue
        mq = QUESTAO_HDR_RE.match(s)
        if mq:
            p = new_para(container)
            style_run(p.add_run(f"Questão {mq.group(1)} (correta: {mq.group(2).lower()})"), bold=True)
            i += 1
            continue
        mf = FEEDBACK_ALT_RE.match(s)
        if mf:
            p = new_para(container)
            style_run(p.add_run(mf.group(1).lower() + ". "))
            add_inline(p, mf.group(2), formula_index, stats)
            feedback_count += 1
            i += 1
            continue
        # fallback: parágrafo comum (não deveria ocorrer no formato esperado)
        p = new_para(container)
        add_inline(p, s, formula_index, stats)
        i += 1
    return feedback_count


# ============================================================ localizar tabelas / cortar parágrafos
def find_table(doc, needle):
    """Busca em doc.tables (tabelas de nível superior do corpo — não entra em
    content controls w:sdt) por substring no texto da 1a célula."""
    needle = needle.lower()
    for t in doc.tables:
        try:
            h = t.rows[0].cells[0].text.strip().lower()
        except Exception:
            continue
        if needle in h:
            return t
    return None


def set_placeholder(doc, prefixos, valor):
    for p in doc.paragraphs:
        low = p.text.strip().lower()
        if any(low.startswith(pre) for pre in prefixos):
            style_run(p.add_run(" " + valor))
            return True
    return False


def cut_body_after(doc, pred, keep_matched=True):
    """Remove parágrafos de nível de corpo a partir do 1o que casa `pred`
    (mantendo-o se keep_matched). Retorna a quantidade removida."""
    paras = doc.paragraphs
    start = None
    for idx, p in enumerate(paras):
        if pred(p.text.strip()):
            start = idx
            break
    if start is None:
        return 0
    if keep_matched:
        start += 1
    removed = 0
    for p in paras[start:]:
        p._p.getparent().remove(p._p)
        removed += 1
    return removed


def remove_paragraphs_with_prefix(doc, prefixos_lower):
    """Remove parágrafos de nível de corpo cujo texto (lowercased) comece com
    algum dos prefixos dados. Usado para tirar avisos de orientação pontuais
    que não estão organizados de forma linear no template (ex.: ENTREGA)."""
    removed = 0
    for p in list(doc.paragraphs):
        low = p.text.strip().lower()
        if low and any(low.startswith(pre) for pre in prefixos_lower):
            p._p.getparent().remove(p._p)
            removed += 1
    return removed


def remove_table(doc, t):
    t._tbl.getparent().remove(t._tbl)


def unlock_rows_and_breaks(doc):
    for tbl in doc.element.body.iter(qn("w:tbl")):
        for tr in tbl.findall(qn("w:tr")):
            trPr = tr.find(qn("w:trPr"))
            if trPr is not None:
                for tag in ("w:cantSplit", "w:tblHeader"):
                    el = trPr.find(qn(tag))
                    if el is not None:
                        trPr.remove(el)
    for p in doc.element.body.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            for tag in ("w:pageBreakBefore", "w:keepNext", "w:keepLines"):
                el = pPr.find(qn(tag))
                if el is not None:
                    pPr.remove(el)


# ============================================================ parsing de markdown por seções
def parse_h2(md_text):
    out = []
    title, buf = None, []
    for ln in md_text.split("\n"):
        if re.match(r"^##\s+", ln):
            if title is not None:
                out.append((title, buf))
            title = ln[2:].strip()
            buf = []
        else:
            if title is not None:
                buf.append(ln)
    if title is not None:
        out.append((title, buf))
    return out


def sections_dict(md_text):
    return dict(parse_h2(md_text))


PARTE_B_RE = re.compile(r"^# Parte B.*$", re.M)


def split_parte_b(md_text):
    """Divide o arquivo-mestre casando '# Parte B' em início de linha (não a
    primeira ocorrência literal — a expressão também aparece entre crases no
    bloco de controle de versão, no topo do arquivo). Devolve
    (texto_parte_a, texto_parte_b_ou_None)."""
    m = PARTE_B_RE.search(md_text)
    if not m:
        return md_text, None
    return md_text[: m.start()], md_text[m.start():]


# ============================================================ verificação
def verify_docx(path: Path) -> dict:
    """Abre o .docx com python-docx e valida o ZIP. Devolve um relatório."""
    import zipfile

    report = {"abre_sem_erro": False, "zip_ok": False, "erro": None}
    try:
        with zipfile.ZipFile(path) as z:
            bad = z.testzip()
            report["zip_ok"] = bad is None
            report["zip_bad_file"] = bad
    except Exception as exc:  # noqa: BLE001
        report["erro"] = f"zip: {exc}"
        return report
    try:
        Document(path)
        report["abre_sem_erro"] = True
    except Exception as exc:  # noqa: BLE001
        report["erro"] = f"python-docx: {exc}"
    return report


ORIENTATION_MARKERS = (
    "xxxx", "excluir", "atenção:", "inserir a resposta esperada",
    "seguir à risca", "10 questões do tipo asserção-razão", "exemplo de questão",
    "exemplo de feedback", "leia as orientações",
)


def find_leftover_orientation(path: Path):
    """Extrai todo o texto (parágrafos + tabelas) e procura por marcadores de
    orientação que não deveriam sobrar no documento final."""
    doc = Document(path)
    achados = []
    todo_texto = []

    for p in doc.paragraphs:
        todo_texto.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                todo_texto.append(cell.text)

    texto_completo = "\n".join(todo_texto)
    baixo = texto_completo.lower()
    for marcador in ORIENTATION_MARKERS:
        if marcador in baixo:
            achados.append(marcador)
    return achados, texto_completo


def extract_all_text(path: Path) -> str:
    doc = Document(path)
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                partes.append(cell.text)
    return "\n".join(partes)


def count_formula_images(path: Path) -> int:
    """Conta imagens embutidas no pacote .docx (proxy para nº de fórmulas
    renderizadas como figura)."""
    import zipfile

    with zipfile.ZipFile(path) as z:
        return sum(1 for n in z.namelist() if n.startswith("word/media/"))
