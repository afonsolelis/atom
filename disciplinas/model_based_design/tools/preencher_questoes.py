#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Preenche os quatro templates institucionais "40 Questões - UNIN" (N = 1..4)
da disciplina Model-Based Design for Cyber-Physical Systems a partir de
unidade_N/questoes_uniN.md, dentro do template Word oficial (UniFECAF /
Núcleo das Engenharias e Tecnologia).

- Preserva as tabelas do coordenador (INTOCÁVEIS).
- Remove as orientações e os exemplos do template (mantendo título e
  identificação).
- Preenche Disciplina / Professor-conteudista.
- Insere as 40 questões (20 asserção-razão + 20 interpretação), seguindo a
  convenção interna do contrato desta disciplina, e não a frase de "10
  questões" do texto de orientação do template (ver ANALISE_MATERIAIS_
  RECEBIDOS.md / DIRETRIZES_PRODUCAO.md).
- A alternativa correta É marcada com asterisco imediatamente antes da
  letra (`*a. ...`), como o modelo institucional exige textualmente e
  demonstra em seus dois exemplos. A correção também aparece na seção de
  gabarito/feedback ao final.
- Os 200 feedbacks (5 por questão x 40 questões) vão ao final do documento.

Sempre parte do backup pristine em tools/_templates_pristine/ e escreve em
entrega_docx/, sobrescrevendo a cada execução.

Uso:
    tools/.venv/bin/python tools/preencher_questoes.py         # unidades 1-4
    tools/.venv/bin/python tools/preencher_questoes.py 1 3     # só 1 e 3
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
import docx_comum as dc


def fill_unit(n: int) -> dict:
    template_name = f"40 Questões - UNI{n}_{dc.DISCIPLINA_ARQUIVO}.docx"
    pristine = dc.PRISTINE / template_name
    md_path = dc.DISC / f"unidade_{n}" / f"questoes_uni{n}.md"
    out_path = dc.SAIDA / template_name

    if not pristine.exists():
        print(f"  ! template não encontrado: {pristine}")
        return {}
    if not md_path.exists():
        print(f"  ! markdown não encontrado: {md_path}")
        return {}

    md_text = md_path.read_text(encoding="utf-8")
    print(f"  renderizando fórmulas de UNI{n}...")
    formula_index = dc.build_formula_index(md_text)
    stats = dc.new_stats()

    doc = Document(pristine)

    # 1) Disciplina / Professor-conteudista
    dc.set_placeholder(doc, ("disciplina:",), dc.DISCIPLINA)
    dc.set_placeholder(doc, ("professor-conteudista:",), dc.CONTEUDISTA)

    # 2) remove orientações + exemplos do template (tudo após a linha
    #    "Professor-conteudista:")
    removed = dc.cut_body_after(
        doc, lambda t: t.strip().lower().startswith("professor-conteudista"),
        keep_matched=True,
    )

    # 3) seções do markdown
    secoes = dc.sections_dict(md_text)
    questoes_lines = secoes.get("Questões")
    gabarito_lines = secoes.get("Gabarito e feedbacks")
    if questoes_lines is None or gabarito_lines is None:
        print(f"  ! seções 'Questões'/'Gabarito e feedbacks' não encontradas em {md_path.name}")
        return {}

    # 4) divisor + questões (asterisco preservado na alternativa correta)
    p = dc.new_para(doc)
    dc.style_run(p.add_run("Questões"), bold=True)
    dc.render_blocks(doc, questoes_lines, formula_index, stats, first=False, alts=True)

    # 5) divisor + gabarito/feedbacks (200 linhas: 5 por questão x 40 questões)
    p = dc.new_para(doc)
    dc.style_run(p.add_run("Gabarito e feedbacks"), bold=True)
    feedback_count = dc.render_gabarito(doc, gabarito_lines, formula_index, stats)

    dc.unlock_rows_and_breaks(doc)
    dc.SAIDA.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    n_questoes_md = len(re.findall(r"(?m)^\*\*(\d+)\.\*\*", "\n".join(questoes_lines)))

    print(f"  OK -> {out_path.name}  (par. orientação/exemplo removidos: {removed}; "
          f"questões no md: {n_questoes_md}; linhas de feedback: {feedback_count}; "
          f"fórmulas: {dc.stats_total_imagens(stats)} ok / {len(stats['faltantes'])} faltantes)")
    if stats["faltantes"]:
        print("     fórmulas ausentes do índice (entraram como texto):",
              stats["faltantes"][:5], "..." if len(stats["faltantes"]) > 5 else "")

    return {
        "path": out_path,
        "pristine": pristine,
        "n_questoes_md": n_questoes_md,
        "feedback_count": feedback_count,
        "stats": stats,
    }


def main(argv):
    nums = [int(a) for a in argv[1:]] or [1, 2, 3, 4]
    print("QUESTIONÁRIOS (40 Questões - UNI1..4):", dc.DISCIPLINA)
    print("-" * 60)
    resultados = {}
    for n in nums:
        print(f"Unidade {n}:")
        resultados[n] = fill_unit(n)
    print("-" * 60)
    print("Pronto.")
    return resultados


if __name__ == "__main__":
    main(sys.argv)
