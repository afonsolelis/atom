#!/usr/bin/env python3
"""Audita o pacote institucional DOCX contra as regras dos modelos oficiais.

Verifica os 12 documentos de `entrega_docx/` — 4 unidades, 4 questionários,
avaliação final (mestre e estudante) e entrega de trabalho (mestre e estudante)
— contra o que os próprios modelos da UniFECAF determinam, e contra as regras
de `DIRETRIZES_PRODUCAO.md`.

Uso:
    tools/.venv/bin/python tools/validar_docx.py
    tools/.venv/bin/python tools/validar_docx.py --detalhe
"""

from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image

RAIZ = Path(__file__).resolve().parent.parent
SAIDA = RAIZ / "entrega_docx"
PRISTINE = RAIZ / "tools" / "_templates_pristine"
DISC_ARQ = "Model Based Design for Cyber-Physical Systems"
DISCIPLINA = "Model-Based Design for Cyber-Physical Systems"
CONTEUDISTA = "Afonso Cesar Lelis Brandão"

NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

FALHAS: list[str] = []
AVISOS: list[str] = []
OKS: list[str] = []


def falha(m: str) -> None:
    FALHAS.append(m)


def aviso(m: str) -> None:
    AVISOS.append(m)


def ok(m: str) -> None:
    OKS.append(m)


# --------------------------------------------------------------- utilidades

def texto_completo(doc: Document) -> str:
    """Texto de todo o documento: parágrafos soltos e conteúdo de tabelas.

    Percorre o XML do corpo, e não `doc.paragraphs`/`doc.tables`, porque os
    modelos das Unidades 3 e 4 escondem quase todas as caixas dentro de
    controles de conteúdo `w:sdt`, invisíveis para a API de alto nível.
    """
    partes = []
    for t in doc.element.body.iter():
        if t.tag.endswith("}t") and t.text:
            partes.append(t.text)
    return "\n".join(partes)


def tabelas_xml(doc: Document):
    return list(doc.element.body.iter(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"))


def texto_tabela(tbl) -> str:
    return " ".join(t.text or "" for t in tbl.iter() if t.tag.endswith("}t"))


def abrir(nome: str):
    p = SAIDA / nome
    if not p.exists():
        falha(f"{nome}: arquivo ausente")
        return None, None
    try:
        z = zipfile.ZipFile(p)
        if z.testzip() is not None:
            falha(f"{nome}: pacote ZIP corrompido")
            return None, None
        d = Document(p)
    except Exception as exc:  # noqa: BLE001
        falha(f"{nome}: não abre ({type(exc).__name__}: {exc})")
        return None, None
    return d, z


# ------------------------------------------------------- verificações comuns

# Frases de orientação ao conteudista que o modelo manda apagar. "Caro
# coordenador" NÃO entra aqui: é instrução dirigida ao coordenador, na caixa
# "Plano de Ensino", e deve permanecer.
RESIDUOS = [
    "XXXX", "A pessoa conteudista deve", "O conteudista deve",
    "Apagar as orientações", "elabore um título que gere curiosidade",
    "Exclusão de Exemplos", "EXEMPLO DE QUESTÃO", "EXEMPLO DE FEEDBACK",
    "Leia as Orientações", "SEGUIR À RISCA", "EXCLUIR AS ORIENTAÇÕES",
    "10 questões do tipo asserção-razão", "Nenhuma das anteriores",
    "Primeira alternativa", "insira o plano de ensino da disciplina aqui.  ",
]

# LaTeX que sobreviveu como texto em vez de virar imagem.
LATEX_CRU = [r"\frac", r"\mathrm", r"\begin{", r"\times", r"\approx", r"\cdot"]


def checar_comum(nome: str, doc: Document, z: zipfile.ZipFile,
                 pristine_nome: str) -> str:
    txt = texto_completo(doc)

    # 1) identificação preenchida (apenas nos modelos que possuem o campo)
    tem_campos = "Disciplina:" in Document(PRISTINE / pristine_nome).element.body.xml \
        if (PRISTINE / pristine_nome).exists() else True
    if tem_campos and DISCIPLINA not in txt:
        falha(f"{nome}: campo da disciplina não preenchido com '{DISCIPLINA}'")
    if tem_campos and CONTEUDISTA not in txt:
        falha(f"{nome}: campo do conteudista não preenchido")

    # 2) resíduos de orientação do modelo
    achados = [r for r in RESIDUOS if r in txt]
    if achados:
        falha(f"{nome}: resíduo de orientação do modelo: {achados}")

    # 3) LaTeX cru
    cru = [c for c in LATEX_CRU if c in txt]
    if cru:
        falha(f"{nome}: LaTeX não convertido em imagem: {cru}")
    dolares = txt.count("$")
    if dolares:
        falha(f"{nome}: {dolares} delimitador(es) '$' remanescente(s) no texto")

    # 4) tabelas do coordenador idênticas ao modelo pristino
    pp = PRISTINE / pristine_nome
    if pp.exists():
        pd = Document(pp)
        tp, tg = tabelas_xml(pd), tabelas_xml(doc)
        for idx, rotulo in ((0, "uso exclusivo do coordenador"), (1, "Data da Validação")):
            if idx >= len(tp) or idx >= len(tg):
                falha(f"{nome}: tabela do coordenador [{idx}] ausente")
                continue
            if texto_tabela(tp[idx]).split() != texto_tabela(tg[idx]).split():
                falha(f"{nome}: tabela do coordenador '{rotulo}' foi alterada")
    else:
        aviso(f"{nome}: modelo pristino não encontrado para comparação")

    # 5) imagens de fórmula em branco
    brancas, total_img = 0, 0
    for item in z.namelist():
        if not item.startswith("word/media/") or not item.lower().endswith(".png"):
            continue
        total_img += 1
        try:
            with z.open(item) as fh:
                im = Image.open(fh).convert("RGBA")
                px = im.load()
                visivel = any(px[x, y][3] > 10
                              for y in range(0, im.height, 3)
                              for x in range(0, im.width, 3))
                if not visivel:
                    brancas += 1
        except Exception:  # noqa: BLE001
            aviso(f"{nome}: imagem ilegível em {item}")
    if brancas:
        falha(f"{nome}: {brancas} de {total_img} imagens estão em branco")
    elif total_img:
        ok(f"{nome}: {total_img} imagens, nenhuma em branco")

    # 6) cabeçalho institucional preservado
    cabecalhos = [n for n in z.namelist() if re.match(r"word/header\d+\.xml", n)]
    if not cabecalhos:
        falha(f"{nome}: cabeçalho institucional ausente")

    return txt


# ------------------------------------------------------------------ unidades

def caixas_do_modelo(pristine_path: Path) -> list[str]:
    """Rótulos das caixas, lidos do próprio modelo pristino.

    Não dá para fixar a lista em código: o modelo numera as caixas com a
    numeração contínua da disciplina — a Unidade 3 traz "TEXTO BASE AULA 9" a
    "AULA 12", não "AULA 1" a "AULA 4" —, e as caixas de relação profissional e
    de AAI só existem em algumas unidades.
    """
    d = Document(pristine_path)
    rotulos = []
    for tbl in tabelas_xml(d):
        txt = texto_tabela(tbl).strip()
        m = re.match(r"^(TEXTO BASE AULA \d+|ROTEIRO VIDEOAULA \d+|QUIZ|AAI|"
                     r"DIRETO DA FONTE|PARA MERGULHAR|PODCAST|ARTIGO CIENTÍFICO|"
                     r"Relação da disciplina com atuação profissional)", txt)
        if m:
            rotulos.append(m.group(1))
    return rotulos


def caixas_esperadas(pristine_path: Path, n: int) -> list[str]:
    """Rótulos que a cópia gerada deve ter, já com a numeração contínua.

    O modelo da Unidade 2 é uma cópia defeituosa do da Unidade 1 e rotula suas
    caixas como "AULA 1" a "AULA 4", quando deveria trazer "AULA 5" a "AULA 8".
    O preenchimento corrige esses rótulos na cópia gerada; a expectativa aqui
    precisa seguir a mesma regra, senão a auditoria acusaria como falha
    justamente a correção.
    """
    primeira = (n - 1) * 4 + 1
    rotulos, k_aula, k_rot = [], 0, 0
    for bruto in caixas_do_modelo(pristine_path):
        if bruto.startswith("TEXTO BASE AULA"):
            rotulos.append(f"TEXTO BASE AULA {primeira + k_aula}")
            k_aula += 1
        elif bruto.startswith("ROTEIRO VIDEOAULA"):
            rotulos.append(f"ROTEIRO VIDEOAULA {primeira + k_rot}")
            k_rot += 1
        elif bruto.startswith("Relação da disciplina"):
            # o vídeo introdutório é único e pertence à Unidade 1
            rotulos.append("Relação da disciplina com atuação profissional" if n == 1
                           else "Relação da unidade com atuação profissional")
        else:
            rotulos.append(bruto)
    return rotulos


def validar_unidade(n: int) -> None:
    nome = f"TEMPLATE - Unidade {n}_{DISC_ARQ}.docx"
    doc, z = abrir(nome)
    if doc is None:
        return
    txt = checar_comum(nome, doc, z, nome)

    esperadas = caixas_esperadas(PRISTINE / nome, n)
    faltando = [c for c in esperadas if c.lower() not in txt.lower()]
    if faltando:
        falha(f"{nome}: caixas ausentes: {faltando}")
    else:
        ok(f"{nome}: {len(esperadas)} caixas do modelo presentes e preenchidas")

    # o número de tabelas precisa bater com o do modelo (nenhuma caixa perdida);
    # tabelas de conteúdo criadas a partir do Markdown podem acrescentar algumas
    n_prist = len(tabelas_xml(Document(PRISTINE / nome)))
    n_ger = len(tabelas_xml(doc))
    if n_ger < n_prist - 1:
        falha(f"{nome}: {n_ger} tabelas contra {n_prist} do modelo — caixa perdida")

    for sec in ("O que você verá", "Síntese da unidade", "Referências da unidade"):
        if sec.lower() not in txt.lower():
            falha(f"{nome}: seção obrigatória ausente: '{sec}'")

    if "Caro coordenador" not in txt:
        falha(f"{nome}: caixa 'Plano de Ensino' perdeu a instrução ao coordenador")
    if f"Plano de Ensino - Unidade {n}" not in txt:
        falha(f"{nome}: cabeçalho do plano de ensino não aponta para a Unidade {n}")

    tem_aai = "AAI" in txt or "Atividade avaliativa individual" in txt.lower()
    if n == 1 and not tem_aai:
        falha(f"{nome}: a Unidade 1 exige a AAI")

    # referências ABNT: ao menos uma entrada com autor em caixa alta e ano
    if not re.search(r"[A-ZÁÉÍÓÚÃÕÇ]{3,},\s+[A-Z]", txt):
        aviso(f"{nome}: nenhuma referência em formato ABNT reconhecida")

    if len(txt) < 60000:
        falha(f"{nome}: apenas {len(txt):,} caracteres — conteúdo aparentemente truncado")
    else:
        ok(f"{nome}: {len(txt):,} caracteres")


# ------------------------------------------------------------ questionários

def validar_questionario(n: int) -> None:
    nome = f"40 Questões - UNI{n}_{DISC_ARQ}.docx"
    doc, z = abrir(nome)
    if doc is None:
        return
    checar_comum(nome, doc, z, nome)

    ps = doc.paragraphs
    try:
        corte = next(i for i, p in enumerate(ps) if p.text.strip() == "Gabarito e feedbacks")
    except StopIteration:
        falha(f"{nome}: seção 'Gabarito e feedbacks' ausente")
        return

    atual, alts = None, {}
    for p in ps[:corte]:
        t = p.text.strip()
        m = re.match(r"^(\d+)\.(\s|$)", t)
        if m:
            atual = int(m.group(1))
            alts.setdefault(atual, [])
        elif atual is not None and re.match(r"^\*?[a-e]\.", t):
            tem_img = len(p._p.findall(f".//{NS_A}blip")) > 0
            mA = re.match(r"^(\*?)([a-e])\.", t)
            alts[atual].append(
                {"marcada": bool(mA.group(1)), "letra": mA.group(2),
                 "conteudo": tem_img or len(t) > 3})

    if len(alts) != 40:
        falha(f"{nome}: {len(alts)} questões (exigidas 40)")
    else:
        ok(f"{nome}: 40 questões")

    letras: dict[str, int] = {}
    for q in sorted(alts):
        v = alts[q]
        if len(v) != 5:
            falha(f"{nome} Q{q}: {len(v)} alternativas (exigidas 5)")
        marcadas = [a["letra"] for a in v if a["marcada"]]
        if len(marcadas) != 1:
            falha(f"{nome} Q{q}: {len(marcadas)} alternativas marcadas com '*' (exigida 1)")
        else:
            letras[marcadas[0]] = letras.get(marcadas[0], 0) + 1
        for a in v:
            if not a["conteudo"]:
                falha(f"{nome} Q{q}: alternativa '{a['letra']}' está vazia")
        letras_ordem = [a["letra"] for a in v]
        if letras_ordem != ["a", "b", "c", "d", "e"]:
            falha(f"{nome} Q{q}: alternativas fora da ordem a-e ({letras_ordem})")

    if letras and sorted(letras.values()) != [8, 8, 8, 8, 8]:
        aviso(f"{nome}: distribuição da correta {dict(sorted(letras.items()))} (alvo 8 por letra)")
    elif letras:
        ok(f"{nome}: asterisco em 40 questões, 8 por letra")

    gab = ps[corte:]
    cab = [p for p in gab if re.match(r"^Questão \d+ \(correta: [a-e]\)", p.text.strip())]
    fb = [p for p in gab if re.match(r"^[a-e]\.", p.text.strip())]
    if len(cab) != 40:
        falha(f"{nome}: {len(cab)} cabeçalhos de gabarito (exigidos 40)")
    if len(fb) != 200:
        falha(f"{nome}: {len(fb)} linhas de devolutiva (exigidas 200)")
    else:
        ok(f"{nome}: 200 devolutivas")

    # coerência entre o asterisco e a letra declarada no gabarito
    declaradas = {int(re.match(r"^Questão (\d+)", p.text.strip()).group(1)):
                  re.search(r"\(correta: ([a-e])\)", p.text).group(1) for p in cab}
    for q, v in alts.items():
        m = [a["letra"] for a in v if a["marcada"]]
        if m and q in declaradas and m[0] != declaradas[q]:
            falha(f"{nome} Q{q}: asterisco em '{m[0]}' mas gabarito declara '{declaradas[q]}'")


# --------------------------------------------------------- instrumentos

VAZAMENTOS = ["Resposta esperada", "Solução esperada", "Critérios de correção",
              "Rubrica", "critério de correção", "SOLUÇÃO:"]


def validar_avaliacao() -> None:
    base = f"Avaliação final_(10 discursivas)_{DISC_ARQ}"
    prist = f"{base}.docx"
    for sufixo in ("MESTRE", "ESTUDANTE"):
        nome = f"{base} - {sufixo}.docx"
        doc, z = abrir(nome)
        if doc is None:
            continue
        txt = checar_comum(nome, doc, z, prist)

        n_q = len(re.findall(r"(?m)^\s*Questão \d+", txt))
        if n_q < 10:
            falha(f"{nome}: {n_q} questões localizadas (exigidas 10)")
        else:
            ok(f"{nome}: {n_q} marcações de questão")

        vaz = [v for v in VAZAMENTOS if v.lower() in txt.lower()]
        if sufixo == "ESTUDANTE":
            if vaz:
                falha(f"{nome}: vazamento de conteúdo do tutor: {vaz}")
            else:
                ok(f"{nome}: sem respostas nem rubricas")
        else:
            if not vaz:
                falha(f"{nome}: versão mestre sem respostas esperadas")
            else:
                ok(f"{nome}: respostas do tutor presentes")


def validar_avaliacao_final() -> None:
    """Avaliação final no padrão Átomo 3.0 — o mesmo de
    `disciplinas/portos_aeroportos_e_ferrovias/`: documento único, 30 objetivas
    (15 asserção-razão + 15 interpretação) e 10 discursivas, com feedback para
    todas as 5 alternativas de cada objetiva, ao final do documento."""
    nome = f"Avaliação final_(30 questões múltipla escolha + 10 discursivas)_{DISC_ARQ}.docx"
    doc, z = abrir(nome)
    if doc is None:
        return
    txt = checar_comum(nome, doc, z, nome)

    n_obj = len(re.findall(r"(?m)^Questão \d+ \((?:Asserção-Razão|Interpretação)\)", txt))
    n_dis = len(re.findall(r"(?m)^Questão \d+ \(Discursiva\)", txt))
    if n_obj != 30 or n_dis != 10:
        falha(f"{nome}: {n_obj} objetivas e {n_dis} discursivas (exigidas 30 e 10)")
    else:
        ok(f"{nome}: 30 objetivas + 10 discursivas")

    n_ast = len(re.findall(r"(?m)^\*[a-e]\.", txt))
    if n_ast != 30:
        falha(f"{nome}: {n_ast} alternativas marcadas com '*' (exigidas 30, uma por objetiva)")
    else:
        ok(f"{nome}: gabarito marcado com asterisco nas 30 objetivas")

    # o feedback vai em lista com marcador; as alternativas das questões, não
    n_fb = len(re.findall(r"(?m)^•\s+[a-e]\.\s", txt))
    if n_fb != 150:
        falha(f"{nome}: {n_fb} linhas de feedback (exigidas 150)")
    else:
        ok(f"{nome}: feedback para as 150 alternativas")

    n_resp = txt.count("Resposta esperada")
    if n_resp != 10:
        falha(f"{nome}: {n_resp} respostas esperadas nas discursivas (exigidas 10)")
    else:
        ok(f"{nome}: as 10 discursivas com resposta esperada")

    for outra in ("Indústria 4.0", "Digitalização de Processos"):
        if outra in txt:
            falha(f"{nome}: resíduo de conteúdo de outra disciplina ('{outra}')")


def validar_pbl() -> None:
    base = f"TEMPLATE ENTREGA DE TRABALHO - {DISC_ARQ}"
    prist = f"{base}.docx"
    contagem = {}
    for sufixo in ("MESTRE", "ESTUDANTE"):
        nome = f"{base} - {sufixo}.docx"
        doc, z = abrir(nome)
        if doc is None:
            continue
        txt = checar_comum(nome, doc, z, prist)
        contagem[sufixo] = len(tabelas_xml(doc))

        for caixa in ("TÍTULO", "DESAFIO", "FONTE DE PESQUISA", "ENTREGÁVEL"):
            if caixa.lower() not in txt.lower():
                falha(f"{nome}: caixa '{caixa}' ausente")

        tem_solucao = bool(re.search(r"\d\.\s*SOLU[ÇC][ÃA]O\s*:", txt, re.I))
        if sufixo == "ESTUDANTE":
            if tem_solucao:
                falha(f"{nome}: a caixa SOLUÇÃO precisa ser removida da versão do estudante")
            vaz = [v for v in VAZAMENTOS if v.lower() in txt.lower()]
            if vaz:
                falha(f"{nome}: vazamento de conteúdo do tutor: {vaz}")
            else:
                ok(f"{nome}: sem solução nem rubrica")
        else:
            if not tem_solucao:
                falha(f"{nome}: versão mestre sem a caixa SOLUÇÃO")
            else:
                ok(f"{nome}: caixa SOLUÇÃO presente")

    if contagem.get("MESTRE") and contagem.get("ESTUDANTE"):
        if contagem["ESTUDANTE"] >= contagem["MESTRE"]:
            falha("PBL: a versão do estudante deveria ter uma tabela a menos "
                  f"(mestre={contagem['MESTRE']}, estudante={contagem['ESTUDANTE']})")
        else:
            ok(f"PBL: mestre {contagem['MESTRE']} tabelas, estudante "
               f"{contagem['ESTUDANTE']} (SOLUÇÃO removida)")


# ---------------------------------------------------------------------- main

def main(argv: list[str]) -> int:
    detalhe = "--detalhe" in argv

    esperados = (
        [f"TEMPLATE - Unidade {n}_{DISC_ARQ}.docx" for n in (1, 2, 3, 4)]
        + [f"40 Questões - UNI{n}_{DISC_ARQ}.docx" for n in (1, 2, 3, 4)]
        + [f"Avaliação final_(10 discursivas)_{DISC_ARQ} - {s}.docx"
           for s in ("MESTRE", "ESTUDANTE")]
        + [f"Avaliação final_(30 questões múltipla escolha + 10 discursivas)_{DISC_ARQ}.docx"]
        + [f"TEMPLATE ENTREGA DE TRABALHO - {DISC_ARQ} - {s}.docx"
           for s in ("MESTRE", "ESTUDANTE")]
    )
    presentes = sorted(p.name for p in SAIDA.glob("*.docx"))
    extras = sorted(set(presentes) - set(esperados))
    if extras:
        falha(f"arquivos não previstos no pacote institucional: {extras}")
    else:
        ok(f"pacote com exatamente {len(presentes)} documentos institucionais")

    for n in (1, 2, 3, 4):
        validar_unidade(n)
        validar_questionario(n)
    validar_avaliacao()
    validar_avaliacao_final()
    validar_pbl()

    print("=" * 78)
    print("AUDITORIA DO PACOTE INSTITUCIONAL DOCX")
    print("=" * 78)
    if detalhe:
        print(f"\nAPROVADOS ({len(OKS)}):")
        for m in OKS:
            print(f"  [ok]     {m}")
    if AVISOS:
        print(f"\nAVISOS ({len(AVISOS)}):")
        for m in AVISOS:
            print(f"  [aviso]  {m}")
    if FALHAS:
        print(f"\nFALHAS ({len(FALHAS)}):")
        for m in FALHAS:
            print(f"  [FALHA]  {m}")
    print("\n" + "-" * 78)
    print(f"Aprovados: {len(OKS)}  |  Avisos: {len(AVISOS)}  |  Falhas: {len(FALHAS)}")
    print("-" * 78)
    return 1 if FALHAS else 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
