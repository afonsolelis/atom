#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Preenche os templates institucionais Word (UniFECAF / Núcleo das Engenharias
e Tecnologia) da AVALIAÇÃO FINAL (10 discursivas) e da ENTREGA DE TRABALHO
(PBL) da disciplina Model-Based Design for Cyber-Physical Systems.

Cada instrumento é um arquivo-mestre de produção com "# Parte A — Versão do
estudante" e "# Parte B — Versão exclusiva do professor tutor" (respostas
esperadas / solução / rubrica). A divisão é feita casando o cabeçalho
"# Parte B" em início de linha (regex com MULTILINE) — NÃO pela primeira
ocorrência literal da string, que também aparece entre crases no bloco de
controle de versão no topo do arquivo. Usar a primeira ocorrência ingênua
cortaria o documento no lugar errado e deixaria a versão do estudante quase
vazia.

Gera sempre dois arquivos por instrumento:
  - "... - MESTRE.docx"    -> Parte A + Parte B completas.
  - "... - ESTUDANTE.docx" -> apenas Parte A (nenhuma resposta, rubrica ou
    critério de correção; na Entrega de Trabalho, a caixa [7] SOLUÇÃO é
    removida por inteiro, não apenas esvaziada).

Sempre parte do backup pristine em tools/_templates_pristine/ e escreve em
entrega_docx/, sobrescrevendo a cada execução.

Uso:
    tools/.venv/bin/python tools/preencher_instrumentos.py            # os dois
    tools/.venv/bin/python tools/preencher_instrumentos.py avaliacao  # só a avaliação final
    tools/.venv/bin/python tools/preencher_instrumentos.py entrega    # só a entrega de trabalho
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
import docx_comum as dc


# ============================================================ Avaliação final (10 discursivas)
def build_avaliacao(doc, secoes_a, secoes_b, formula_index, incluir_parte_b: bool) -> dict:
    stats = dc.new_stats()

    dc.set_placeholder(doc, ("disciplina:",), dc.DISCIPLINA)
    dc.set_placeholder(doc, ("professor-conteudista:",), dc.CONTEUDISTA)
    removed = dc.cut_body_after(
        doc, lambda t: t.strip().lower().startswith("professor-conteudista"),
        keep_matched=True,
    )

    orientacoes = secoes_a.get("Orientações")
    questoes = secoes_a.get("Questões")

    if orientacoes:
        p = dc.new_para(doc)
        dc.style_run(p.add_run("Orientações"), bold=True)
        dc.render_blocks(doc, orientacoes, formula_index, stats, first=False)

    if questoes:
        p = dc.new_para(doc)
        dc.style_run(p.add_run("Questões"), bold=True)
        dc.render_blocks(doc, questoes, formula_index, stats, first=False)

    if incluir_parte_b:
        respostas = secoes_b.get("Respostas esperadas e critérios de correção")
        p = dc.new_para(doc)
        dc.style_run(p.add_run(
            "NÃO DISTRIBUIR AOS ESTUDANTES — uso exclusivo do professor tutor."),
            bold=True, italic=True)
        p = dc.new_para(doc)
        dc.style_run(p.add_run("Respostas esperadas e critérios de correção"), bold=True)
        if respostas:
            dc.render_blocks(doc, respostas, formula_index, stats, first=False)

    dc.unlock_rows_and_breaks(doc)
    return {"stats": stats, "orientacao_removida": removed}


def fill_avaliacao() -> dict:
    template_name = f"Avaliação final_(10 discursivas)_{dc.DISCIPLINA_ARQUIVO}.docx"
    pristine = dc.PRISTINE / template_name
    md_path = dc.DISC / "instrumentos_avaliativos" / "avaliacao_dissertativa.md"

    if not pristine.exists():
        print(f"  ! template não encontrado: {pristine}")
        return {}
    if not md_path.exists():
        print(f"  ! markdown não encontrado: {md_path}")
        return {}

    md_text = md_path.read_text(encoding="utf-8")
    parte_a_text, parte_b_text = dc.split_parte_b(md_text)
    secoes_a = dc.sections_dict(parte_a_text)
    secoes_b = dc.sections_dict(parte_b_text) if parte_b_text else {}

    print("  renderizando fórmulas da avaliação final...")
    formula_index = dc.build_formula_index(md_text)

    base = pristine.stem
    out_mestre = dc.SAIDA / f"{base} - MESTRE.docx"
    out_estudante = dc.SAIDA / f"{base} - ESTUDANTE.docx"

    doc_mestre = Document(pristine)
    r_mestre = build_avaliacao(doc_mestre, secoes_a, secoes_b, formula_index, incluir_parte_b=True)
    dc.SAIDA.mkdir(parents=True, exist_ok=True)
    doc_mestre.save(out_mestre)

    doc_estudante = Document(pristine)
    r_estudante = build_avaliacao(doc_estudante, secoes_a, secoes_b, formula_index, incluir_parte_b=False)
    doc_estudante.save(out_estudante)

    print(f"  OK -> {out_mestre.name}  (fórmulas: {dc.stats_total_imagens(r_mestre['stats'])} ok / "
          f"{len(r_mestre['stats']['faltantes'])} faltantes)")
    print(f"  OK -> {out_estudante.name}  (fórmulas: {dc.stats_total_imagens(r_estudante['stats'])} ok / "
          f"{len(r_estudante['stats']['faltantes'])} faltantes)")

    return {
        "pristine": pristine,
        "mestre": {"path": out_mestre, **r_mestre},
        "estudante": {"path": out_estudante, **r_estudante},
    }


# ============================================================ Entrega de trabalho (PBL)
BOX_MAP = [
    ("título", "1. Título"),
    ("desafio", "2. Desafio"),
    ("fonte de pesquisa", "3. Fontes de pesquisa"),
    ("entregável", "4. Componentes avaliativos, submissão e pontuação"),
]

ENTREGA_ORIENTACAO_PREFIXOS = (
    "atenção:", "excluir as orientações", "todos os itens são obrigatório",
)


def build_entrega(doc, secoes_a, solucao_lines, formula_index, estudante: bool) -> dict:
    stats = dc.new_stats()
    filled = []

    for needle, titulo in BOX_MAP:
        t = dc.find_table(doc, needle)
        lines = secoes_a.get(titulo)
        if t is not None and lines is not None:
            dc.render_blocks(t.rows[1].cells[0], lines, formula_index, stats, first=True)
            filled.append(titulo)
        elif t is None:
            print(f"    ! caixa '{needle}' não encontrada no template")
        elif lines is None:
            print(f"    ! seção '{titulo}' não encontrada no markdown")

    t_solucao = dc.find_table(doc, "solução")
    if estudante:
        if t_solucao is not None:
            dc.remove_table(doc, t_solucao)
            filled.append("solução-REMOVIDA (versão estudante)")
        else:
            print("    ! caixa 'solução' já ausente ao tentar remover")
    else:
        if t_solucao is not None and solucao_lines is not None:
            dc.render_blocks(t_solucao.rows[1].cells[0], solucao_lines, formula_index, stats, first=True)
            filled.append("6. Solução (mestre)")
        elif t_solucao is None:
            print("    ! caixa 'solução' não encontrada no template (mestre)")
        elif solucao_lines is None:
            print("    ! seção 'Solução esperada e critérios de correção' não encontrada no markdown")

    orient_removidos = dc.remove_paragraphs_with_prefix(doc, ENTREGA_ORIENTACAO_PREFIXOS)

    roteiro_lines = secoes_a.get("Roteiro do estudante")
    roteiro_removidos = dc.cut_body_after(
        doc, lambda t: t.strip().lower() == "roteiro do estudante", keep_matched=True,
    )
    if roteiro_lines:
        dc.render_blocks(doc, roteiro_lines, formula_index, stats, first=False)
        filled.append("roteiro do estudante")

    dc.unlock_rows_and_breaks(doc)
    return {
        "stats": stats, "filled": filled,
        "orientacao_removida": orient_removidos, "roteiro_antigo_removido": roteiro_removidos,
    }


def fill_entrega() -> dict:
    template_name = f"TEMPLATE ENTREGA DE TRABALHO - {dc.DISCIPLINA_ARQUIVO}.docx"
    pristine = dc.PRISTINE / template_name
    md_path = dc.DISC / "instrumentos_avaliativos" / "entrega_trabalho.md"

    if not pristine.exists():
        print(f"  ! template não encontrado: {pristine}")
        return {}
    if not md_path.exists():
        print(f"  ! markdown não encontrado: {md_path}")
        return {}

    md_text = md_path.read_text(encoding="utf-8")
    parte_a_text, parte_b_text = dc.split_parte_b(md_text)
    secoes_a = dc.sections_dict(parte_a_text)
    secoes_b = dc.sections_dict(parte_b_text) if parte_b_text else {}
    solucao_lines = secoes_b.get("Solução esperada e critérios de correção")

    print("  renderizando fórmulas da entrega de trabalho...")
    formula_index = dc.build_formula_index(md_text)

    base = pristine.stem
    out_mestre = dc.SAIDA / f"{base} - MESTRE.docx"
    out_estudante = dc.SAIDA / f"{base} - ESTUDANTE.docx"

    doc_mestre = Document(pristine)
    r_mestre = build_entrega(doc_mestre, secoes_a, solucao_lines, formula_index, estudante=False)
    dc.SAIDA.mkdir(parents=True, exist_ok=True)
    doc_mestre.save(out_mestre)

    doc_estudante = Document(pristine)
    r_estudante = build_entrega(doc_estudante, secoes_a, solucao_lines, formula_index, estudante=True)
    doc_estudante.save(out_estudante)

    print(f"  OK -> {out_mestre.name}  (seções: {', '.join(r_mestre['filled'])}; "
          f"fórmulas: {dc.stats_total_imagens(r_mestre['stats'])} ok / {len(r_mestre['stats']['faltantes'])} faltantes)")
    print(f"  OK -> {out_estudante.name}  (seções: {', '.join(r_estudante['filled'])}; "
          f"fórmulas: {dc.stats_total_imagens(r_estudante['stats'])} ok / {len(r_estudante['stats']['faltantes'])} faltantes)")

    return {
        "pristine": pristine,
        "mestre": {"path": out_mestre, **r_mestre},
        "estudante": {"path": out_estudante, **r_estudante},
    }


def main(argv):
    what = argv[1] if len(argv) > 1 else "all"
    print("INSTRUMENTOS AVALIATIVOS:", dc.DISCIPLINA)
    print("-" * 60)
    resultados = {}
    if what in ("all", "avaliacao"):
        print("Avaliação final (10 discursivas):")
        resultados["avaliacao"] = fill_avaliacao()
    if what in ("all", "entrega"):
        print("Entrega de trabalho (PBL):")
        resultados["entrega"] = fill_entrega()
    print("-" * 60)
    print("Pronto.")
    return resultados


if __name__ == "__main__":
    main(sys.argv)
