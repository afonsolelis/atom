#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Preenche os TEMPLATES Word institucionais (UniFECAF / Nucleo das Engenharias e
Tecnologia) das quatro unidades de "Model-Based Design for Cyber-Physical
Systems" com o conteudo dos markdowns correspondentes.

Trabalha sempre a partir das copias pristinas em tools/_templates_pristine/,
para que o script seja reexecutavel a vontade. Saida em entrega_docx/,
sobrescrita a cada execucao.

Uso:
    tools/.venv/bin/python tools/preencher_unidades.py        # todas as unidades
    tools/.venv/bin/python tools/preencher_unidades.py 1 3    # so 1 e 3

Diferencas de layout desta disciplina frente a data_engineering_and_pipelines
(cuja implementacao tools/preencher_docx.py serviu de referencia):
    - roteiros das videoaulas vivem em unidade_N/roteiros_20min.md (arquivo
      proprio), nao embutidos em unidade_N.md;
    - o roteiro do video introdutorio so existe (e so entra na caixa) na
      Unidade 1, em roteiro_video_introdutorio.md;
    - Quiz/AAI/Sintese/Material complementar vivem como subsecoes (### e
      ####) dentro de uma unica secao de nivel 2
      "## Atividades, sintese e material complementar", nao como secoes de
      nivel 2 separadas;
    - a AAI so existe (no conteudo e na caixa do template) nas Unidades 1 e 2;
      na Unidade 2 a caixa recebe uma nota explicita, no lugar de uma AAI
      inventada.
"""
from __future__ import annotations

import io
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.table import Table

TOOLS = Path(__file__).resolve().parent
DISC = TOOLS.parent
sys.path.insert(0, str(TOOLS))
import formulas as fm  # coletar_formulas / renderizar, cache em tools/_cache_formulas/

DISCIPLINA = "Model-Based Design for Cyber-Physical Systems"
# Nome exato usado nos arquivos do template (sem hífen) — difere do nome
# "oficial" da disciplina usado no corpo do texto.
DISCIPLINA_ARQUIVO = "Model Based Design for Cyber-Physical Systems"
CONTEUDISTA = "Afonso Cesar Lelis Brandão"
PRISTINE = TOOLS / "_templates_pristine"
OUT_DIR = DISC / "entrega_docx"

AVISOS: list[str] = []


def avisar(msg: str) -> None:
    AVISOS.append(msg)
    print("  [aviso]", msg)


# ============================================================ fórmulas
def build_formula_index(*textos: str) -> dict:
    achadas: list[tuple[str, bool]] = []
    vistas = set()
    for texto in textos:
        for tex, disp in fm.coletar_formulas(texto):
            if (tex, disp) not in vistas:
                vistas.add((tex, disp))
                achadas.append((tex, disp))
    if not achadas:
        return {}
    return fm.renderizar(achadas)


def _formula_png_bytes(meta: dict) -> bytes:
    return meta["png"].read_bytes()


def _latex_legivel(latex: str) -> str:
    """Converte a notação LaTeX de vírgula decimal '{,}' para ',' comum, para
    uso em texto alternativo (acessibilidade) da imagem da fórmula."""
    return latex.replace("{,}", ",")


def _set_pic_alt(run, latex: str) -> None:
    """Define o texto alternativo (descr) da imagem da fórmula recém-inserida
    no run, com o LaTeX de origem em forma legível — acessibilidade e também
    o que torna a fórmula localizável ao converter o DOCX para texto puro."""
    drawing = run._r.find(qn("w:drawing"))
    if drawing is None:
        return
    inline = drawing.find(qn("wp:inline"))
    if inline is None:
        inline = drawing.find(qn("wp:anchor"))
    if inline is None:
        return
    docPr = inline.find(qn("wp:docPr"))
    if docPr is not None:
        alt = _latex_legivel(latex)
        docPr.set("descr", alt[:500])
        docPr.set("name", "Fórmula: " + alt[:50])


# ============================================================ helpers docx genéricos
def clear_container(cell) -> None:
    """Remove tudo (paragrafos, tabelas aninhadas, sdt de orientacao) menos w:tcPr,
    e destrava a linha (cantSplit/tblHeader) para o Word nao empurrar a caixa
    inteira para a pagina seguinte quando o conteudo e longo."""
    tc = cell._tc
    tr = tc.getparent()
    if tr is not None and tr.tag == qn("w:tr"):
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None:
            for tag in ("w:cantSplit", "w:tblHeader"):
                el = trPr.find(qn(tag))
                if el is not None:
                    trPr.remove(el)
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)


def ensure_trailing_p(cell) -> None:
    tc = cell._tc
    if len(tc) == 0 or tc[-1].tag != qn("w:p"):
        cell.add_paragraph("")


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rpr.append(sz)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), "0563C1")
    rpr.append(col)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


def style_body_paragraph(p) -> None:
    """Aplica a formatacao de corpo exigida pelo template: Times New Roman 12,
    espacamento 1,15, sem espaco apos paragrafo, alinhado a esquerda. So se
    aplica a paragrafos de estilo 'normal' (paragrafos de titulo/heading
    mantem o estilo proprio do template, ja correto)."""
    if p.style is None or (p.style.name or "").lower() not in ("normal", "default paragraph style"):
        return
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    if pf.alignment is None:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def mkpara(container, style=None):
    p = container.add_paragraph(style=style)
    style_body_paragraph(p)
    return p


def mkrun(p, text: str = "", mono: bool = False):
    r = p.add_run(text)
    if mono:
        r.font.name = "Consolas"
        r.font.size = Pt(10)
    else:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return r


INLINE_RE = re.compile(
    r"\*\*(?P<b>.+?)\*\*"
    r"|`(?P<c>[^`]+)`"
    r"|\[(?P<lt>[^\]]+)\]\((?P<lu>[^)]+)\)"
    r"|(?<![\*\w])\*(?P<i>[^*]+?)\*(?![\*\w])"
)


def add_md_runs(p, text: str) -> None:
    """Adiciona runs com **negrito**, `codigo`, [link](url), *italico* (palavras
    estrangeiras), em Times New Roman 12."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            mkrun(p, text[pos:m.start()])
        if m.group("b") is not None:
            mkrun(p, m.group("b")).bold = True
        elif m.group("c") is not None:
            mkrun(p, m.group("c"), mono=True)
        elif m.group("lt") is not None:
            add_hyperlink(p, m.group("lu"), m.group("lt"))
        elif m.group("i") is not None:
            mkrun(p, m.group("i")).italic = True
        pos = m.end()
    if pos < len(text):
        mkrun(p, text[pos:])


INLINE_MATH_RE = re.compile(r"(?<!\$)\$([^$\n]+)\$(?!\$)")


def add_inline(p, text: str, formula_idx: dict) -> None:
    """Como add_md_runs, mas tratando $formula$ inline como imagem (ou, na
    ausencia no indice, como texto LaTeX, sem quebrar)."""
    pos = 0
    for m in INLINE_MATH_RE.finditer(text):
        if m.start() > pos:
            add_md_runs(p, text[pos:m.start()])
        latex = m.group(1).strip()
        meta = formula_idx.get((latex, False))
        if meta:
            r = p.add_run()
            w = Pt(min(meta["largura_pt"], 460))
            h = Pt(meta["altura_pt"] * (min(meta["largura_pt"], 460) / meta["largura_pt"]))
            r.add_picture(io.BytesIO(_formula_png_bytes(meta)), width=w, height=h)
            _set_pic_alt(r, latex)
        else:
            avisar(f"formula inline fora do indice, inserida como texto: {latex[:60]!r}")
            add_md_runs(p, f"${latex}$")
        pos = m.end()
    if pos < len(text):
        add_md_runs(p, text[pos:])


def shade_paragraph(p, hexcolor: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def set_table_full_width(t, ncols: int) -> None:
    tblPr = t._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "autofit")
    col_w = Inches(6.0 / ncols)
    for row in t.rows:
        for c in row.cells:
            c.width = col_w


def add_markdown_table(cell, tbl_lines: list[str], formula_idx: dict) -> None:
    rows = []
    for k, ln in enumerate(tbl_lines):
        if k == 1:
            continue  # linha separadora |---|
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = cell.add_table(rows=len(rows), cols=ncols)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    set_table_full_width(t, ncols)
    for ri, r in enumerate(rows):
        for ci in range(ncols):
            tc = t.rows[ri].cells[ci]
            tc.text = ""
            p = tc.paragraphs[0]
            style_body_paragraph(p)
            val = r[ci] if ci < len(r) else ""
            add_inline(p, val, formula_idx)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    mkpara(cell)


# ============================================================ renderer de blocos markdown
def render_blocks(cell, lines: list[str], formula_idx: dict, base_label: str = "",
                   first: bool = False, alts: bool = False) -> None:
    """Renderiza linhas de Markdown dentro de `cell`. Se first=True, limpa a
    celula antes. Se alts=True, trata linhas 'a.'/'*a.' como alternativas de
    questao, destacando a correta com '*'."""
    if first:
        clear_container(cell)
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if s == "":
            i += 1
            continue

        # --- heading
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            style = "Heading %d" % min(max(level, 2), 4)
            p = cell.add_paragraph(style=style)
            add_inline(p, m.group(2), formula_idx)
            i += 1
            continue

        # --- regra horizontal
        if re.match(r"^-{3,}$", s):
            i += 1
            continue

        # --- bloco de código
        if s.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = cell.add_paragraph()
            style_body_paragraph(p)
            shade_paragraph(p, "F2F2F2")
            for k, cl in enumerate(code):
                r = mkrun(p, cl, mono=True)
                if k < len(code) - 1:
                    r.add_break()
            continue

        # --- fórmula de bloco ($$ ... $$)
        if s == "$$":
            i += 1
            expr = []
            while i < n and lines[i].strip() != "$$":
                expr.append(lines[i])
                i += 1
            i += 1
            latex = "\n".join(expr).strip()
            meta = formula_idx.get((latex, True))
            p = cell.add_paragraph()
            style_body_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if meta:
                w = min(meta["largura_pt"], 460)
                h = meta["altura_pt"] * (w / meta["largura_pt"])
                r = p.add_run()
                r.add_picture(io.BytesIO(_formula_png_bytes(meta)), width=Pt(w), height=Pt(h))
                _set_pic_alt(r, latex)
            else:
                avisar(f"formula em bloco fora do indice, inserida como texto: {latex[:60]!r}")
                add_md_runs(p, latex)
            continue

        # --- fórmula de bloco delimitada por linhas soltas com "$"
        #
        # Esta é a forma que o material da disciplina usa (o `$$ ... $$` acima
        # praticamente não ocorre). Sem este ramo, as três linhas saíam como
        # três parágrafos de texto — "$", o LaTeX cru e "$" —, que foi o defeito
        # encontrado na auditoria dos DOCX.
        #
        # A junção precisa ser por espaço, e não por quebra de linha, para casar
        # exatamente com a chave gerada em `formulas.coletar_formulas`, que usa
        # `" ".join(buffer)` neste caso.
        if s == "$":
            i += 1
            expr = []
            while i < n and lines[i].strip() != "$":
                expr.append(lines[i])
                i += 1
            i += 1
            latex = " ".join(x.strip() for x in expr).strip()
            meta = formula_idx.get((latex, True))
            p = cell.add_paragraph()
            style_body_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if meta:
                w = min(meta["largura_pt"], 460)
                h = meta["altura_pt"] * (w / meta["largura_pt"])
                r = p.add_run()
                r.add_picture(io.BytesIO(_formula_png_bytes(meta)), width=Pt(w), height=Pt(h))
                _set_pic_alt(r, latex)
            else:
                avisar(f"formula em bloco fora do indice, inserida como texto: {latex[:60]!r}")
                add_md_runs(p, latex)
            continue

        # --- blockquote ("Recurso visual N" — descrição para a equipe de edição)
        if s.startswith(">"):
            quote = []
            while i < n and (lines[i].strip().startswith(">") or lines[i].strip() == ""):
                if lines[i].strip() == "":
                    if i + 1 < n and lines[i + 1].strip().startswith(">"):
                        quote.append("")
                        i += 1
                        continue
                    break
                quote.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            label = cell.add_paragraph()
            style_body_paragraph(label)
            shade_paragraph(label, "FCE4D6")
            lr = mkrun(label, "RECURSO VISUAL — PLACEHOLDER PARA A EQUIPE DE EDIÇÃO (não é uma imagem existente)")
            lr.bold = True
            lr.font.size = Pt(10)
            para_buf: list[str] = []

            def flush_quote():
                if para_buf:
                    p = cell.add_paragraph()
                    style_body_paragraph(p)
                    shade_paragraph(p, "FCE4D6")
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.right_indent = Inches(0.2)
                    add_inline(p, " ".join(para_buf), formula_idx)
                    for r in p.runs:
                        r.italic = True
                    para_buf.clear()

            for ql in quote:
                qs = ql.strip()
                if qs == "":
                    flush_quote()
                else:
                    para_buf.append(qs)
            flush_quote()
            continue

        # --- tabela
        if s.startswith("|") and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(cell, tbl_lines, formula_idx)
            continue

        # --- imagem isolada (defensivo — este conteúdo não usa imagens reais)
        mi = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", s)
        if mi:
            alt = mi.group(1)
            p = cell.add_paragraph()
            style_body_paragraph(p)
            r = mkrun(p, "[imagem: %s]" % alt)
            r.italic = True
            avisar(f"imagem markdown encontrada e não migrada (fora de escopo): {alt!r}")
            i += 1
            continue

        # --- alternativa de questão (a. / *a.) — preserva o '*' da correta
        if alts:
            malt = re.match(r"^(\*?)([a-eA-E])\.\s+(.*)$", s)
            if malt:
                star, letter, rest = malt.group(1), malt.group(2), malt.group(3)
                p = cell.add_paragraph()
                style_body_paragraph(p)
                if star:
                    rs = mkrun(p, "* ")
                    rs.bold = True
                    rs.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    rl = mkrun(p, letter + ". ")
                    rl.bold = True
                else:
                    mkrun(p, letter + ". ")
                add_inline(p, rest, formula_idx)
                i += 1
                continue

        # --- lista (marcadores / numerada)
        if re.match(r"^[-*]\s+", s) or re.match(r"^\d+\.\s+", s):
            while i < n and (re.match(r"^\s*[-*]\s+", lines[i]) or re.match(r"^\s*\d+\.\s+", lines[i])
                              or (lines[i].strip() and lines[i].startswith("  "))):
                raw = lines[i]
                indent = len(raw) - len(raw.lstrip(" "))
                item = raw.strip()
                mnum = re.match(r"^(\d+)\.\s+(.*)$", item)
                mbul = re.match(r"^[-*]\s+(.*)$", item)
                p = cell.add_paragraph()
                style_body_paragraph(p)
                p.paragraph_format.left_indent = Inches(0.25 + (0.25 if indent >= 2 else 0))
                if mbul:
                    mkrun(p, "•  ")
                    add_inline(p, mbul.group(1), formula_idx)
                elif mnum:
                    mkrun(p, mnum.group(1) + ".  ")
                    add_inline(p, mnum.group(2), formula_idx)
                else:
                    add_inline(p, item, formula_idx)
                i += 1
            continue

        # --- parágrafo normal
        p = cell.add_paragraph()
        style_body_paragraph(p)
        add_inline(p, s, formula_idx)
        i += 1

    if hasattr(cell, "_tc"):
        ensure_trailing_p(cell)


# ============================================================ localizar caixas no docx
def _box_header(el) -> str | None:
    tr = el.find(qn("w:tr"))
    if tr is None:
        return None
    tc = tr.find(qn("w:tc"))
    if tc is None:
        return None
    return "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip().lower()


def find_table(doc, needle: str):
    """Busca em TODAS as w:tbl do corpo (inclusive as embrulhadas em
    content-controls w:sdt — nas Unidades 3 e 4 quase todas as caixas estão
    escondidas em sdt, invisíveis a doc.tables/iter_inner_content). Casa por
    substring no texto da 1ª célula da 1ª linha."""
    needle = needle.lower()
    for el in doc.element.body.iter(qn("w:tbl")):
        h = _box_header(el)
        if h and needle in h:
            return Table(el, doc)
    return None


def find_box_num(doc, prefix: str, number: int):
    """Acha a caixa '<prefix> ... <number>' com número EXATO (não seguido de
    outro dígito — senão 'aula 1' casaria com 'aula 10')."""
    pat = re.compile(r"%s\s*0*%d(?!\d)" % (re.escape(prefix.lower()), number))
    for el in doc.element.body.iter(qn("w:tbl")):
        h = _box_header(el)
        if h and pat.search(h):
            return Table(el, doc)
    return None


def set_roteiro_titulo(table, titulo: str) -> None:
    """Substitui 'Título do vídeo (elabore ...)' pelo título real, preservando
    o prefixo em negrito 'ROTEIRO VIDEOAULA N: ' que já está correto."""
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    if len(p.runs) >= 2:
        p.runs[1].text = titulo
        for extra in p.runs[2:]:
            extra.text = ""
    elif len(p.runs) == 1:
        # fallback defensivo: não deveria ocorrer, mas evita perder o cabeçalho
        txt = p.runs[0].text
        novo = re.sub(r":\s*Título do vídeo.*$", ": " + titulo, txt)
        p.runs[0].text = novo


# ============================================================ parsing dos markdowns
def sections_at_level(lines: list[str], level: int):
    """Divide `lines` em (titulo, corpo) por cabeçalhos '#'*level + espaço.
    O conteúdo antes do primeiro cabeçalho desse nível vira (None, corpo)."""
    marker = "#" * level + " "
    out = []
    title, buf = None, []
    for ln in lines:
        if ln.startswith(marker) and not ln.startswith(marker + "#"):
            out.append((title, buf))
            title = ln[len(marker):].strip()
            buf = []
        else:
            buf.append(ln)
    out.append((title, buf))
    return out


def get_section(sections, predicate):
    for title, body in sections:
        if title is not None and predicate(title.lower()):
            return title, body
    return None, None


def strip_trailing_blank(lines: list[str]) -> list[str]:
    while lines and lines[-1].strip() == "":
        lines.pop()
    while lines and lines[0].strip() == "":
        lines.pop(0)
    return lines


# ============================================================ preenchimento por unidade
AULA_HEADING_RE = re.compile(r"^Aula\s+(\d+)\s*[—-]\s*(.*)$", re.IGNORECASE)
ROTEIRO_HEADING_RE = re.compile(r"^Roteiro da Videoaula\s+(\d+)\s*[—-]\s*\"?(.*?)\"?$", re.IGNORECASE)


def parse_unidade_md(nnum: int):
    md_path = DISC / f"unidade_{nnum}/unidade_{nnum}.md"
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    top = sections_at_level(lines, 2)

    _, relacao_body = get_section(top, lambda t: "relação da unidade com a atuação profissional" in t)
    if relacao_body is None:
        avisar(f"Unidade {nnum}: seção 'Relação da unidade...' não encontrada no markdown")
        relacao_body = []

    _, overview_body = get_section(top, lambda t: t.strip() == "o que você verá nesta unidade")
    if overview_body is None:
        avisar(f"Unidade {nnum}: seção 'O que você verá nesta unidade' não encontrada no markdown")
        overview_body = []

    _, referencias_body = get_section(top, lambda t: t.strip() == "referências da unidade")
    if referencias_body is None:
        avisar(f"Unidade {nnum}: seção 'Referências da unidade' não encontrada no markdown")
        referencias_body = []

    aulas = {}  # numero_absoluto -> (titulo, body_sem_roteiro)
    for title, body in top:
        if title is None:
            continue
        m = AULA_HEADING_RE.match(title.strip())
        if not m:
            continue
        num = int(m.group(1))
        titulo_aula = m.group(2).strip()
        subs = sections_at_level(body, 3)
        kept_lines: list[str] = []
        for sub_t, sub_body in subs:
            if sub_t is not None and re.match(r"^Roteiro da Videoaula\s+\d+", sub_t.strip(), re.IGNORECASE):
                continue  # remissão a roteiros_20min.md — não entra na caixa de texto-base
            if sub_t is None:
                kept_lines.extend(sub_body)
            else:
                kept_lines.append(f"### {sub_t}")
                kept_lines.append("")
                kept_lines.extend(sub_body)
        aulas[num] = (titulo_aula, strip_trailing_blank(kept_lines))

    _, ativ_body = get_section(top, lambda t: "atividades" in t and "material complementar" in t)
    quiz_body, aai_body, material_body, sintese_body = [], None, None, []
    if ativ_body is not None:
        subs3 = sections_at_level(ativ_body, 3)
        for sub_t, sub_body in subs3:
            if sub_t is None:
                continue
            tl = sub_t.lower()
            if "quiz" in tl:
                quiz_body = sub_body
            elif "aai" in tl or "atividade avaliativa individual" in tl:
                aai_body = sub_body
            elif tl.strip() == "material complementar":
                material_body = sub_body
            elif "síntese da unidade" in tl:
                sintese_body = sub_body
    else:
        avisar(f"Unidade {nnum}: seção 'Atividades, síntese e material complementar' não encontrada")

    if not sintese_body:
        avisar(f"Unidade {nnum}: seção 'Síntese da unidade' não encontrada no markdown")
        sintese_body = []

    material_items = {}
    if material_body:
        subs4 = sections_at_level(material_body, 4)
        for sub_t, sub_body in subs4:
            if sub_t is not None:
                material_items[sub_t.lower()] = sub_body

    return {
        "relacao": relacao_body,
        "overview": strip_trailing_blank(list(overview_body)),
        "aulas": aulas,
        "quiz": quiz_body,
        "aai": aai_body,
        "material": material_items,
        "sintese": strip_trailing_blank(list(sintese_body)),
        "referencias": strip_trailing_blank(list(referencias_body)),
        "full_text": text,
    }


def parse_roteiros_md(nnum: int):
    path = DISC / f"unidade_{nnum}/roteiros_20min.md"
    text = path.read_text(encoding="utf-8")
    lines = text.split("\n")
    top = sections_at_level(lines, 2)
    roteiros = {}
    for title, body in top:
        if title is None:
            continue
        m = ROTEIRO_HEADING_RE.match(title.strip())
        if not m:
            continue
        num = int(m.group(1))
        titulo = m.group(2).strip()
        roteiros[num] = (titulo, strip_trailing_blank(list(body)))
    return roteiros, text


def parse_roteiro_intro():
    path = DISC / "roteiro_video_introdutorio.md"
    text = path.read_text(encoding="utf-8")
    lines = text.split("\n")
    top = sections_at_level(lines, 2)
    _, narracao = get_section(top, lambda t: t.strip() == "narração")
    return strip_trailing_blank(list(narracao or [])), text


# ============================================================ quiz / aai / material
QUESTAO_RE = re.compile(r"^\*\*Quest(ã|a)o\s*\d+\.\*\*", re.IGNORECASE)
FEEDBACK_RE = re.compile(r"^\*Feedback", re.IGNORECASE)


def split_quiz(lines: list[str]):
    blocos = []
    cur: list[str] = []
    for ln in lines:
        if QUESTAO_RE.match(ln.strip()):
            if cur:
                blocos.append(cur)
            cur = [ln]
        else:
            cur.append(ln)
    if cur:
        blocos.append(cur)
    perguntas, respostas = [], []
    for bloco in blocos:
        q_part, a_part, in_ans = [], [], False
        for ln in bloco:
            if FEEDBACK_RE.match(ln.strip()):
                in_ans = True
            (a_part if in_ans else q_part).append(ln)
        perguntas.extend(q_part)
        perguntas.append("")
        respostas.extend(a_part)
        respostas.append("")
    return strip_trailing_blank(perguntas), strip_trailing_blank(respostas)


def fill_quiz(doc, quiz_body, formula_idx, nnum):
    t = find_table(doc, "quiz")
    if not t:
        avisar(f"Unidade {nnum}: caixa QUIZ não encontrada no template")
        return
    if not quiz_body:
        avisar(f"Unidade {nnum}: sem conteúdo de quiz no markdown")
        return
    perguntas, respostas = split_quiz(quiz_body)
    render_blocks(t.rows[1].cells[0], perguntas, formula_idx, alts=True, first=True)
    if len(t.rows) > 2:
        render_blocks(t.rows[2].cells[0], respostas, formula_idx, first=True)


AAI_NOTA = (
    "Nota de produção: conforme o Plano de Aprendizagem Proposto desta disciplina, "
    "a Atividade Avaliativa Individual (AAI) é única para as quatro unidades e está "
    "concentrada na Unidade 1 (caixa \"AAI – Atividade avaliativa individual\" de "
    "TEMPLATE - Unidade 1_Model Based Design for Cyber-Physical Systems.docx). "
    "Esta unidade não repete, substitui nem complementa aquela atividade — trata-se "
    "de uma decisão deliberada de desenho instrucional, não de uma lacuna de produção."
)


def fill_aai(doc, aai_body, formula_idx, nnum):
    t = find_table(doc, "aai")
    if not t:
        t = find_table(doc, "atividade avaliativa individual")
    if not t:
        if nnum in (3, 4):
            avisar(f"Unidade {nnum}: template não possui caixa AAI (esperado — AAI concentrada na Unidade 1)")
        else:
            avisar(f"Unidade {nnum}: caixa AAI esperada não encontrada no template")
        return
    if aai_body:
        render_blocks(t.rows[1].cells[0], aai_body, formula_idx, first=True)
    else:
        render_blocks(t.rows[1].cells[0], [AAI_NOTA], formula_idx, first=True)


MATERIAL_TARGETS = [
    ("direto da fonte", "direto da fonte"),
    ("para mergulhar", "para mergulhar no assunto"),
    ("podcast", "podcast"),
    ("artigo", "artigo científico"),
]


def fill_material(doc, material_items: dict, formula_idx, nnum):
    for key_md, needle_tbl in MATERIAL_TARGETS:
        body = None
        for k, v in material_items.items():
            if key_md in k:
                body = v
                break
        t = find_table(doc, needle_tbl)
        if not t:
            avisar(f"Unidade {nnum}: caixa de material complementar '{needle_tbl}' não encontrada")
            continue
        if body is None:
            avisar(f"Unidade {nnum}: subseção de material complementar '{key_md}' não encontrada no markdown")
            continue
        cell = t.rows[1].cells[0]
        render_blocks(cell, body, formula_idx, first=True)
        for ri in range(len(t.rows) - 1, 1, -1):
            tr = t.rows[ri]._tr
            tr.getparent().remove(tr)


# ============================================================ orientações / placeholders
def fill_disciplina_conteudista(doc) -> None:
    for p in doc.paragraphs:
        for run in p.runs:
            if "XXXX" in run.text:
                low = p.text.lower()
                if "disciplina" in low:
                    run.text = run.text.replace("XXXX", DISCIPLINA)
                elif "conteudista" in low:
                    run.text = run.text.replace("XXXX", CONTEUDISTA)


def remove_orientations(doc) -> int:
    keep_prefix = ("nome da disciplina", "conteudista", "template de produção")
    removed = 0
    for p in list(doc.paragraphs):
        style = p.style.name if p.style else ""
        txt = p.text.strip()
        if not txt:
            continue
        if style.lower().startswith("heading"):
            continue
        if txt.lower().startswith(keep_prefix):
            continue
        p._p.getparent().remove(p._p)
        removed += 1
    return removed


def unlock_rows_and_breaks(doc) -> None:
    for tbl in doc.element.body.iter(qn("w:tbl")):
        for tr in tbl.findall(qn("w:tr")):
            trPr = tr.find(qn("w:trPr"))
            if trPr is not None:
                for tag in ("w:cantSplit", "w:tblHeader"):
                    el = trPr.find(qn(tag))
                    if el is not None:
                        trPr.remove(el)
    for p in doc.element.body.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            for tag in ("w:pageBreakBefore", "w:keepNext", "w:keepLines"):
                el = pPr.find(qn(tag))
                if el is not None:
                    pPr.remove(el)


# ============================================================ preenchimento de uma unidade
def normalizar_rotulos(doc, nnum: int) -> list[str]:
    """Corrige rótulos de caixa herdados com defeito do modelo recebido.

    O modelo da Unidade 2 é uma cópia do da Unidade 1: traz "Plano de Ensino -
    Unidade 1", numera as caixas como "TEXTO BASE AULA 1" a "AULA 4" e mantém a
    caixa de vídeo introdutório, que só existe na Unidade 1. Os modelos das
    Unidades 3 e 4 provam qual é a convenção pretendida — numeração contínua da
    disciplina ("AULA 9" a "AULA 12", "AULA 13" a "AULA 16").

    Sem esta correção, a Unidade 2 sairia com a Aula 5 sob o rótulo "AULA 1", o
    que induziria a coordenação e a equipe de edição ao erro. A correção é feita
    apenas na cópia gerada; o modelo pristino permanece intacto.

    Para as Unidades 1, 3 e 4 esta função é inofensiva: os rótulos já estão
    corretos e nada é alterado.
    """
    primeira = (nnum - 1) * 4 + 1
    aplicadas: list[str] = []

    # A renumeração é POSICIONAL: a k-ésima caixa de aula do documento recebe
    # o número primeira+k. Substituição textual não serve — "TEXTO BASE AULA 1"
    # é prefixo de "AULA 10", "AULA 11" e "AULA 12", e trocá-la corromperia os
    # rótulos das Unidades 3 e 4, que já estão corretos.
    tbl_tag = qn("w:tbl")
    contadores = {"TEXTO BASE AULA": 0, "ROTEIRO VIDEOAULA": 0}

    for tbl in doc.element.body.iter(tbl_tag):
        for t in tbl.iter():
            if not t.tag.endswith("}t") or not t.text:
                continue
            m = re.match(r"^\s*(TEXTO BASE AULA|ROTEIRO VIDEOAULA)\s*(\d+)", t.text)
            if not m:
                continue
            prefixo, atual = m.group(1), int(m.group(2))
            idx = contadores[prefixo]
            contadores[prefixo] += 1
            if idx > 3:            # caixa inesperada; não mexe
                break
            certo = primeira + idx
            if atual != certo:
                t.text = re.sub(rf"^(\s*{prefixo}\s*)\d+", rf"\g<1>{certo}", t.text, count=1)
                aplicadas.append(f"{prefixo} {atual} -> {prefixo} {certo}")
            break                  # um rótulo por caixa

    # Rótulos textuais que não envolvem numeração de aula.
    textuais: list[tuple[str, str]] = [
        ("Plano de Ensino - Unidade 1", f"Plano de Ensino - Unidade {nnum}"),
    ]
    if nnum != 1:
        # o vídeo introdutório é único e pertence à Unidade 1
        textuais.append((
            "Relação da disciplina com atuação profissional + Roteiro do vídeo introdutório",
            "Relação da unidade com atuação profissional",
        ))

    for t in doc.element.body.iter():
        if not t.tag.endswith("}t") or not t.text:
            continue
        for antigo, novo in textuais:
            if antigo != novo and antigo in t.text:
                t.text = t.text.replace(antigo, novo)
                aplicadas.append(f"{antigo!r} -> {novo!r}")

    if aplicadas:
        print("     rótulos corrigidos (defeito do modelo recebido):")
        for a in dict.fromkeys(aplicadas):
            print("       ·", a)
    return aplicadas


def fill_unit(nnum: int) -> Path | None:
    md_path = DISC / f"unidade_{nnum}/unidade_{nnum}.md"
    roteiros_path = DISC / f"unidade_{nnum}/roteiros_20min.md"
    if not md_path.exists():
        avisar(f"Unidade {nnum}: markdown não encontrado em {md_path}")
        return None
    if not roteiros_path.exists():
        avisar(f"Unidade {nnum}: roteiros_20min.md não encontrado em {roteiros_path}")
        return None

    docx_name = f"TEMPLATE - Unidade {nnum}_{DISCIPLINA_ARQUIVO}.docx"
    pristine = PRISTINE / docx_name
    if not pristine.exists():
        avisar(f"Unidade {nnum}: template pristino não encontrado em {pristine}")
        return None

    conteudo = parse_unidade_md(nnum)
    roteiros, roteiros_text = parse_roteiros_md(nnum)

    textos_formula = [conteudo["full_text"], roteiros_text]
    intro_narracao_lines = None
    if nnum == 1:
        intro_narracao_lines, intro_text = parse_roteiro_intro()
        textos_formula.append(intro_text)

    print("  renderizando fórmulas da unidade...")
    formula_idx = build_formula_index(*textos_formula)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUT_DIR / docx_name

    doc = Document(pristine)  # sempre parte do template limpo
    filled = []

    # 1) placeholders disciplina / conteudista
    fill_disciplina_conteudista(doc)

    # 2) caixa intro (relação com atuação profissional + "O que você verá
    #    nesta unidade" [+ roteiro do vídeo introdutório, só na Unidade 1])
    t_intro = find_table(doc, "relação da disciplina com atuação profissional")
    overview_placed = False
    if t_intro is not None:
        cell = t_intro.rows[1].cells[0]
        render_blocks(cell, conteudo["relacao"], formula_idx, first=True)
        if conteudo["overview"]:
            sub = cell.add_paragraph()
            style_body_paragraph(sub)
            mkrun(sub, "O que você verá nesta unidade").bold = True
            render_blocks(cell, conteudo["overview"], formula_idx, first=False)
            overview_placed = True
        if intro_narracao_lines:
            head = cell.add_paragraph(style="Heading 3")
            add_inline(head, "Roteiro do vídeo introdutório", formula_idx)
            render_blocks(cell, intro_narracao_lines, formula_idx, first=False)
        filled.append("intro")
    elif nnum in (1, 2):
        avisar(f"Unidade {nnum}: caixa de relação com atuação profissional esperada, não encontrada")

    # 3) TEXTO BASE AULA i / ROTEIRO VIDEOAULA i, i = 1..4 por unidade
    for i in range(1, 5):
        abs_num = (nnum - 1) * 4 + i
        titulo_aula, body_aula = conteudo["aulas"].get(abs_num, (None, None))
        if titulo_aula is None:
            avisar(f"Unidade {nnum}: Aula {abs_num} não encontrada no markdown")
        else:
            t = find_box_num(doc, "texto base aula", i) or find_box_num(doc, "texto base aula", abs_num)
            if t:
                header = [f"### {titulo_aula}", ""]
                render_blocks(t.rows[1].cells[0], header + body_aula, formula_idx, first=True)
                filled.append(f"texto{abs_num}")
            else:
                avisar(f"Unidade {nnum}: caixa TEXTO BASE AULA {i}/{abs_num} não encontrada")

        titulo_rot, body_rot = roteiros.get(abs_num, (None, None))
        if titulo_rot is None:
            avisar(f"Unidade {nnum}: roteiro da Videoaula {abs_num} não encontrado em roteiros_20min.md")
            continue
        t = find_box_num(doc, "roteiro videoaula", i) or find_box_num(doc, "roteiro videoaula", abs_num)
        if not t:
            avisar(f"Unidade {nnum}: caixa ROTEIRO VIDEOAULA {i}/{abs_num} não encontrada")
            continue
        set_roteiro_titulo(t, titulo_rot)
        render_blocks(t.rows[1].cells[0], body_rot, formula_idx, first=True)
        filled.append(f"roteiro{abs_num}")

    # 4) QUIZ / AAI / material complementar
    fill_quiz(doc, conteudo["quiz"], formula_idx, nnum)
    filled.append("quiz")
    fill_aai(doc, conteudo["aai"], formula_idx, nnum)
    filled.append("aai")
    fill_material(doc, conteudo["material"], formula_idx, nnum)
    filled.append("material")

    # 5) limpar orientações soltas do corpo do documento
    removed = remove_orientations(doc)

    # 6) destravar linhas/paragrafos para evitar páginas em branco
    unlock_rows_and_breaks(doc)

    # 7) seções sem caixa dedicada no template: "O que você verá nesta
    #    unidade" (só quando não coube na caixa intro — Unidades 3 e 4, que
    #    não têm essa caixa), "Síntese da unidade" e "Referências da
    #    unidade" (ABNT, exigência institucional) — todas ao final do corpo
    #    do documento, depois da última caixa (ARTIGO CIENTÍFICO), com
    #    títulos no mesmo estilo (Heading 1) dos demais títulos do template.
    #    Adicionadas DEPOIS de remove_orientations() para não serem
    #    removidas por ela (que varre parágrafos "Normal" soltos no corpo).
    if not overview_placed and conteudo["overview"]:
        h = doc.add_paragraph(style="Heading 1")
        h.add_run("O QUE VOCÊ VERÁ NESTA UNIDADE")
        render_blocks(doc, conteudo["overview"], formula_idx, first=False)
        filled.append("overview-fim-de-documento")

    h = doc.add_paragraph(style="Heading 1")
    h.add_run("SÍNTESE DA UNIDADE")
    render_blocks(doc, conteudo["sintese"], formula_idx, first=False)
    filled.append("sintese")

    h = doc.add_paragraph(style="Heading 1")
    h.add_run("REFERÊNCIAS DA UNIDADE")
    render_blocks(doc, conteudo["referencias"], formula_idx, first=False)
    filled.append("referencias")

    normalizar_rotulos(doc, nnum)

    doc.save(docx_path)
    print("  OK ->", docx_path.name)
    print("     caixas preenchidas:", ", ".join(filled))
    print("     parágrafos de orientação removidos:", removed)
    return docx_path


# ============================================================ verificação
def verificar(nnum: int, docx_path: Path) -> dict:
    import zipfile

    resultado = {"unidade": nnum, "arquivo": str(docx_path)}
    resultado["tamanho_bytes"] = docx_path.stat().st_size

    with zipfile.ZipFile(docx_path) as z:
        bad = z.testzip()
        resultado["zip_ok"] = bad is None
        resultado["zip_bad_file"] = bad

    doc = Document(docx_path)
    resultado["abre_python_docx"] = True

    pristine_path = PRISTINE / f"TEMPLATE - Unidade {nnum}_{DISCIPLINA_ARQUIVO}.docx"
    doc_pristine = Document(pristine_path)
    coord_ok = True
    for idx in (0, 1):
        a = doc.tables[idx]
        b = doc_pristine.tables[idx]
        text_a = "\n".join(c.text for row in a.rows for c in row.cells)
        text_b = "\n".join(c.text for row in b.rows for c in row.cells)
        if text_a != text_b:
            coord_ok = False
    resultado["tabelas_coordenador_intactas"] = coord_ok

    orient_markers = ["xxxx", "caro coordenador", "o conteudista deve", "o(a) conteudista deve",
                       "apagar as orientações", "elabore um título"]
    problemas = []
    caixas_vazias = []
    for el in doc.element.body.iter(qn("w:tbl")):
        header = _box_header(el)
        if header is None:
            continue
        t = Table(el, doc)
        if header in ("tabela para uso exclusivo do(a) coordenador(a)", "data da validação"):
            continue
        # ignora tabelas nao-caixa (VALIDADO, Plano de Ensino, Texto Dialógico...)
        is_box = any(k in header for k in (
            "relação da disciplina", "texto base aula", "roteiro videoaula", "quiz",
            "aai", "atividade avaliativa", "direto da fonte", "para mergulhar", "podcast",
            "artigo científico"))
        if not is_box:
            continue
        body_text = "\n".join(c.text for row in t.rows[1:] for c in row.cells).strip().lower()
        if not body_text:
            caixas_vazias.append(header)
            continue
        for marker in orient_markers:
            if marker in body_text:
                problemas.append((header, marker))
    resultado["caixas_vazias"] = caixas_vazias
    resultado["orientacoes_residuais"] = problemas

    n_formulas = 0
    for shape in doc.element.body.iter(qn("w:drawing")):
        n_formulas += 1
    for shape in doc.element.body.iter(qn("w:pict")):
        n_formulas += 1
    resultado["imagens_embutidas"] = n_formulas

    # doc.tables (e iter_inner_content) não enxergam as tabelas escondidas em
    # w:sdt — nas Unidades 3 e 4 é assim que quase todas as caixas estão
    # embrulhadas. Iterar w:tbl diretamente no corpo pega todas.
    todas_tabelas = [Table(el, doc) for el in doc.element.body.iter(qn("w:tbl"))]
    resultado["paragrafos"] = len(doc.paragraphs) + sum(
        len(c.paragraphs) for t in todas_tabelas for row in t.rows for c in row.cells)
    resultado["caracteres"] = sum(len(p.text) for p in doc.paragraphs) + sum(
        len(c.text) for t in todas_tabelas for row in t.rows for c in row.cells)

    md_path = DISC / f"unidade_{nnum}/unidade_{nnum}.md"
    roteiros_path = DISC / f"unidade_{nnum}/roteiros_20min.md"
    md_chars = len(md_path.read_text(encoding="utf-8")) + len(roteiros_path.read_text(encoding="utf-8"))
    resultado["caracteres_markdown_fonte"] = md_chars

    # "Plano de Ensino" é caixa de uso do(a) coordenador(a) (aninhada em
    # w:sdt dentro da tabela "Data da Validação"), nunca tocada por este
    # script — a instrução "Caro coordenador, insira o plano de ensino..."
    # deve permanecer intacta ali.
    t_plano = find_table(doc, "plano de ensino")
    resultado["plano_ensino_intacto"] = bool(
        t_plano is not None and "caro coordenador" in t_plano.rows[1].cells[0].text.strip().lower()
    )

    # As três seções sem caixa dedicada no template (visão geral, síntese e
    # referências da unidade) — buscadas no documento inteiro, incluindo
    # tabelas escondidas em w:sdt.
    texto_completo = extrair_texto_completo(doc)
    resultado["secoes_sem_caixa_presentes"] = {
        secao: secao.lower() in texto_completo.lower() for secao in SECOES_SEM_CAIXA
    }

    return resultado


TRECHOS_CARACTERISTICOS = {
    1: ["NexaBot", "-342,857", "21,2164"],
    2: ["NexaBot"],
    3: ["NexaBot"],
    4: ["NexaBot"],
}

SECOES_SEM_CAIXA = ["O que você verá nesta unidade", "Síntese da unidade", "Referências da unidade"]


def extrair_texto_completo(doc) -> str:
    """'Converte' o DOCX inteiro para texto puro: parágrafos de corpo, TODAS as
    tabelas (inclusive as escondidas em w:sdt, que doc.tables não alcança) e
    os textos alternativos (descr) das imagens de fórmula — onde vivem números
    como '-342,857' que, por regra desta disciplina, entram como imagem, não
    como texto corrido."""
    partes = [p.text for p in doc.paragraphs]
    for el in doc.element.body.iter(qn("w:tbl")):
        t = Table(el, doc)
        for row in t.rows:
            for c in row.cells:
                partes.append(c.text)
    for docPr in doc.element.body.iter(qn("wp:docPr")):
        descr = docPr.get("descr")
        if descr:
            partes.append(descr)
    return "\n".join(partes)


def verificar_texto_extraido(nnum: int, docx_path: Path) -> dict:
    doc = Document(docx_path)
    texto = extrair_texto_completo(doc)
    achados = {}
    for trecho in TRECHOS_CARACTERISTICOS.get(nnum, []):
        achados[trecho] = trecho in texto
    return achados


def main():
    nums = [int(a) for a in sys.argv[1:]] or [1, 2, 3, 4]
    print("Disciplina:", DISCIPLINA)
    print("-" * 70)

    gerados = {}
    for nnum in nums:
        print(f"Unidade {nnum}:")
        docx_path = fill_unit(nnum)
        if docx_path:
            gerados[nnum] = docx_path
        print()

    print("=" * 70)
    print("VERIFICAÇÃO")
    print("=" * 70)
    for nnum, docx_path in gerados.items():
        print(f"\nUnidade {nnum} — {docx_path}")
        r = verificar(nnum, docx_path)
        print(f"  tamanho: {r['tamanho_bytes']} bytes")
        print(f"  zip_ok: {r['zip_ok']} (bad_file={r['zip_bad_file']})")
        print(f"  abre com python-docx: {r['abre_python_docx']}")
        print(f"  tabelas [0]/[1] (coordenador) intactas: {r['tabelas_coordenador_intactas']}")
        print(f"  caixas vazias: {r['caixas_vazias'] or 'nenhuma'}")
        print(f"  orientações residuais: {r['orientacoes_residuais'] or 'nenhuma'}")
        print(f"  imagens embutidas (fórmulas): {r['imagens_embutidas']}")
        print(f"  parágrafos: {r['paragrafos']} | caracteres: {r['caracteres']} "
              f"(markdown fonte: {r['caracteres_markdown_fonte']} caracteres)")
        print(f"  Plano de Ensino (caixa do coordenador) intacta: {r['plano_ensino_intacto']}")
        print(f"  seções sem caixa dedicada presentes no documento: {r['secoes_sem_caixa_presentes']}")
        achados = verificar_texto_extraido(nnum, docx_path)
        print(f"  trechos característicos encontrados: {achados}")

    if AVISOS:
        print("\n" + "=" * 70)
        print(f"AVISOS ({len(AVISOS)}):")
        for a in AVISOS:
            print("  -", a)


if __name__ == "__main__":
    main()
