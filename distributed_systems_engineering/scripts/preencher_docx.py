#!/usr/bin/env python3
"""Preenche os quatro modelos institucionais de unidade com o conteúdo autoral.

O modelo recebido em `documentos/Unidade N/` é aberto sem ser alterado; cada
caixa colorida recebe a parte correspondente do Markdown; as orientações do
modelo são removidas; e o resultado é gravado em `entrega_final/docx/`.

Uso:
    PYTHONPATH=/tmp/dse-docx-libs python3 scripts/preencher_docx.py        # 1 a 4
    PYTHONPATH=/tmp/dse-docx-libs python3 scripts/preencher_docx.py 2 3    # só 2 e 3
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))

import docx_comum  # noqa: E402

from docx_comum import (  # noqa: E402
    CONTEUDISTA,
    DISCIPLINA,
    DOCUMENTOS,
    FONTE_UNIDADE,
    RAIZ,
    achar_caixa,
    achar_caixa_numerada,
    arquivo_unidade,
    carregar,
    destravar_linhas_e_quebras,
    formatar_run,
    reescrever_cabecalho,
    remover_caixa,
    renderizar,
    secoes,
    sem_secao,
)

TITULO_UNIDADE = {
    1: "Fundamentos, comunicação, tempo e falhas",
    2: "Dados distribuídos, consistência e coordenação",
    3: "Serviços, eventos e plataformas cloud-native",
    4: "Operação, validação e evolução",
}

CAIXAS_MATERIAL = [
    ("Direto da Fonte", "direto da fonte"),
    ("Para Mergulhar no Assunto", "para mergulhar"),
    ("Podcast", "podcast"),
    ("Artigo científico", "artigo"),
]

# Parágrafos de orientação do modelo que devem sair da versão final.
MANTER_PREFIXOS = (
    "template de produção",
    "nome da disciplina",
    "conteudista",
)


def aulas_da_unidade(unidade: int) -> list[int]:
    return list(range((unidade - 1) * 4 + 1, unidade * 4 + 1))


def caixa_de_aula(documento, prefixo: str, aula: int, unidade: int):
    """Os modelos numeram as caixas por unidade (1-4) ou de forma contínua."""
    local = aula - 4 * (unidade - 1)
    return achar_caixa_numerada(documento, prefixo, local) or achar_caixa_numerada(
        documento, prefixo, aula
    )


def substituir_placeholders(documento, unidade: int) -> None:
    for paragrafo in documento.paragraphs:
        minusculo = paragrafo.text.lower()
        for run in paragrafo.runs:
            if "XXXX" not in run.text:
                continue
            if "disciplina" in minusculo:
                run.text = run.text.replace("XXXX", DISCIPLINA)
            elif "conteudista" in minusculo:
                run.text = run.text.replace("XXXX", CONTEUDISTA)
    # O modelo da Unidade 2 é uma cópia do da Unidade 1 e traz o número errado.
    plano = achar_caixa(documento, "plano de ensino")
    if plano is not None:
        reescrever_cabecalho(
            plano, f"Plano de Ensino - Unidade {unidade}", FONTE_UNIDADE
        )
        renderizar(
            plano.rows[1].cells[0],
            [
                "O plano de ensino oficial da disciplina não integrava o pacote "
                "recebido. Enquanto ele não for disponibilizado pela coordenação, "
                "o conteúdo desta unidade segue a proposta provisória registrada em "
                "`PLANO_APRENDIZAGEM_PROPOSTO.md`, conforme a regra de governança de "
                "`DIRETRIZES_PRODUCAO.md`.",
            ],
            limpar=True,
        )


def remover_titulo(documento, texto: str) -> None:
    """Remove um título de seção do modelo que ficou sem caixa correspondente."""
    for paragrafo in list(documento.paragraphs):
        if paragrafo.text.strip().lower() == texto:
            paragrafo._p.getparent().remove(paragrafo._p)
            return


def tem_figura(paragrafo) -> bool:
    elemento = paragrafo._p
    return (
        elemento.find(".//w:drawing", namespaces=elemento.nsmap) is not None
        or elemento.find(".//w:pict", namespaces=elemento.nsmap) is not None
    )


def remover_orientacoes(documento) -> int:
    """Remove os parágrafos de orientação soltos no corpo do modelo.

    Depois de apagá-los sobram sequências de parágrafos vazios que empurravam as
    caixas para a página seguinte; elas são reduzidas a um único separador.
    """
    removidos = 0
    for paragrafo in list(documento.paragraphs):
        texto = paragrafo.text.strip()
        if not texto:
            continue
        estilo = (paragrafo.style.name or "").lower()
        if estilo.startswith("heading"):
            continue
        if texto.lower().startswith(MANTER_PREFIXOS):
            continue
        paragrafo._p.getparent().remove(paragrafo._p)
        removidos += 1

    anterior_vazio = False
    for paragrafo in list(documento.paragraphs):
        vazio = not paragrafo.text.strip() and not tem_figura(paragrafo)
        if vazio and anterior_vazio:
            paragrafo._p.getparent().remove(paragrafo._p)
            removidos += 1
            continue
        anterior_vazio = vazio
    return removidos


def campos(linhas: list[str]) -> dict[str, str]:
    """Lê pares `**Rótulo:** valor` de um bloco de material complementar."""
    resultado: dict[str, str] = {}
    for linha in linhas:
        achado = re.match(r"^\*\*(.+?):\*\*\s*(.*)$", linha.strip())
        if achado:
            resultado[achado.group(1).strip().lower()] = achado.group(2).strip()
    return resultado


def preencher_material(documento, bloco: list[str]) -> list[str]:
    preenchidas = []
    blocos = dict(secoes("\n".join(bloco), nivel=4))
    for rotulo, trecho in CAIXAS_MATERIAL:
        caixa = achar_caixa(documento, trecho)
        if caixa is None or rotulo not in blocos:
            continue
        dados = campos(blocos[rotulo])
        renderizar(
            caixa.rows[0].cells[0] if False else caixa.rows[1].cells[0],
            [dados.get("texto provocativo", "")],
            limpar=True,
        )
        identificacao = []
        if "referência" in dados and rotulo != "Artigo científico":
            identificacao.append(f"**Referência:** {dados['referência']}")
        if "link de acesso" in dados:
            identificacao.append(f"**Link de acesso:** {dados['link de acesso']}")
        if "trecho obrigatório" in dados:
            identificacao.append(f"**Trecho obrigatório:** {dados['trecho obrigatório']}")
        if "aula indicada" in dados:
            identificacao.append(
                f"**Em qual aula o material deverá entrar:** {dados['aula indicada']}"
            )
        if len(caixa.rows) > 2:
            renderizar(caixa.rows[2].cells[0], identificacao, limpar=True)
        if rotulo == "Artigo científico" and len(caixa.rows) > 3:
            renderizar(
                caixa.rows[3].cells[0],
                [
                    "**Referência bibliográfica do artigo no formato ABNT:**",
                    dados.get("referência", ""),
                ],
                limpar=True,
            )
        preenchidas.append(rotulo)
    return preenchidas


def preencher_quiz(documento, bloco: list[str]) -> bool:
    caixa = achar_caixa(documento, "quiz")
    if caixa is None:
        return False
    perguntas: list[str] = []
    respostas: list[str] = []
    atual = perguntas
    numero = 0
    for linha in bloco:
        texto = linha.strip()
        if re.match(r"^\*\*Questão \d+\.?\*\*", texto):
            numero += 1
            perguntas.append(f"**Questão {numero}.**" + texto.split("**", 2)[-1])
            atual = perguntas
            continue
        if texto.startswith("*Feedback"):
            respostas.append(f"**Questão {numero} — devolutiva**")
            respostas.append(texto)
            atual = respostas
            continue
        atual.append(linha)
    renderizar(caixa.rows[1].cells[0], perguntas, limpar=True, alternativas=True)
    if len(caixa.rows) > 2:
        renderizar(caixa.rows[2].cells[0], respostas, limpar=True)
    return True


def preencher_unidade(unidade: int) -> Path:
    modelo = (
        DOCUMENTOS / f"Unidade {unidade}" / f"TEMPLATE - Unidade {unidade}_nome da disciplina.docx"
    )
    if not modelo.exists():
        raise FileNotFoundError(modelo)

    fonte_unidade = carregar(RAIZ / f"unidade_{unidade}" / f"unidade_{unidade}.md")
    fonte_roteiros = carregar(RAIZ / f"unidade_{unidade}" / "roteiros_20min.md")

    # os caminhos de imagem do Markdown são relativos ao arquivo da unidade
    docx_comum.BASE_IMAGENS = RAIZ / f"unidade_{unidade}"

    blocos_unidade = dict(secoes(fonte_unidade))
    roteiros = {
        int(re.match(r"Roteiro da Videoaula (\d+)", titulo).group(1)): (titulo, corpo)
        for titulo, corpo in secoes(fonte_roteiros)
        if re.match(r"Roteiro da Videoaula \d+", titulo)
    }
    aulas = {
        int(re.match(r"Aula (\d+)", titulo).group(1)): (titulo, corpo)
        for titulo, corpo in blocos_unidade.items()
        if re.match(r"Aula \d+", titulo)
    }
    fechamento = dict(
        secoes("\n".join(blocos_unidade.get("Atividades, síntese e material complementar", [])), nivel=3)
    )

    documento = Document(modelo)
    preenchidas: list[str] = []

    substituir_placeholders(documento, unidade)

    # --- caixa de abertura: só existe (e só faz sentido) na Unidade 1
    abertura = achar_caixa(documento, "relação da disciplina")
    if abertura is not None:
        if unidade == 1:
            corpo = ["### Relação da disciplina com a atuação profissional", ""]
            corpo += blocos_unidade.get("Relação da unidade com a atuação profissional", [])
            corpo += ["", "### Roteiro do vídeo introdutório (até 2 minutos)", ""]
            introducao = carregar(RAIZ / "roteiro_video_introdutorio.md")
            for titulo, linhas in secoes(introducao):
                corpo += [f"#### {titulo}", ""] + linhas
            renderizar(abertura.rows[1].cells[0], corpo, limpar=True)
            preenchidas.append("abertura")
        else:
            remover_caixa(abertura)
            preenchidas.append("abertura removida")

    # --- textos-base e roteiros
    lista_aulas = aulas_da_unidade(unidade)
    for posicao, numero in enumerate(lista_aulas):
        titulo_aula, corpo_aula = aulas[numero]
        corpo = []
        if posicao == 0:
            corpo += ["### O que você verá nesta unidade", ""]
            corpo += blocos_unidade.get("O que você verá nesta unidade", [])
            if unidade > 1:
                corpo += ["", "### Relação da unidade com a atuação profissional", ""]
                corpo += blocos_unidade.get(
                    "Relação da unidade com a atuação profissional", []
                )
            corpo.append("")
        corpo += [f"### {titulo_aula}", ""]
        corpo += sem_secao(corpo_aula, ["Roteiro da Videoaula"])
        if posicao == len(lista_aulas) - 1 and "Síntese da unidade" in fechamento:
            corpo += ["", "### Síntese da unidade", ""] + fechamento["Síntese da unidade"]

        caixa = caixa_de_aula(documento, "texto base aula", numero, unidade)
        if caixa is None:
            raise RuntimeError(f"caixa TEXTO BASE AULA {numero} não encontrada")
        # O modelo da Unidade 2 é cópia do da Unidade 1 e numera as caixas de 1 a 4.
        if caixa.rows[0].cells[0].text.strip().upper() != f"TEXTO BASE AULA {numero}":
            reescrever_cabecalho(caixa, f"TEXTO BASE AULA {numero}", FONTE_UNIDADE)
        renderizar(caixa.rows[1].cells[0], corpo, limpar=True)
        preenchidas.append(f"texto base {numero}")

        titulo_roteiro, corpo_roteiro = roteiros[numero]
        rotulo = titulo_roteiro.split("—", 1)[1].strip().strip("“”")
        caixa = caixa_de_aula(documento, "roteiro videoaula", numero, unidade)
        if caixa is None:
            raise RuntimeError(f"caixa ROTEIRO VIDEOAULA {numero} não encontrada")
        reescrever_cabecalho(
            caixa, f"ROTEIRO VIDEOAULA {numero}: {rotulo}", FONTE_UNIDADE
        )
        renderizar(caixa.rows[1].cells[0], corpo_roteiro, limpar=True)
        preenchidas.append(f"roteiro {numero}")

    # --- quiz, AAI e material complementar
    if "Quiz não avaliativo" in fechamento and preencher_quiz(
        documento, fechamento["Quiz não avaliativo"]
    ):
        preenchidas.append("quiz")

    aai = achar_caixa(documento, "aai") or achar_caixa(documento, "atividade avaliativa")
    if aai is not None:
        bloco_aai = fechamento.get("Atividade Avaliativa Individual (AAI)")
        if bloco_aai:
            renderizar(aai.rows[1].cells[0], bloco_aai, limpar=True)
            preenchidas.append("aai")
        else:
            # O contrato prevê AAI apenas na Unidade 1.
            remover_caixa(aai)
            remover_titulo(documento, "atividade verificadora")
            preenchidas.append("aai removida")

    if "Material complementar" in fechamento:
        preenchidas += preencher_material(documento, fechamento["Material complementar"])

    removidos = remover_orientacoes(documento)
    destravar_linhas_e_quebras(documento)

    documento.core_properties.title = f"Unidade {unidade} — {TITULO_UNIDADE[unidade]}"
    documento.core_properties.subject = DISCIPLINA
    documento.core_properties.author = CONTEUDISTA
    documento.core_properties.comments = ""

    destino = arquivo_unidade(unidade)
    destino.parent.mkdir(parents=True, exist_ok=True)
    documento.save(destino)
    print(f"  OK -> {destino.name}")
    print(f"     caixas preenchidas: {', '.join(preenchidas)}")
    print(f"     parágrafos de orientação removidos: {removidos}")
    return destino


def main() -> int:
    unidades = [int(a) for a in sys.argv[1:]] or [1, 2, 3, 4]
    print("Disciplina:", DISCIPLINA)
    print("-" * 66)
    for unidade in unidades:
        print(f"Unidade {unidade}:")
        preencher_unidade(unidade)
    print("-" * 66)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
