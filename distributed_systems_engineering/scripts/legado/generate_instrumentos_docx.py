#!/usr/bin/env python3
"""Gera os instrumentos e fichas de validação sem alterar os originais."""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ORIGINAIS = ROOT / "documentos"
SAIDA = ROOT / "entrega_final" / "docx"

DISCIPLINA = "Distributed Systems Engineering"
CONTEUDISTA = "Afonso Cesar Lelis Brandão"


def clear_paragraph(paragraph):
    """Remove o conteúdo, mas preserva as propriedades do parágrafo."""
    for child in list(paragraph._p):
        if child.tag == qn("w:pPr"):
            continue
        paragraph._p.remove(child)


def clear_cell(cell):
    paragraphs = cell.paragraphs
    first = paragraphs[0]
    clear_paragraph(first)
    for paragraph in paragraphs[1:]:
        cell._tc.remove(paragraph._p)
    return first


def style_run(run, *, bold=None, italic=None, code=False, size=12, font_name="Times New Roman"):
    selected_font = "Courier New" if code else font_name
    run.font.name = selected_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), selected_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def normalize_inline(text: str) -> str:
    replacements = {
        r"\lambda_{\text{pico}}": "λ_pico",
        r"C_{\text{instância}}": "C_instância",
        r"U_{\text{alvo}}": "U_alvo",
        r"\lambda_{\text{eventos}}": "λ_eventos",
        r"\lambda_{\text{HTTP}}": "λ_HTTP",
        r"p_{\text{pedidos}}": "p_pedidos",
        r"e_{\text{eventos por pedido}}": "e_eventos_por_pedido",
        r"C_{\text{mês}}": "C_mês",
        r"C_{\text{fixo}}": "C_fixo",
        r"C_{\text{instância-hora}}": "C_instância-hora",
        r"C_{\text{dados}}": "C_dados",
        r"C_{\text{mensageria}}": "C_mensageria",
        r"C_{\text{atual}}": "C_atual",
        r"L_{\text{local}}": "L_local",
        r"L_{\text{recebido}}": "L_recebido",
        r"\leftarrow": "←",
        r"\rightarrow": "→",
        r"\Rightarrow": "⇒",
        r"\max": "max",
        r"\times": "×",
        r"\sum_j": "Σj",
        r"\leq": "≤",
        r"\text{ eventos/s}": " eventos/s",
        r"\text{ instâncias}": " instâncias",
        r"\left\lceil": "⌈",
        r"\right\rceil": "⌉",
        r"\frac": "",
        "{": "",
        "}": "",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text.replace("$", "")


INLINE_TOKEN = re.compile(
    r"(\[[^\]]+\]\(https?://[^)]+\)|<https?://[^>]+>|https?://[^\s<>)]+|"
    r"\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+`)"
)


def add_hyperlink(paragraph, label: str, url: str, *, size=12, font_name="Times New Roman"):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = paragraph.add_run(label)
    style_run(run, size=size, font_name=font_name)
    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    run.font.underline = True
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, *, default_bold=False, size=12, font_name="Times New Roman"):
    text = normalize_inline(text)
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            style_run(run, bold=default_bold, size=size, font_name=font_name)
        token = match.group(0)
        if token.startswith("["):
            link = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2), size=size, font_name=font_name)
        elif token.startswith("<http"):
            run = paragraph.add_run("<")
            style_run(run, bold=default_bold, size=size, font_name=font_name)
            add_hyperlink(paragraph, token[1:-1], token[1:-1], size=size, font_name=font_name)
            run = paragraph.add_run(">")
            style_run(run, bold=default_bold, size=size, font_name=font_name)
        elif token.startswith("http"):
            add_hyperlink(paragraph, token, token, size=size, font_name=font_name)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            style_run(run, bold=True, size=size, font_name=font_name)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            style_run(run, bold=default_bold, code=True, size=size, font_name=font_name)
        else:
            run = paragraph.add_run(token[1:-1])
            style_run(run, bold=default_bold, italic=True, size=size, font_name=font_name)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        style_run(run, bold=default_bold, size=size, font_name=font_name)


def set_spacing(paragraph, *, after=6, line=1.15):
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_paragraph(container, text="", style=None):
    try:
        paragraph = container.add_paragraph(style=style)
    except (KeyError, TypeError):
        paragraph = container.add_paragraph()
    if text:
        add_inline(paragraph, text)
    set_spacing(paragraph)
    return paragraph


def add_heading(container, text: str, level: int):
    if hasattr(container, "add_heading"):
        paragraph = container.add_heading(level=min(level, 3))
    else:
        paragraph = container.add_paragraph()
    add_inline(paragraph, text, default_bold=True, size=14 if level == 1 else 12)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def simplify_math(lines: list[str]) -> str:
    raw = " ".join(line.strip() for line in lines)
    if ("lambda_{\\text{pico}}" in raw or "frac{6.000}" in raw) and "frac" in raw:
        if "6.000" in raw:
            return "N = ⌈6.000 / (200 × 0,70)⌉ = ⌈42,86⌉ = 43 instâncias"
        return "N = ⌈λ_pico / (C_instância × U_alvo)⌉"
    if "lambda_{\\text{eventos}}" in raw:
        if "4.800" in raw:
            return "λ_eventos = 6.000 × 0,20 × 4 = 4.800 eventos/s"
        return "λ_eventos = λ_HTTP × p_pedidos × e_eventos_por_pedido"
    if "C_{\\text{mês}}" in raw:
        return (
            "C_mês = C_fixo + Σj (N_j × h_j × C_instância-hora) "
            "+ C_dados + C_mensageria"
        )
    return normalize_inline(raw)


def render_markdown(container, markdown: str):
    lines = markdown.strip().splitlines()
    index = 0
    math_lines: list[str] | None = None
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        index += 1
        if stripped == "$$":
            if math_lines is None:
                math_lines = []
            else:
                paragraph = add_paragraph(container)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                add_inline(paragraph, simplify_math(math_lines))
                math_lines = None
            continue
        if math_lines is not None:
            math_lines.append(stripped)
            continue
        if not stripped or stripped == "---":
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            add_heading(container, heading.group(2), len(heading.group(1)))
            continue
        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            text = quote.group(1)
            paragraph = add_paragraph(container, style="Quote")
            add_inline(paragraph, text)
            continue
        bullet = re.match(r"^(\s*)-\s+(.+)$", line)
        if bullet:
            level = 2 if len(bullet.group(1)) >= 2 else 1
            style = "List Bullet 2" if level == 2 else "List Bullet"
            paragraph = add_paragraph(container, style=style)
            add_inline(paragraph, bullet.group(2))
            continue
        numbered = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if numbered:
            level = 2 if len(numbered.group(1)) >= 2 else 1
            style = "List Number 2" if level == 2 else "List Number"
            paragraph = add_paragraph(container, style=style)
            add_inline(paragraph, numbered.group(2))
            continue
        paragraph = add_paragraph(container)
        add_inline(paragraph, stripped)


def extract(markdown: str, start: str, end: str | None) -> str:
    body = markdown.split(start, 1)[1]
    if end is not None:
        body = body.split(end, 1)[0]
    return body.strip()


def set_text(paragraph, text: str, *, bold=False, size=12):
    clear_paragraph(paragraph)
    add_inline(paragraph, text, default_bold=bold, size=size)
    set_spacing(paragraph)


def fill_pending_validation_tables(document, *, after_recording=False):
    pending = "Pendente após gravação" if after_recording else "Pendente de validação pela coordenação"
    if document.tables:
        table = document.tables[0]
        for row in table.rows[2:]:
            if len(row.cells) >= 2:
                set_text(clear_cell(row.cells[1]), pending)
    if len(document.tables) > 1 and len(document.tables[1].rows) > 1:
        for cell in document.tables[1].rows[1].cells:
            set_text(clear_cell(cell), pending)


def move_table_to_end(document, table):
    table_xml = table._tbl
    parent = table_xml.getparent()
    parent.remove(table_xml)
    body = document._body._element
    section_properties = body.sectPr
    body.insert(body.index(section_properties), table_xml)


def set_core_properties(document, title: str, subject: str):
    props = document.core_properties
    props.title = title
    props.subject = subject
    props.author = CONTEUDISTA
    props.keywords = "Distributed Systems Engineering; NEaD; UniFECAF"
    props.comments = "Gerado a partir de cópia imutável do modelo institucional."


def replace_paragraph_properties(paragraph, properties):
    current = paragraph._p.pPr
    if current is not None:
        paragraph._p.remove(current)
    if properties is not None:
        paragraph._p.insert(0, deepcopy(properties))


def table_paragraphs(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def apply_evaluation_typography(document, formats):
    """Reaplica Arial 12 e os três padrões de parágrafo do modelo recebido."""

    part_b = False
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("PARTE B —"):
            part_b = True

        if index == 4:
            selected = formats["title"]
        elif index in (5, 6):
            selected = formats["metadata"]
        elif index == 8:
            selected = formats["overview"]
        elif paragraph.style.name.startswith("Heading 1") or paragraph.style.name.startswith("Heading 2"):
            selected = formats["section_heading"]
        elif part_b and not paragraph.style.name.startswith("Heading") and not text.startswith("NÃO DISTRIBUIR"):
            selected = formats["answer"]
        else:
            selected = formats["question"]
        replace_paragraph_properties(paragraph, selected)

        for run in paragraph.runs:
            style_run(run, size=12, font_name="Arial")

    for paragraph in table_paragraphs(document):
        for run in paragraph.runs:
            style_run(run, size=12, font_name="Arial")


def generate_work():
    source = ORIGINAIS / "Instrumentos Avaliativos" / "TEMPLATE ENTREGA DE TRABALHO - nomedadisciplina.docx"
    target = SAIDA / "Entrega de Trabalho PBL - Distributed Systems Engineering - versão institucional.docx"
    markdown = (ROOT / "instrumentos_avaliativos" / "entrega_trabalho.md").read_text(encoding="utf-8")
    document = Document(source)
    fill_pending_validation_tables(document)

    title = extract(markdown, "## 1. Título", "---").replace("**", "").strip()
    challenge = extract(markdown, "## 2. Desafio", "---")
    sources = extract(markdown, "## 3. Fontes de pesquisa", "---")
    deliverable = extract(markdown, "## 4. Componentes avaliativos, submissão e pontuação", "---")
    roadmap = extract(markdown, "## Roteiro do estudante", "---\n\n# Parte B")
    solution = extract(markdown, "## Solução esperada e critérios de correção", None)

    # Metadados e avisos de produção substituem as instruções internas soltas.
    set_text(document.paragraphs[0], "Roteiro para elaboração com Problem-Based Learning", bold=True)
    set_text(document.paragraphs[6], f"Disciplina: {DISCIPLINA}")
    set_text(document.paragraphs[7], f"Professor-conteudista: {CONTEUDISTA}")
    set_text(
        document.paragraphs[8],
        "PARTE A — VERSÃO DO ESTUDANTE. A solução exclusiva do tutor está delimitada ao final.",
        bold=True,
    )
    set_text(
        document.tables[2].cell(0, 0).paragraphs[0],
        "Este caso aplica os conteúdos da disciplina a uma situação realista do mercado de trabalho.",
    )

    boxes = [
        (3, "PARTE A — TÍTULO", title),
        (4, "PARTE A — DESAFIO", challenge),
        (5, "PARTE A — FONTES DE PESQUISA", sources),
        (6, "PARTE A — COMPONENTES, SUBMISSÃO E PONTUAÇÃO", deliverable),
        (7, "PARTE B — SOLUÇÃO EXCLUSIVA DO PROFESSOR TUTOR", solution),
    ]
    for table_index, heading, content in boxes:
        table = document.tables[table_index]
        set_text(clear_cell(table.cell(0, 0)), heading, bold=True)
        clear_cell(table.cell(1, 0))
        render_markdown(table.cell(1, 0), content)

    # Remove o roteiro genérico do modelo e inclui a versão completa antes da solução.
    for paragraph in document.paragraphs[27:]:
        clear_paragraph(paragraph)
    add_heading(document, "Roteiro do estudante", 1)
    render_markdown(document, roadmap)
    document.add_page_break()
    add_heading(document, "PARTE B — VERSÃO EXCLUSIVA DO PROFESSOR TUTOR", 1)
    warning = add_paragraph(document)
    add_inline(
        warning,
        "NÃO DISTRIBUIR AOS ESTUDANTES. O quadro seguinte contém a solução esperada e os critérios de correção.",
        default_bold=True,
    )
    move_table_to_end(document, document.tables[7])

    set_core_properties(document, f"Entrega de Trabalho PBL — {DISCIPLINA}", "Instrumento avaliativo institucional")
    document.save(target)
    return target


def generate_evaluation():
    source = ORIGINAIS / "Instrumentos Avaliativos" / "Avaliação final_(10 discursivas)_nomedadisciplina.docx"
    target = SAIDA / "Avaliação final - 10 discursivas - Distributed Systems Engineering - versão institucional.docx"
    markdown = (ROOT / "instrumentos_avaliativos" / "avaliacao_dissertativa.md").read_text(encoding="utf-8")
    document = Document(source)
    fill_pending_validation_tables(document)

    original_paragraphs = list(document.paragraphs)
    formats = {
        "title": deepcopy(original_paragraphs[4]._p.pPr),
        "metadata": deepcopy(original_paragraphs[5]._p.pPr),
        "overview": deepcopy(original_paragraphs[8]._p.pPr),
        "section_heading": deepcopy(original_paragraphs[22]._p.pPr),
        "question": deepcopy(original_paragraphs[24]._p.pPr),
        "answer": deepcopy(original_paragraphs[26]._p.pPr),
    }

    # Remove integralmente exemplos e instruções do corpo, preservando tabelas, estilos, cabeçalho e rodapé.
    for paragraph in original_paragraphs:
        clear_paragraph(paragraph)
    for paragraph in original_paragraphs[9:]:
        paragraph._p.getparent().remove(paragraph._p)
    set_text(document.paragraphs[4], "AVALIAÇÃO FINAL", bold=True, size=12)
    set_text(document.paragraphs[5], f"Disciplina: {DISCIPLINA}")
    set_text(document.paragraphs[6], f"Professor-conteudista: {CONTEUDISTA}")
    set_text(document.paragraphs[8], "10 questões dissertativas — versão institucional com respostas ao final.", bold=True)

    part_a = extract(markdown, "# Parte A — Versão do estudante", "---\n\n# Parte B")
    # O checklist de exportação permanece apenas na fonte técnica em Markdown.
    part_b = extract(
        markdown,
        "## Respostas esperadas e critérios de correção",
        "## Conferência antes da exportação",
    )
    add_heading(document, "PARTE A — VERSÃO DO ESTUDANTE", 1)
    render_markdown(document, part_a)
    document.add_page_break()
    add_heading(document, "PARTE B — VERSÃO EXCLUSIVA DO PROFESSOR TUTOR", 1)
    warning = add_paragraph(document)
    add_inline(
        warning,
        "NÃO DISTRIBUIR AOS ESTUDANTES. As respostas esperadas e as devolutivas estão reunidas ao final, conforme o modelo institucional.",
        default_bold=True,
    )
    add_heading(document, "Respostas esperadas e critérios de correção", 2)
    render_markdown(document, part_b)
    apply_evaluation_typography(document, formats)

    set_core_properties(document, f"Avaliação final dissertativa — {DISCIPLINA}", "Instrumento avaliativo institucional")
    document.save(target)
    return target


def collect_video_metadata():
    metadata: dict[int, tuple[str, str]] = {}
    heading_pattern = re.compile(
        r'^## (?:Roteiro da )?Videoaula (\d+) — [“"]?(.+?)[”"]?$',
        re.MULTILINE,
    )
    objective_pattern = re.compile(r"\*\*Objetivo(?: da videoaula)?:\*\*\s*(.+)")
    for unit in range(1, 5):
        text = (ROOT / f"unidade_{unit}" / "roteiros_20min.md").read_text(encoding="utf-8")
        headings = list(heading_pattern.finditer(text))
        for index, match in enumerate(headings):
            end = headings[index + 1].start() if index + 1 < len(headings) else len(text)
            section = text[match.end() : end]
            objective = objective_pattern.search(section)
            metadata[int(match.group(1))] = (
                match.group(2).strip(),
                objective.group(1).strip() if objective else "Objetivo a confirmar na revisão do roteiro.",
            )

    intro_text = (ROOT / "roteiro_video_introdutorio.md").read_text(encoding="utf-8")
    intro_objective = extract(intro_text, "## Objetivo", "## Roteiro falado").replace("\n", " ").strip()
    metadata[0] = ("Vídeo introdutório: mercado e jornada de aprendizagem", intro_objective)
    return metadata


VALIDATION_TEMPLATES = [
    (
        ORIGINAIS / "Unidade 1" / "Videoaula_ Introdutória + 1 a 4" / "Videoaulas Introdutória + 1 a 4 - Validação.docx",
        "Validação - Vídeo introdutório e videoaulas 1 a 4 - Distributed Systems Engineering.docx",
        [0, 1, 2, 3, 4],
    ),
    (
        ORIGINAIS / "Unidade 2" / "Videoaulas 5 a 8" / "Videoaulas 5 a 8 - Validação.docx",
        "Validação - Videoaulas 5 a 8 - Distributed Systems Engineering.docx",
        [5, 6, 7, 8],
    ),
    (
        ORIGINAIS / "Unidade 3" / "Videoaula 9 a 12" / "Videoaulas 9 a 12 - Validação.docx",
        "Validação - Videoaulas 9 a 12 - Distributed Systems Engineering.docx",
        [9, 10, 11, 12],
    ),
    (
        ORIGINAIS / "Unidade 4" / "Videoaula 13 a 16" / "Videoaulas 13 a 16 - Validação.docx",
        "Validação - Videoaulas 13 a 16 - Distributed Systems Engineering.docx",
        [13, 14, 15, 16],
    ),
]


def generate_validation_forms():
    metadata = collect_video_metadata()
    targets = []
    for source, filename, videos in VALIDATION_TEMPLATES:
        target = SAIDA / filename
        document = Document(source)

        general = (
            f"Disciplina: {DISCIPLINA}. Status geral: PENDENTE APÓS GRAVAÇÃO. "
            "Documento de uso exclusivo da coordenação e do profissional de vídeo; nenhum parecer foi simulado."
        )
        for paragraph in document.paragraphs:
            clear_paragraph(paragraph)
        set_text(document.paragraphs[0], general, bold=True)

        coordinator_table = document.tables[0]
        videomaker_table = document.tables[3]
        for offset, video_number in enumerate(videos, start=2):
            title, objective = metadata[video_number]
            label = "Vídeo introdutório" if video_number == 0 else f"Videoaula {video_number}"
            title_sentence = title if title.endswith((".", "!", "?")) else title + "."
            neutral = (
                f"Disciplina: {DISCIPLINA}. {label}. Título: {title_sentence} Objetivo: {objective} "
                "Duração prevista: "
                + ("até 2 minutos" if video_number == 0 else "20 minutos")
                + ". PENDENTE APÓS GRAVAÇÃO."
            )
            set_text(clear_cell(coordinator_table.cell(offset, 1)), neutral)
            set_text(
                clear_cell(videomaker_table.cell(offset, 1)),
                neutral + " Imagem, áudio, edição e oratória ainda não avaliados.",
            )

        # Datas e responsáveis somente podem ser informados depois da gravação.
        for table_index in (1, 4):
            table = document.tables[table_index]
            for cell in table.rows[1].cells:
                set_text(clear_cell(cell), "Pendente após gravação")

        # Tabelas de status (VALIDADO/NÃO VALIDADO/CANCELADO/AJUSTAR) permanecem sem marcação.
        set_core_properties(document, filename.removesuffix(".docx"), "Ficha institucional de validação de videoaulas")
        document.save(target)
        targets.append(target)
    return targets


def validate_evaluation_typography(document, path: Path):
    paragraphs = [p for p in document.paragraphs if p.text.strip()]
    paragraphs.extend(p for p in table_paragraphs(document) if p.text.strip())
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run.font.name != "Arial" or run.font.size is None or abs(run.font.size.pt - 12) > 0.01:
                raise ValueError(f"Tipografia divergente do modelo em {path.name}: {paragraph.text[:30]}")

    prompt_count = 0
    rubric_count = 0
    body = document.paragraphs
    for index, paragraph in enumerate(body):
        text = paragraph.text.strip()
        if re.match(r"Questão \d+ —", text):
            prompt_count += 1
            prompt = next((candidate for candidate in body[index + 1 :] if candidate.text.strip()), None)
            if prompt is None or prompt.paragraph_format.line_spacing != 1.5:
                raise ValueError(f"Espaçamento de enunciado divergente em {path.name}: {text}")
        if re.match(r"0 a \d+ pontos:", text):
            rubric_count += 1
            if (
                paragraph.paragraph_format.line_spacing != 1.0
                or paragraph.alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            ):
                raise ValueError(f"Formatação de resposta divergente em {path.name}: {text[:30]}")
    if prompt_count != 10 or rubric_count != 40:
        raise ValueError(
            f"Estrutura tipográfica incompleta em {path.name}: "
            f"{prompt_count} enunciados e {rubric_count} critérios"
        )
    if document.paragraphs[4].alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
        raise ValueError(f"Título não centralizado em {path.name}")
    return {"font": "Arial", "size_pt": 12, "prompts_1_5": prompt_count, "rubrics_1_0": rubric_count}


def validate_work_hyperlinks(document, path: Path):
    hyperlink_elements = len(document.element.xpath(".//w:hyperlink"))
    targets = {
        relation.target_ref
        for relation in document.part.rels.values()
        if relation.reltype == RT.HYPERLINK and relation.target_ref.startswith(("http://", "https://"))
    }
    if hyperlink_elements < 5 or len(targets) < 5:
        raise ValueError(
            f"Referências sem hyperlinks suficientes em {path.name}: "
            f"{hyperlink_elements} elementos e {len(targets)} relações HTTP"
        )
    return {"hyperlink_elements": hyperlink_elements, "http_relationships": len(targets)}


def validate_docx(path: Path, *, forbidden=()):
    document = Document(path)
    if not document.sections:
        raise ValueError(f"Documento sem seção: {path}")
    with ZipFile(path) as archive:
        corrupt_member = archive.testzip()
        package_names = archive.namelist()
    if corrupt_member:
        raise ValueError(f"Membro ZIP inválido em {path.name}: {corrupt_member}")
    forbidden_parts = [
        name
        for name in package_names
        if name.startswith(("customXml/", "word/comments", "word/people"))
        or name == "docProps/custom.xml"
    ]
    if forbidden_parts:
        raise ValueError(f"Partes internas residuais em {path.name}: {len(forbidden_parts)}")
    text = "\n".join(p.text for p in document.paragraphs)
    text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    for phrase in forbidden:
        if phrase in text:
            raise ValueError(f"Texto residual proibido em {path.name}: {phrase}")
    latex_residuals = re.findall(r"\\[A-Za-z]+|\\\s|\$|[_^]\{", text)
    if latex_residuals:
        raise ValueError(f"Sintaxe LaTeX residual em {path.name}: {len(latex_residuals)} ocorrência(s)")
    markdown_residuals = re.findall(
        r"\*\*|(?<!\*)\*[^*\n]{1,160}\*(?!\*)|`|\[[^\]]+\]\(https?://",
        text,
    )
    if markdown_residuals:
        raise ValueError(
            f"Marcação Markdown residual em {path.name}: {len(markdown_residuals)} ocorrência(s)"
        )
    formula_paragraphs = list(document.paragraphs) + [
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    ]
    for paragraph in formula_paragraphs:
        formula = paragraph.text.strip()
        if formula.startswith(("N =", "λ_eventos =", "C_mês =")) and formula.endswith("."):
            raise ValueError(f"Pontuação isolada após fórmula em {path.name}: {formula[:20]}")
    metrics = {
        "path": str(path.relative_to(ROOT)),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "bytes": path.stat().st_size,
    }
    if path.name.startswith("Avaliação final"):
        metrics["typography"] = validate_evaluation_typography(document, path)
    if path.name.startswith("Entrega de Trabalho PBL"):
        metrics["hyperlinks"] = validate_work_hyperlinks(document, path)
    return metrics


def main():
    SAIDA.mkdir(parents=True, exist_ok=True)
    outputs = [generate_work(), generate_evaluation(), *generate_validation_forms()]
    forbidden = (
        "EXCLUIR OS EXEMPLOS",
        "Olá, Conteudista",
        "EXEMPLOS:",
        "Feedbacks - Exemplos",
        "Conferência antes da exportação",
    )
    results = []
    for path in outputs:
        phrases = forbidden if "Validação" not in path.name else ()
        results.append(validate_docx(path, forbidden=phrases))
    for result in results:
        print(
            f"OK {result['path']} | {result['bytes']} bytes | "
            f"{result['paragraphs']} parágrafos | {result['tables']} tabelas | {result['sections']} seção"
        )


if __name__ == "__main__":
    main()
