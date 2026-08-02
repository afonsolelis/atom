#!/usr/bin/env python3
"""Renderizador Markdown → DOCX usado ao preencher os modelos institucionais.

Mesma abordagem adotada em `data_engineering_and_pipelines/tools`: o modelo
oficial recebido em `documentos/` nunca é alterado; ele é aberto, cada caixa
colorida recebe o conteúdo autoral correspondente, as orientações internas são
removidas e o resultado é gravado como cópia em `entrega_final/docx/`.

Diferença em relação àquela disciplina: aqui não há navegador disponível para
rasterizar fórmulas com MathJax, então o LaTeX das fontes é convertido para
notação Unicode legível no Word (`limpar_formula`).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell

RAIZ = Path(__file__).resolve().parents[1]
DOCUMENTOS = RAIZ / "documentos"
ENTREGA = RAIZ / "entrega_final"

DISCIPLINA = "Distributed Systems Engineering"
CONTEUDISTA = "Afonso Cesar Lelis Brandão"

# O pacote final espelha a árvore de pastas dos originais recebidos em
# `documentos/`, inclusive os nomes irregulares das subpastas de validação.
SUBPASTA_VALIDACAO = {
    1: "Videoaula_ Introdutória + 1 a 4",
    2: "Videoaulas 5 a 8",
    3: "Videoaula 9 a 12",
    4: "Videoaula 13 a 16",
}
FICHA_VALIDACAO = {
    1: "Videoaulas Introdutória + 1 a 4 - Validação",
    2: "Videoaulas 5 a 8 - Validação",
    3: "Videoaulas 9 a 12 - Validação",
    4: "Videoaulas 13 a 16 - Validação",
}
SUBPASTA_SLIDES = {
    1: "SLIDES - Videoaulas Introdutória + 1 a 4",
    2: "SLIDES - Videoaulas 5 a 8",
    3: "SLIDES - Videoaulas 9 a 12",
    4: "SLIDES - Videoaulas 13 a 16",
}


def pasta_unidade(unidade: int) -> Path:
    return ENTREGA / f"Unidade {unidade}"


def pasta_instrumentos() -> Path:
    return ENTREGA / "Instrumentos Avaliativos"


def pasta_validacao(unidade: int) -> Path:
    return pasta_unidade(unidade) / SUBPASTA_VALIDACAO[unidade]


def pasta_slides(unidade: int) -> Path:
    return pasta_unidade(unidade) / SUBPASTA_SLIDES[unidade]


def arquivo_unidade(unidade: int) -> Path:
    return pasta_unidade(unidade) / f"TEMPLATE - Unidade {unidade}_{DISCIPLINA}.docx"


def arquivo_questoes(unidade: int) -> Path:
    return pasta_unidade(unidade) / f"40 Questões - UNI{unidade}_{DISCIPLINA}.docx"


def arquivo_avaliacao() -> Path:
    return pasta_instrumentos() / f"Avaliação final_(10 discursivas)_{DISCIPLINA}.docx"


def arquivo_entrega_trabalho() -> Path:
    return pasta_instrumentos() / f"TEMPLATE ENTREGA DE TRABALHO - {DISCIPLINA}.docx"


def arquivo_ficha(unidade: int) -> Path:
    return pasta_validacao(unidade) / f"{FICHA_VALIDACAO[unidade]} - {DISCIPLINA}.docx"

# Tipografia exigida pelos modelos (ver DIRETRIZES_PRODUCAO.md).
# Largura útil dentro das caixas coloridas do modelo (A4 menos as margens do
# modelo e as margens internas da célula).
LARGURA_CAIXA_IN = 6.0
LARGURA_IMAGEM_IN = 5.7

# Diretório de referência para os caminhos relativos das imagens do Markdown.
# `preencher_docx.py` ajusta antes de renderizar cada arquivo-fonte.
BASE_IMAGENS = RAIZ

FONTE_UNIDADE = "Times New Roman"
FONTE_AVALIACAO = "Arial"
CORPO_PT = Pt(12)


# ============================================================ fórmulas
def limpar_formula(expressao: str) -> str:
    """Converte o subconjunto de LaTeX usado nas fontes em texto legível."""
    resultado = " ".join(expressao.split())
    for _ in range(5):
        novo = re.sub(r"\\(?:text|mathrm|mathit)\{([^{}]*)\}", r"\1", resultado)
        if novo == resultado:
            break
        resultado = novo
    resultado = re.sub(r"_\{([^{}]*)\}", r"_\1", resultado)
    resultado = re.sub(r"\^\{([^{}]*)\}", r"^\1", resultado)
    resultado = resultado.replace("{,}", ",")
    for _ in range(5):
        novo = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1) / (\2)", resultado)
        if novo == resultado:
            break
        resultado = novo
    simbolos = {
        r"\leftarrow": "←",
        r"\Rightarrow": "⇒",
        r"\rightarrow": "→",
        r"\times": "×",
        r"\cdot": "·",
        r"\approx": "≈",
        r"\neq": "≠",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\lceil": "⌈",
        r"\rceil": "⌉",
        r"\lfloor": "⌊",
        r"\rfloor": "⌋",
        r"\sum": "∑",
        r"\Delta": "Δ",
        r"\delta": "δ",
        r"\lambda": "λ",
        r"\rho": "ρ",
        r"\max": "max",
        r"\min": "min",
        r"\quad": " ",
        r"\left": "",
        r"\right": "",
    }
    for antigo, novo in simbolos.items():
        resultado = resultado.replace(antigo, novo)
    resultado = resultado.replace(r"\%", "%").replace(r"\,", " ").replace(r"\ ", " ")
    resultado = resultado.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", resultado).strip()


# ============================================================ formatação básica
def formatar_run(run, fonte: str, tamanho=CORPO_PT, monoespacada: bool = False) -> None:
    run.font.name = "Consolas" if monoespacada else fonte
    run.font.size = Pt(10.5) if monoespacada else tamanho
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for atributo in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(atributo), run.font.name)


def formatar_paragrafo(paragrafo) -> None:
    formato = paragrafo.paragraph_format
    formato.line_spacing = 1.15
    formato.space_after = Pt(0)
    formato.space_before = Pt(0)
    if paragrafo.alignment is None:
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT


def sombrear(paragrafo, cor: str) -> None:
    ppr = paragrafo._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), cor)
    ppr.append(shd)


def adicionar_hyperlink(paragrafo, url: str, texto: str, fonte: str) -> None:
    r_id = paragrafo.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    cor = OxmlElement("w:color")
    cor.set(qn("w:val"), "0563C1")
    rpr.append(cor)
    sublinhado = OxmlElement("w:u")
    sublinhado.set(qn("w:val"), "single")
    rpr.append(sublinhado)
    rfonts = OxmlElement("w:rFonts")
    for atributo in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(atributo), fonte)
    rpr.append(rfonts)
    tamanho = OxmlElement("w:sz")
    tamanho.set(qn("w:val"), "24")
    rpr.append(tamanho)
    run.append(rpr)
    texto_el = OxmlElement("w:t")
    texto_el.set(qn("xml:space"), "preserve")
    texto_el.text = texto
    run.append(texto_el)
    link.append(run)
    paragrafo._p.append(link)


# ============================================================ inline markdown
EM_LINHA = re.compile(
    r"\*\*\*(?P<bi>.+?)\*\*\*"
    r"|\*\*(?P<b>.+?)\*\*"
    r"|`(?P<c>[^`]+)`"
    r"|\[(?P<lt>[^\]]+)\]\((?P<lu>[^)]+)\)"
    r"|<(?P<au>https?://[^>]+)>"
    r"|\$(?P<m>[^$]+)\$"
    r"|(?<![\*\w])\*(?P<i>[^*]+?)\*(?![\*\w])"
)


def adicionar_inline(paragrafo, texto: str, fonte: str) -> None:
    posicao = 0
    for achado in EM_LINHA.finditer(texto):
        if achado.start() > posicao:
            formatar_run(paragrafo.add_run(texto[posicao : achado.start()]), fonte)
        if achado.group("bi") is not None:
            run = paragrafo.add_run(achado.group("bi"))
            run.bold = run.italic = True
            formatar_run(run, fonte)
        elif achado.group("b") is not None:
            run = paragrafo.add_run(achado.group("b"))
            run.bold = True
            formatar_run(run, fonte)
        elif achado.group("c") is not None:
            formatar_run(paragrafo.add_run(achado.group("c")), fonte, monoespacada=True)
        elif achado.group("lt") is not None:
            adicionar_hyperlink(paragrafo, achado.group("lu"), achado.group("lt"), fonte)
        elif achado.group("au") is not None:
            adicionar_hyperlink(paragrafo, achado.group("au"), achado.group("au"), fonte)
        elif achado.group("m") is not None:
            formatar_run(paragrafo.add_run(limpar_formula(achado.group("m"))), fonte)
        elif achado.group("i") is not None:
            run = paragrafo.add_run(achado.group("i"))
            run.italic = True
            formatar_run(run, fonte)
        posicao = achado.end()
    if posicao < len(texto):
        formatar_run(paragrafo.add_run(texto[posicao:]), fonte)


# ============================================================ células e caixas
def limpar_celula(celula: _Cell) -> None:
    """Remove todo o conteúdo da célula, inclusive tabelas de orientação em w:sdt."""
    tc = celula._tc
    tr = tc.getparent()
    if tr is not None and tr.tag == qn("w:tr"):
        trpr = tr.find(qn("w:trPr"))
        if trpr is not None:
            for etiqueta in ("w:cantSplit", "w:tblHeader"):
                elemento = trpr.find(qn(etiqueta))
                if elemento is not None:
                    trpr.remove(elemento)
    for filho in list(tc):
        if filho.tag != qn("w:tcPr"):
            tc.remove(filho)


def garantir_paragrafo_final(celula: _Cell) -> None:
    if len(celula._tc) == 0 or celula._tc[-1].tag != qn("w:p"):
        celula.add_paragraph("")


def cabecalho_caixa(elemento) -> str | None:
    tr = elemento.find(qn("w:tr"))
    if tr is None:
        return None
    tc = tr.find(qn("w:tc"))
    if tc is None:
        return None
    return "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip().lower()


def achar_caixa(documento, trecho: str) -> Table | None:
    trecho = trecho.lower()
    for elemento in documento.element.body.iter(qn("w:tbl")):
        cabecalho = cabecalho_caixa(elemento)
        if cabecalho and trecho in cabecalho:
            return Table(elemento, documento)
    return None


def achar_caixa_numerada(documento, prefixo: str, numero: int) -> Table | None:
    """Casa o número exato: 'aula 1' não pode casar com 'aula 10'."""
    padrao = re.compile(rf"{re.escape(prefixo.lower())}\s*0*{numero}(?!\d)")
    for elemento in documento.element.body.iter(qn("w:tbl")):
        cabecalho = cabecalho_caixa(elemento)
        if cabecalho and padrao.search(cabecalho):
            return Table(elemento, documento)
    return None


def remover_caixa(tabela: Table) -> None:
    tabela._tbl.getparent().remove(tabela._tbl)


def reescrever_cabecalho(tabela: Table, texto: str, fonte: str) -> None:
    """Troca o texto do cabeçalho da caixa preservando a formatação do 1º run."""
    celula = tabela.rows[0].cells[0]
    paragrafos = celula.paragraphs
    if not paragrafos:
        return
    primeiro = paragrafos[0]
    for extra in paragrafos[1:]:
        extra._p.getparent().remove(extra._p)
    runs = primeiro.runs
    if runs:
        runs[0].text = texto
        for run in runs[1:]:
            run.text = ""
    else:
        formatar_run(primeiro.add_run(texto), fonte)


def aparar_linhas(tabela: Table, quantidade: int) -> None:
    """Mantém apenas as `quantidade` primeiras linhas da caixa."""
    for indice in range(len(tabela.rows) - 1, quantidade - 1, -1):
        linha = tabela.rows[indice]._tr
        linha.getparent().remove(linha)


def destravar_linhas_e_quebras(documento) -> None:
    """Ajusta a paginação das caixas preenchidas.

    Três correções: (1) as linhas deixam de ser indivisíveis, para que uma caixa
    longa flua entre páginas em vez de saltar inteira para a seguinte;
    (2) quebras e "manter junto" herdados do modelo são removidos; e (3) o
    cabeçalho de cada caixa recebe *keep with next*, para nunca ficar órfão no
    fim de uma página com o conteúdo na página seguinte.
    """
    for tabela in documento.element.body.iter(qn("w:tbl")):
        for tr in tabela.findall(qn("w:tr")):
            trpr = tr.find(qn("w:trPr"))
            if trpr is None:
                continue
            for etiqueta in ("w:cantSplit", "w:tblHeader"):
                elemento = trpr.find(qn(etiqueta))
                if elemento is not None:
                    trpr.remove(elemento)

    for paragrafo in documento.element.body.iter(qn("w:p")):
        ppr = paragrafo.find(qn("w:pPr"))
        if ppr is None:
            continue
        for etiqueta in ("w:pageBreakBefore", "w:keepNext", "w:keepLines"):
            elemento = ppr.find(qn(etiqueta))
            if elemento is not None:
                ppr.remove(elemento)

    for tabela in documento.element.body.iter(qn("w:tbl")):
        linhas = tabela.findall(qn("w:tr"))
        if len(linhas) < 2:
            continue
        for paragrafo in linhas[0].iter(qn("w:p")):
            ppr = paragrafo.find(qn("w:pPr"))
            if ppr is None:
                ppr = OxmlElement("w:pPr")
                paragrafo.insert(0, ppr)
            ppr.insert(0, OxmlElement("w:keepNext"))


# ============================================================ blocos markdown
def largura_total(tabela: Table, colunas: int, pesos: list[float] | None = None) -> None:
    """Ajusta a tabela à largura da caixa, distribuindo as colunas.

    `pesos` permite dar mais espaço às colunas de texto corrido; sem ele as
    colunas ficam iguais. A grade (`w:tblGrid`) é reescrita junto das células —
    sem isso o Word recebe uma grade de uma coluna só para linhas de várias
    células e refaz a largura por conta própria, espremendo o conteúdo.
    """
    tblpr = tabela._tbl.tblPr
    largura = tblpr.find(qn("w:tblW"))
    if largura is None:
        largura = OxmlElement("w:tblW")
        tblpr.append(largura)
    largura.set(qn("w:type"), "pct")
    largura.set(qn("w:w"), "5000")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    if not pesos or len(pesos) != colunas:
        pesos = [1.0] * colunas
    total = sum(pesos)
    larguras = [LARGURA_CAIXA_IN * peso / total for peso in pesos]

    grade = tabela._tbl.find(qn("w:tblGrid"))
    if grade is not None:
        tabela._tbl.remove(grade)
    grade = OxmlElement("w:tblGrid")
    for medida in larguras:
        coluna = OxmlElement("w:gridCol")
        coluna.set(qn("w:w"), str(int(medida * 1440)))
        grade.append(coluna)
    tabela._tbl.insert(list(tabela._tbl).index(tblpr) + 1, grade)

    for linha in tabela.rows:
        for indice, celula in enumerate(linha.cells):
            celula.width = Inches(larguras[min(indice, colunas - 1)])


def adicionar_imagem(container, alternativo: str, origem: str, fonte: str) -> None:
    """Insere a figura centralizada, com o texto alternativo de acessibilidade.

    O DOCX não aceita SVG: o Markdown aponta para o PNG gerado por
    `figuras_unidade1.py`, e o SVG correspondente fica ao lado como fonte
    editável. Quando o arquivo não existe, o parágrafo denuncia a ausência em
    vez de gerar um documento silenciosamente incompleto.
    """
    caminho = Path(origem)
    if not caminho.is_absolute():
        caminho = (BASE_IMAGENS / origem).resolve()

    paragrafo = container.add_paragraph()
    formatar_paragrafo(paragrafo)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_before = Pt(8)

    if not caminho.exists():
        run = paragrafo.add_run(f"[imagem ausente: {origem}]")
        run.italic = True
        formatar_run(run, fonte)
        return

    run = paragrafo.add_run()
    run.add_picture(str(caminho), width=Inches(LARGURA_IMAGEM_IN))
    for descricao in run._element.xpath(".//wp:docPr"):
        descricao.set("descr", alternativo)
        descricao.set("title", alternativo[:255])


def adicionar_legenda(container, texto: str, fonte: str) -> None:
    paragrafo = container.add_paragraph()
    formatar_paragrafo(paragrafo)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_after = Pt(8)
    adicionar_inline(paragrafo, texto, fonte)
    for run in paragrafo.runs:
        formatar_run(run, fonte, tamanho=Pt(10))
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def adicionar_tabela(container, linhas: list[str], fonte: str) -> None:
    dados = []
    for indice, linha in enumerate(linhas):
        if indice == 1:
            continue  # separador |---|
        dados.append([c.strip() for c in linha.strip().strip("|").split("|")])
    if not dados:
        return
    colunas = max(len(linha) for linha in dados)
    tabela = container.add_table(rows=len(dados), cols=colunas)
    try:
        tabela.style = "Table Grid"
    except KeyError:
        pass
    # colunas de texto corrido recebem mais espaço que colunas numéricas
    pesos = []
    for j in range(colunas):
        maior = max((len(linha[j]) for linha in dados if j < len(linha)), default=1)
        pesos.append(float(min(max(maior, 8), 46)))
    largura_total(tabela, colunas, pesos)
    for i, linha in enumerate(dados):
        for j in range(colunas):
            celula = tabela.rows[i].cells[j]
            celula.text = ""
            paragrafo = celula.paragraphs[0]
            formatar_paragrafo(paragrafo)
            adicionar_inline(paragrafo, linha[j] if j < len(linha) else "", fonte)
            if i == 0:
                for run in paragrafo.runs:
                    run.bold = True
    container.add_paragraph("")


ALTERNATIVA = re.compile(r"^(\*?)([a-eA-E])\.\s+(.*)$")


def renderizar(
    container,
    linhas,
    fonte: str = FONTE_UNIDADE,
    limpar: bool = False,
    alternativas: bool = False,
) -> None:
    """Escreve `linhas` de Markdown dentro de uma célula ou do corpo do documento.

    Com `alternativas=True`, linhas no formato `*a. …` preservam e destacam o
    asterisco que marca a alternativa correta, conforme o modelo institucional.
    """
    if limpar and isinstance(container, _Cell):
        limpar_celula(container)

    indice, total = 0, len(linhas)
    while indice < total:
        linha = linhas[indice]
        texto = linha.strip()
        if not texto:
            indice += 1
            continue

        titulo = re.match(r"^(#{1,6})\s+(.*)$", texto)
        if titulo:
            nivel = min(max(len(titulo.group(1)), 2), 4)
            paragrafo = container.add_paragraph(style=f"Heading {nivel}")
            adicionar_inline(paragrafo, titulo.group(2), fonte)
            indice += 1
            continue

        if re.match(r"^-{3,}$", texto):
            indice += 1
            continue

        if texto.startswith("```"):
            indice += 1
            codigo = []
            while indice < total and not linhas[indice].strip().startswith("```"):
                codigo.append(linhas[indice])
                indice += 1
            indice += 1
            paragrafo = container.add_paragraph()
            formatar_paragrafo(paragrafo)
            sombrear(paragrafo, "F2F2F2")
            for posicao, conteudo in enumerate(codigo):
                run = paragrafo.add_run(conteudo)
                formatar_run(run, fonte, monoespacada=True)
                if posicao < len(codigo) - 1:
                    run.add_break()
            continue

        if texto == "$$":
            indice += 1
            expressao = []
            while indice < total and linhas[indice].strip() != "$$":
                expressao.append(linhas[indice])
                indice += 1
            indice += 1
            paragrafo = container.add_paragraph()
            formatar_paragrafo(paragrafo)
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragrafo.add_run(limpar_formula("\n".join(expressao)))
            run.italic = True
            formatar_run(run, fonte)
            continue

        if texto.startswith(">"):
            citacao = []
            while indice < total and (
                linhas[indice].strip().startswith(">") or linhas[indice].strip() == ""
            ):
                if linhas[indice].strip() == "":
                    proxima = indice + 1 < total and linhas[indice + 1].strip().startswith(">")
                    indice += 1
                    if proxima:
                        citacao.append("")
                        continue
                    break
                citacao.append(re.sub(r"^\s*>\s?", "", linhas[indice]))
                indice += 1
            for conteudo in citacao:
                if not conteudo.strip():
                    continue
                paragrafo = container.add_paragraph()
                formatar_paragrafo(paragrafo)
                paragrafo.paragraph_format.left_indent = Inches(0.3)
                paragrafo.paragraph_format.right_indent = Inches(0.2)
                adicionar_inline(paragrafo, conteudo.strip(), fonte)
                for run in paragrafo.runs:
                    run.italic = True
            continue

        if texto.startswith("|") and indice + 1 < total and re.match(
            r"^\s*\|?[\s:|-]+\|?\s*$", linhas[indice + 1]
        ):
            bloco = []
            while indice < total and linhas[indice].strip().startswith("|"):
                bloco.append(linhas[indice].strip())
                indice += 1
            adicionar_tabela(container, bloco, fonte)
            continue

        imagem = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", texto)
        if imagem:
            adicionar_imagem(container, imagem.group(1), imagem.group(2), fonte)
            indice += 1
            # a legenda é a citação imediatamente seguinte, iniciada por "Figura"
            proxima = indice
            while proxima < total and not linhas[proxima].strip():
                proxima += 1
            if proxima < total and re.match(
                r"^>\s*\*\*Figura", linhas[proxima].strip()
            ):
                legenda = []
                indice = proxima
                while indice < total and linhas[indice].strip().startswith(">"):
                    legenda.append(re.sub(r"^\s*>\s?", "", linhas[indice]).strip())
                    indice += 1
                adicionar_legenda(container, " ".join(legenda), fonte)
            continue

        if alternativas:
            achado = ALTERNATIVA.match(texto)
            if achado:
                estrela, letra, resto = achado.groups()
                paragrafo = container.add_paragraph()
                formatar_paragrafo(paragrafo)
                if estrela:
                    marca = paragrafo.add_run("*")
                    marca.bold = True
                    marca.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    formatar_run(marca, fonte)
                    rotulo = paragrafo.add_run(f"{letra}. ")
                    rotulo.bold = True
                    formatar_run(rotulo, fonte)
                else:
                    formatar_run(paragrafo.add_run(f"{letra}. "), fonte)
                adicionar_inline(paragrafo, resto, fonte)
                indice += 1
                continue

        if re.match(r"^[-*]\s+", texto) or re.match(r"^\d+\.\s+", texto):
            while indice < total and (
                re.match(r"^\s*[-*]\s+", linhas[indice])
                or re.match(r"^\s*\d+\.\s+", linhas[indice])
                or (linhas[indice].strip() and linhas[indice].startswith("  "))
            ):
                # blocos aninhados (fórmula, código, tabela) saem da lista
                aninhado = linhas[indice].strip()
                if aninhado in ("$$",) or aninhado.startswith(("```", "|")):
                    break
                bruto = linhas[indice]
                recuo = len(bruto) - len(bruto.lstrip(" "))
                item = bruto.strip()
                paragrafo = container.add_paragraph()
                formatar_paragrafo(paragrafo)
                paragrafo.paragraph_format.left_indent = Inches(
                    0.25 + (0.25 if recuo >= 2 else 0)
                )
                numerado = re.match(r"^(\d+)\.\s+(.*)$", item)
                marcador = re.match(r"^[-*]\s+(.*)$", item)
                if marcador:
                    formatar_run(paragrafo.add_run("•  "), fonte)
                    adicionar_inline(paragrafo, marcador.group(1), fonte)
                elif numerado:
                    formatar_run(paragrafo.add_run(f"{numerado.group(1)}.  "), fonte)
                    adicionar_inline(paragrafo, numerado.group(2), fonte)
                else:
                    adicionar_inline(paragrafo, item, fonte)
                indice += 1
            continue

        paragrafo = container.add_paragraph()
        formatar_paragrafo(paragrafo)
        adicionar_inline(paragrafo, texto, fonte)
        indice += 1

    if isinstance(container, _Cell):
        garantir_paragrafo_final(container)


# ============================================================ leitura do markdown
def secoes(markdown: str, nivel: int = 2) -> list[tuple[str, list[str]]]:
    """Divide o texto em (título, linhas) pelos cabeçalhos do nível indicado."""
    marca = "#" * nivel
    padrao = re.compile(rf"^{marca}\s+(?!#)(.*)$")
    resultado: list[tuple[str, list[str]]] = []
    titulo: str | None = None
    corpo: list[str] = []
    for linha in markdown.split("\n"):
        achado = padrao.match(linha)
        if achado:
            if titulo is not None:
                resultado.append((titulo, corpo))
            titulo, corpo = achado.group(1).strip(), []
        elif titulo is not None:
            corpo.append(linha)
    if titulo is not None:
        resultado.append((titulo, corpo))
    return resultado


def sem_secao(linhas: list[str], titulos: list[str]) -> list[str]:
    """Remove subseções `### <título>` (e o respectivo corpo) de um bloco."""
    alvos = {t.lower() for t in titulos}
    saida, pulando = [], False
    for linha in linhas:
        achado = re.match(r"^###\s+(.*)$", linha.strip())
        if achado:
            pulando = any(achado.group(1).strip().lower().startswith(a) for a in alvos)
        if not pulando:
            saida.append(linha)
    return saida


def carregar(caminho: Path) -> str:
    return caminho.read_text(encoding="utf-8")
