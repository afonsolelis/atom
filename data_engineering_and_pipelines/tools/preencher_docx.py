#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Preenche o TEMPLATE Word (Atomo 3.0 / UniFECAF) de uma unidade com o conteudo do
markdown correspondente (unidade_N.md), colocando cada parte na "caixa" certa:
intro, TEXTO BASE AULA 1-4, ROTEIRO VIDEOAULA 1-4, QUIZ, AAI e material complementar.

- Formulas ($..$ / $$..$$) entram como IMAGEM (reusa o renderizador MathJax de
  montar_para_google_docs.py e o cache tools/_cache_formulas/).
- Negrito, italico, codigo, listas, tabelas, blockquotes e imagens sao preservados.
- Os textos de orientacao do template sao removidos (dentro e fora das caixas).
- Trabalha sempre a partir de um backup "pristine" do template, entao pode ser
  re-executado a vontade.

Uso:
    python tools/preencher_docx.py 1        # preenche a Unidade 1
    python tools/preencher_docx.py 1 2 3 4  # varias unidades

O .docx de saida e o proprio arquivo na Area de Trabalho:
    C:\\Users\\<user>\\Desktop\\Unidade N\\TEMPLATE - Unidade N_Data Engineering and Pipelines.docx
"""
import sys
import os
import re
import io
import shutil
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

TOOLS = Path(__file__).resolve().parent
DISC = TOOLS.parent
sys.path.insert(0, str(TOOLS))
import montar_para_google_docs as mg   # reusa render de formulas / fetch / rasterize

DISCIPLINA = "Data Engineering and Pipelines"
CONTEUDISTA = "Afonso Cesar Lelis Brandão"
DESKTOP = Path(os.path.expanduser("~")) / "Desktop"
PRISTINE = TOOLS / "_templates_pristine"


# ============================================================ formulas (cache)
def build_formula_index(md_text):
    formulas = mg.collect_formulas(md_text)          # [(latex, display), ...]
    uniq = {}
    for latex, disp in formulas:
        uniq[(latex.strip(), disp)] = None
    return mg.render_all_formulas(list(uniq.keys())) if uniq else {}


# ============================================================ helpers docx
def clear_container(cell):
    # remove TUDO menos as propriedades da celula (w:tcPr): paragrafos, tabelas
    # aninhadas e tambem content-controls w:sdt (onde o template esconde tabelas
    # de orientacao como a "Texto Dialogico vs Academico").
    tc = cell._tc
    # destrava a linha: remove cantSplit/tblHeader do <w:tr> pai. Sem isso, como a
    # linha de conteudo e enorme, o Word a joga inteira para a proxima pagina e
    # deixa uma pagina quase em branco depois do cabecalho da caixa.
    tr = tc.getparent()
    if tr is not None and tr.tag == qn("w:tr"):
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None:
            for tag in ("w:cantSplit", "w:tblHeader"):
                el = trPr.find(qn(tag))
                if el is not None:
                    trPr.remove(el)
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)


def ensure_trailing_p(cell):
    tc = cell._tc
    if len(tc) == 0 or tc[-1].tag != qn("w:p"):
        cell.add_paragraph("")


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text; r.append(t)
    hl.append(r); paragraph._p.append(hl)


INLINE_RE = re.compile(
    r"\*\*(?P<b>.+?)\*\*"
    r"|`(?P<c>[^`]+)`"
    r"|\[(?P<lt>[^\]]+)\]\((?P<lu>[^)]+)\)"
    r"|(?<![\*\w])\*(?P<i>[^*]+?)\*(?![\*\w])"
)


def add_md_runs(p, text):
    """Adiciona runs com **negrito**, `codigo`, [link](url), *italico*."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        if m.group("b") is not None:
            p.add_run(m.group("b")).bold = True
        elif m.group("c") is not None:
            r = p.add_run(m.group("c")); r.font.name = "Consolas"; r.font.size = Pt(10.5)
        elif m.group("lt") is not None:
            add_hyperlink(p, m.group("lu"), m.group("lt"))
        elif m.group("i") is not None:
            p.add_run(m.group("i")).italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def add_inline(p, text, formula_pngs):
    """Como add_md_runs, mas tratando $formula$ inline como imagem."""
    for kind, content in mg.iter_math(text):
        if kind == "text":
            add_md_runs(p, content.replace(r"\$", "$"))
        elif kind == "inline":
            png = formula_pngs.get((content.strip(), False))
            if png:
                r = p.add_run()
                r.add_picture(io.BytesIO(png), height=Pt(13))
            else:
                add_md_runs(p, "$%s$" % content)
        # 'display' nao ocorre no meio de uma linha


def get_image_bytes(src, base_dir):
    if src.startswith(("http://", "https://")):
        data, mime = mg.fetch_remote(src)
        if "svg" in mime:
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".svg", delete=False) as tf:
                tf.write(data); tmp = Path(tf.name)
            try:
                data = mg.rasterize_svg(tmp); mime = "image/png"
            finally:
                tmp.unlink(missing_ok=True)
        data, mime = mg.downscale(data, mime)
        return data
    p = (base_dir / src).resolve()
    if not p.exists():
        return None
    if p.suffix.lower() == ".svg":
        return mg.rasterize_svg(p)
    return p.read_bytes()


# ============================================================ renderer de blocos markdown
def _img_size_inch(png, dpi_factor):
    w, h = Image.open(io.BytesIO(png)).size
    return w / dpi_factor, h / dpi_factor


def render_blocks(cell, lines, formula_pngs, base_dir, first=False, alts=False):
    """Renderiza linhas de markdown dentro de `cell` (celula OU o corpo do doc).
    Se first=True, limpa a celula. Se alts=True, trata linhas de alternativa
    '*a.'/'a.' preservando e destacando o '*' da alternativa correta."""
    if first:
        clear_container(cell)
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if s == "":
            i += 1; continue

        # --- heading
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            style = "Heading %d" % min(max(level, 2), 4)
            p = cell.add_paragraph(style=style)
            add_inline(p, m.group(2), formula_pngs)
            i += 1; continue

        # --- horizontal rule
        if re.match(r"^-{3,}$", s):
            i += 1; continue

        # --- fenced code
        if s.startswith("```"):
            i += 1; code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            p = cell.add_paragraph()
            shade_paragraph(p, "F2F2F2")
            for k, cl in enumerate(code):
                r = p.add_run(cl); r.font.name = "Consolas"; r.font.size = Pt(10)
                if k < len(code) - 1:
                    r.add_break()
            continue

        # --- display math ($$ ... $$)
        if s == "$$":
            i += 1; expr = []
            while i < n and lines[i].strip() != "$$":
                expr.append(lines[i]); i += 1
            i += 1
            latex = "\n".join(expr).strip()
            png = formula_pngs.get((latex, True))
            p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if png:
                wi, hi = _img_size_inch(png, 192)
                p.add_run().add_picture(io.BytesIO(png), width=Inches(min(wi, 5.3)))
            else:
                add_md_runs(p, latex)
            continue

        # --- blockquote
        if s.startswith(">"):
            quote = []
            while i < n and (lines[i].strip().startswith(">") or lines[i].strip() == ""):
                if lines[i].strip() == "":
                    if i + 1 < n and lines[i + 1].strip().startswith(">"):
                        quote.append(""); i += 1; continue
                    break
                quote.append(re.sub(r"^\s*>\s?", "", lines[i])); i += 1
            # renderiza o conteudo do blockquote recursivamente, com indentacao/italico
            sub = [ql for ql in quote]
            para_buf = []
            def flush_quote():
                if para_buf:
                    p = cell.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.right_indent = Inches(0.2)
                    for run_text in [" ".join(para_buf)]:
                        add_inline(p, run_text, formula_pngs)
                    for r in p.runs:
                        r.italic = True
                    para_buf.clear()
            for ql in sub:
                qs = ql.strip()
                if qs == "":
                    flush_quote()
                elif re.match(r"^[-*]\s+", qs) or re.match(r"^\d+\.\s+", qs):
                    flush_quote()
                    p = cell.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    txt = re.sub(r"^[-*]\s+", "", qs); txt = re.sub(r"^\d+\.\s+", "", txt)
                    prefix = "•  " if re.match(r"^[-*]\s+", qs) else (qs.split(".")[0] + ".  ")
                    p.add_run(prefix)
                    add_inline(p, txt, formula_pngs)
                    for r in p.runs:
                        r.italic = True
                else:
                    para_buf.append(qs)
            flush_quote()
            continue

        # --- table
        if s.startswith("|") and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip()); i += 1
            add_markdown_table(cell, tbl_lines, formula_pngs)
            continue

        # --- standalone image
        mi = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", s)
        if mi:
            alt, src = mi.group(1), mi.group(2)
            data = None
            try:
                data = get_image_bytes(src, base_dir)
            except Exception as e:
                data = None
            p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if data:
                wi, hi = _img_size_inch(data, 96)
                p.add_run().add_picture(io.BytesIO(data), width=Inches(min(wi, 5.3)))
                if alt:
                    cap = cell.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rc = cap.add_run(alt); rc.italic = True; rc.font.size = Pt(9)
            else:
                r = p.add_run("[imagem: %s]" % alt); r.italic = True
            i += 1; continue

        # --- alternativa de questao (*a. / a.) — preserva o '*' da correta
        if alts:
            malt = re.match(r"^(\*?)([a-eA-E])\.\s+(.*)$", s)
            if malt:
                star, letter, rest = malt.group(1), malt.group(2), malt.group(3)
                p = cell.add_paragraph()
                if star:                      # alternativa CORRETA: '*' em destaque
                    rs = p.add_run("* "); rs.bold = True
                    rs.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    rl = p.add_run(letter + ". "); rl.bold = True
                else:
                    p.add_run(letter + ". ")
                add_inline(p, rest, formula_pngs)
                i += 1; continue

        # --- list (bullets / numbered), possibly multi-line
        if re.match(r"^[-*]\s+", s) or re.match(r"^\d+\.\s+", s):
            while i < n and (re.match(r"^\s*[-*]\s+", lines[i]) or re.match(r"^\s*\d+\.\s+", lines[i])
                             or (lines[i].strip() and lines[i].startswith("  ") )):
                raw = lines[i]
                indent = len(raw) - len(raw.lstrip(" "))
                item = raw.strip()
                mnum = re.match(r"^(\d+)\.\s+(.*)$", item)
                mbul = re.match(r"^[-*]\s+(.*)$", item)
                p = cell.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25 + (0.25 if indent >= 2 else 0))
                if mbul:
                    p.add_run("•  ")
                    add_inline(p, mbul.group(1), formula_pngs)
                elif mnum:
                    p.add_run(mnum.group(1) + ".  ")
                    add_inline(p, mnum.group(2), formula_pngs)
                else:
                    add_inline(p, item, formula_pngs)
                i += 1
            continue

        # --- normal paragraph
        p = cell.add_paragraph()
        add_inline(p, s, formula_pngs)
        i += 1

    if hasattr(cell, "_tc"):      # so celulas exigem paragrafo final; o corpo nao
        ensure_trailing_p(cell)


def shade_paragraph(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def set_table_full_width(t, ncols):
    """Faz a tabela ocupar 100% da largura da celula e distribuir colunas
    (senao o Word colapsa as colunas para ~1 caractere)."""
    tblPr = t._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct"); tblW.set(qn("w:w"), "5000")   # 100%
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); tblPr.append(layout)
    layout.set(qn("w:type"), "autofit")
    # largura-guia por coluna (fallback caso o Word use fixed layout)
    col_w = Inches(6.0 / ncols)
    for row in t.rows:
        for c in row.cells:
            c.width = col_w


def add_markdown_table(cell, tbl_lines, formula_pngs):
    rows = []
    for k, ln in enumerate(tbl_lines):
        if k == 1:
            continue  # separador |---|
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = cell.add_table(rows=len(rows), cols=ncols)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    set_table_full_width(t, ncols)
    for ri, r in enumerate(rows):
        for ci in range(ncols):
            tc = t.rows[ri].cells[ci]
            tc.text = ""
            p = tc.paragraphs[0]
            val = r[ci] if ci < len(r) else ""
            add_inline(p, val, formula_pngs)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    cell.add_paragraph("")


# ============================================================ localizar caixas
def iter_tables(doc):
    out = []
    for block in doc.iter_inner_content():
        if isinstance(block, Table):
            out.append(block)
    return out


def table_header(t):
    try:
        return t.rows[0].cells[0].text.strip()
    except Exception:
        return ""


def _box_header(el):
    tr = el.find(qn("w:tr"))
    if tr is None:
        return None
    tc = tr.find(qn("w:tc"))
    if tc is None:
        return None
    return "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip().lower()


def find_table(doc, needle):
    # busca em TODAS as w:tbl do corpo (inclusive as embrulhadas em content-controls
    # w:sdt — as unidades 3 e 4 escondem quase todas as caixas dentro de sdt, que o
    # iter_inner_content nao enxerga). Casa por substring no texto da 1a celula.
    needle = needle.lower()
    for el in doc.element.body.iter(qn("w:tbl")):
        h = _box_header(el)
        if h and needle in h:
            return Table(el, doc)
    return None


def find_box_num(doc, prefix, number):
    """Acha a caixa cujo cabecalho e '<prefix> ... <number>' com numero EXATO
    (nao seguido de outro digito — senao 'aula 1' casaria com 'aula 10')."""
    pat = re.compile(r"%s\s*0*%d(?!\d)" % (re.escape(prefix.lower()), number))
    for el in doc.element.body.iter(qn("w:tbl")):
        h = _box_header(el)
        if h and pat.search(h):
            return Table(el, doc)
    return None


# ============================================================ parser das secoes do md
def parse_sections(md_text):
    lines = md_text.split("\n")
    sections = []
    title, buf = None, []
    for ln in lines:
        if re.match(r"^##\s+", ln):
            if title is not None:
                sections.append((title, buf))
            title = ln[2:].strip(); buf = []
        else:
            if title is not None:
                buf.append(ln)
    if title is not None:
        sections.append((title, buf))
    return sections


def classify(title):
    t = title.lower()
    if t.startswith("vídeo introdutório") or t.startswith("video introdutório") or "relação da disciplina" in t:
        return ("intro", None)
    m = re.match(r"aula\s+(\d+)\s*[—-].*roteiro", t)
    if m:
        return ("roteiro", int(m.group(1)))
    m = re.match(r"aula\s+(\d+)\s*[—-]", t)
    if m:
        return ("texto", int(m.group(1)))
    if "quiz" in t:
        return ("quiz", None)
    if "atividade verificadora" in t or "aai" in t:
        return ("aai", None)
    if "material complementar" in t:
        return ("material", None)
    return ("outro", None)


def split_subsections(lines):
    """Divide um corpo em subsecoes por '### '. Retorna [(subtitulo|None, linhas)]."""
    out = []
    cur_t, cur = None, []
    for ln in lines:
        if re.match(r"^###\s+", ln):
            out.append((cur_t, cur)); cur_t = ln[3:].strip(); cur = []
        else:
            cur.append(ln)
    out.append((cur_t, cur))
    return out


# ============================================================ preenchimento
def fill_unit(nnum):
    md_path = DISC / ("unidade_%d/unidade_%d.md" % (nnum, nnum))
    if not md_path.exists():
        print("  ! markdown nao encontrado:", md_path); return
    md_text = md_path.read_text(encoding="utf-8")
    base_dir = md_path.parent

    desk_folder = DESKTOP / ("Unidade %d" % nnum)
    docx_name = "TEMPLATE - Unidade %d_%s.docx" % (nnum, DISCIPLINA)
    docx_path = desk_folder / docx_name
    if not docx_path.exists():
        print("  ! docx nao encontrado:", docx_path); return

    # backup pristine (fonte re-executavel)
    PRISTINE.mkdir(parents=True, exist_ok=True)
    pristine = PRISTINE / docx_name
    if not pristine.exists():
        shutil.copy2(docx_path, pristine)
        print("  backup pristine criado:", pristine.name)

    print("  renderizando formulas da unidade...")
    formula_pngs = build_formula_index(md_text)

    doc = Document(pristine)   # sempre parte do template limpo
    sections = parse_sections(md_text)

    # 1) placeholders disciplina / conteudista
    for p in doc.paragraphs:
        for run in p.runs:
            if "XXXX" in run.text:
                low = p.text.lower()
                if "disciplina" in low:
                    run.text = run.text.replace("XXXX", DISCIPLINA)
                elif "conteudista" in low:
                    run.text = run.text.replace("XXXX", CONTEUDISTA)

    # 2) cada secao -> sua caixa
    filled = []
    for title, body in sections:
        kind, num = classify(title)
        if kind == "intro":
            t = find_table(doc, "relação da disciplina")
            if t:
                cell = t.rows[1].cells[0]
                render_blocks(cell, body, formula_pngs, base_dir, first=True)
                filled.append("intro")
        elif kind == "texto":
            # templates inconsistentes: U1/U2 numeram as caixas por-unidade (1-4);
            # U3/U4 usam numeracao continua (9-12, 13-16). Tenta as duas.
            per_unit = num - 4 * (nnum - 1)
            t = find_box_num(doc, "texto base aula", per_unit) or \
                find_box_num(doc, "texto base aula", num)
            if t:
                cell = t.rows[1].cells[0]
                header = ["### " + title, ""]           # titulo real como subheading
                render_blocks(cell, header + body, formula_pngs, base_dir, first=True)
                filled.append("texto%d" % num)
        elif kind == "roteiro":
            per_unit = num - 4 * (nnum - 1)
            t = find_box_num(doc, "roteiro videoaula", per_unit) or \
                find_box_num(doc, "roteiro videoaula", num)
            if t:
                cell = t.rows[1].cells[0]
                render_blocks(cell, body, formula_pngs, base_dir, first=True)
                filled.append("roteiro%d" % num)
        elif kind == "quiz":
            fill_quiz(doc, body, formula_pngs, base_dir); filled.append("quiz")
        elif kind == "aai":
            fill_aai(doc, body, formula_pngs, base_dir); filled.append("aai")
        elif kind == "material":
            fill_material(doc, body, formula_pngs, base_dir); filled.append("material")

    # 2b) caixa "intro" (video introdutorio + relacao da disciplina) so existe na
    #     Unidade 1. Nas demais, o template ainda traz a caixa (com orientacao) —
    #     como nao ha secao correspondente no md, removemos a caixa inteira.
    if "intro" not in filled:
        t = find_table(doc, "relação da disciplina")
        if t is not None:
            t._tbl.getparent().remove(t._tbl)
            filled.append("intro-removida")

    # 3) remover orientacoes soltas (paragrafos body-level normais nao-essenciais)
    removed = remove_orientations(doc)

    # 4) destravar TODAS as linhas (cantSplit/tblHeader) — evita paginas em branco
    #    apos o cabecalho de cada caixa quando o conteudo e longo.
    for tbl in doc.element.body.iter(qn("w:tbl")):
        for tr in tbl.findall(qn("w:tr")):
            trPr = tr.find(qn("w:trPr"))
            if trPr is not None:
                for tag in ("w:cantSplit", "w:tblHeader"):
                    el = trPr.find(qn(tag))
                    if el is not None:
                        trPr.remove(el)
    # 4b) remover quebras/"manter junto" herdadas que empurram a caixa para a
    #     pagina seguinte (pageBreakBefore/keepNext/keepLines).
    for p in doc.element.body.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            for tag in ("w:pageBreakBefore", "w:keepNext", "w:keepLines"):
                el = pPr.find(qn(tag))
                if el is not None:
                    pPr.remove(el)

    doc.save(docx_path)
    print("  OK ->", docx_path.name)
    print("     caixas preenchidas:", ", ".join(filled))
    print("     paragrafos de orientacao removidos:", removed)


def fill_quiz(doc, body, formula_pngs, base_dir):
    t = find_table(doc, "quiz")
    if not t:
        return
    subs = split_subsections(body)   # [(None, pre), ('Questão 1', ...), ('Questão 2', ...)]
    perguntas, respostas = [], []
    for sub_t, sub_lines in subs:
        if not sub_t:
            continue
        # separa a parte de pergunta/alternativas da parte de resposta/feedback
        q_part, a_part, in_ans = [], [], False
        for ln in sub_lines:
            if re.match(r"^\*\*Resposta correta", ln.strip()) or re.match(r"^\*\*Feedback", ln.strip()):
                in_ans = True
            (a_part if in_ans else q_part).append(ln)
        perguntas.append("### " + sub_t); perguntas += q_part
        respostas.append("### " + sub_t); respostas += a_part
    # linha 1 = perguntas ; linha 2 = respostas/feedback
    render_blocks(t.rows[1].cells[0], perguntas, formula_pngs, base_dir, first=True)
    if len(t.rows) > 2:
        render_blocks(t.rows[2].cells[0], respostas, formula_pngs, base_dir, first=True)


def fill_aai(doc, body, formula_pngs, base_dir):
    t = find_table(doc, "aai")
    if not t:
        t = find_table(doc, "atividade avaliativa")
    if not t:
        return
    render_blocks(t.rows[1].cells[0], body, formula_pngs, base_dir, first=True)


def fill_material(doc, body, formula_pngs, base_dir):
    subs = [s for s in split_subsections(body) if s[0]]   # 4 subsecoes com titulo
    targets = ["direto da fonte", "para mergulhar", "podcast", "artigo"]
    for (sub_t, sub_lines), needle in zip(subs, targets):
        t = find_table(doc, needle)
        if not t:
            continue
        # coloca tudo (texto provocativo + campos) na primeira celula de conteudo
        cell = t.rows[1].cells[0]
        render_blocks(cell, sub_lines, formula_pngs, base_dir, first=True)
        # DELETA as demais linhas (orientacao) — nao basta limpar, senao ficam
        # linhas em branco com borda.
        for ri in range(len(t.rows) - 1, 1, -1):
            tr = t.rows[ri]._tr
            tr.getparent().remove(tr)


def remove_orientations(doc):
    keep_prefix = ("nome da disciplina", "conteudista", "template de produção")
    removed = 0
    for p in list(doc.paragraphs):
        style = p.style.name if p.style else ""
        txt = p.text.strip()
        if not txt:
            continue
        if style.lower().startswith("heading"):
            continue
        if txt.lower().startswith(keep_prefix):
            continue
        # paragrafo body-level, normal, nao-essencial => orientacao
        p._p.getparent().remove(p._p)
        removed += 1
    return removed


def main():
    nums = [int(a) for a in sys.argv[1:]] or [1]
    print("Disciplina:", DISCIPLINA)
    print("Navegador (formulas):", mg.BROWSER or "NAO ENCONTRADO")
    print("-" * 60)
    for nnum in nums:
        print("Unidade %d:" % nnum)
        fill_unit(nnum)
    print("-" * 60)
    print("Pronto. Abra o .docx na Area de Trabalho para conferir.")


if __name__ == "__main__":
    main()
